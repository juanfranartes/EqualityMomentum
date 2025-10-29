"""
Generador de Informes - Versión Web
Genera informes Word con gráficos en memoria (BytesIO)
"""

import io
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .utils import formato_numero_es, calcular_brecha, reformatear_etiqueta_escala


# Configuración de colores
COLORES = {
    'mujer': '#1e4389',      # Azul
    'hombre': '#ea5d41',     # Rojo/Naranja
    'neutro': '#000000'
}

COLORES_RGB = {
    'mujer': RGBColor(30, 67, 137),
    'hombre': RGBColor(234, 93, 65),
    'neutro': RGBColor(0, 0, 0)
}

# Mapeo de columnas (sin espacios al final)
COLS = {
    'reg': 'Reg.',
    'sexo': 'Sexo',
    'grupo_prof': 'Grupo profesional',
    'nivel_svpt': 'Nivel SVPT',
    'puesto': 'Puesto de trabajo',
    'nivel_convenio': 'Nivel Convenio Colectivo',
    'sb_efectivo': 'Salario base efectivo Total',
    'sbc_efectivo': 'Salario base anual + complementos Total',
    'sbce_efectivo': 'Salario base anual + complementos + Extrasalariales Total',
    'sb_equiparado': 'salario_base_equiparado',
    'sbc_equiparado': 'sb_mas_comp_salariales_equiparado',
    'sbce_equiparado': 'sb_mas_comp_total_equiparado',
    'comp_efectivo': 'Compltos Salariales efectivo Total',
    'comp_equiparado': 'complementos_salariales_equiparados',
    'extra_efectivo': 'Compltos Extrasalariales efectivo Total',
    'extra_equiparado': 'complementos_extrasalariales_equiparados',
}

# Configurar estilo matplotlib
plt.rcParams['font.size'] = 14
plt.rcParams['font.family'] = 'sans-serif'
sns.set_style("whitegrid")


class GeneradorInformes:
    """Genera informes Word desde Excel procesado"""

    def __init__(self, plantilla_path=None):
        """
        Inicializa el generador

        Args:
            plantilla_path: Ruta a la plantilla Word (opcional)
        """
        self.plantilla_path = plantilla_path

    def generar_informe_completo(self, excel_bytes):
        """
        Genera un informe Word completo desde un Excel procesado

        Args:
            excel_bytes: BytesIO o bytes del Excel procesado

        Returns:
            BytesIO con el documento Word generado
        """
        # Leer Excel
        if isinstance(excel_bytes, bytes):
            excel_bytes = io.BytesIO(excel_bytes)

        excel_bytes.seek(0)
        df = pd.read_excel(excel_bytes, sheet_name='DATOS_PROCESADOS')

        # LIMPIAR ESPACIOS EN NOMBRES DE COLUMNAS
        df.columns = df.columns.str.strip()

        # Crear documento (desde plantilla o en blanco)
        if self.plantilla_path:
            doc = Document(self.plantilla_path)
        else:
            doc = Document()
            doc.add_heading('INFORME DE REGISTRO RETRIBUTIVO', 0)

        # Filtrar datos actuales (excluir 'Ex')
        df_actual = df[df[COLS['reg']] != 'Ex'].copy()

        # Verificar si hay datos suficientes
        if len(df_actual) == 0:
            doc.add_paragraph("No hay datos disponibles para generar el informe.")
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            return output

        # Convertir 'Mujeres' y 'Hombres' a 'M' y 'H' para cálculos
        df_actual.loc[:, COLS['sexo']] = df_actual[COLS['sexo']].map({
            'Mujeres': 'M',
            'Hombres': 'H',
            'M': 'M',
            'H': 'H'
        })

        # SECCIÓN 1: Resumen General (Promedio)
        doc.add_heading('1. Análisis Salarial General (Promedio)', 1)

        # Calcular estadísticos generales
        stats_sb = self._calcular_estadistico(df_actual, COLS['sb_equiparado'])
        stats_sbc = self._calcular_estadistico(df_actual, COLS['sbc_equiparado'])

        if stats_sb and stats_sbc:
            # Crear gráfico donut SB
            grafico_sb = self._crear_grafico_donut(stats_sb, "Salario Base Promedio")
            if grafico_sb:
                doc.add_picture(grafico_sb, width=Inches(3.5))

            doc.add_paragraph()

            # Crear gráfico donut SB+C
            grafico_sbc = self._crear_grafico_donut(stats_sbc, "Salario Base + Complementos Promedio")
            if grafico_sbc:
                doc.add_picture(grafico_sbc, width=Inches(3.5))

        # SECCIÓN 2: Análisis por Grupo Profesional (si existe la columna)
        if COLS['grupo_prof'] in df_actual.columns:
            doc.add_page_break()
            doc.add_heading('2. Análisis por Grupo Profesional', 1)

            grupos = df_actual[COLS['grupo_prof']].dropna().unique()
            datos_grupos = []

            for grupo in grupos:
                stats_grupo = self._calcular_estadistico(
                    df_actual, COLS['sb_equiparado'],
                    grupo_col=COLS['grupo_prof'], grupo_valor=grupo
                )
                if stats_grupo and (stats_grupo['n_M'] > 1 and stats_grupo['n_H'] > 1):
                    datos_grupos.append({
                        'categoria': str(grupo),
                        'M': stats_grupo['M'],
                        'H': stats_grupo['H'],
                        'n_M': stats_grupo['n_M'],
                        'n_H': stats_grupo['n_H']
                    })

            if datos_grupos:
                grafico_grupos = self._crear_grafico_barras(
                    datos_grupos,
                    "Salario Base por Grupo Profesional",
                    "Salario Promedio (€)"
                )
                if grafico_grupos:
                    doc.add_picture(grafico_grupos, width=Inches(6.5))

        # SECCIÓN 3: Análisis por Categoría (si existe la columna)
        if COLS['nivel_convenio'] in df_actual.columns:
            doc.add_page_break()
            doc.add_heading('3. Análisis por Categoría Profesional', 1)

            categorias = df_actual[COLS['nivel_convenio']].dropna().unique()
            datos_categorias = []

            for cat in categorias:
                stats_cat = self._calcular_estadistico(
                    df_actual, COLS['sb_equiparado'],
                    grupo_col=COLS['nivel_convenio'], grupo_valor=cat
                )
                if stats_cat and (stats_cat['n_M'] > 1 and stats_cat['n_H'] > 1):
                    datos_categorias.append({
                        'categoria': str(cat),
                        'M': stats_cat['M'],
                        'H': stats_cat['H'],
                        'n_M': stats_cat['n_M'],
                        'n_H': stats_cat['n_H']
                    })

            if datos_categorias:
                grafico_cat = self._crear_grafico_barras(
                    datos_categorias,
                    "Salario Base por Categoría Profesional",
                    "Salario Promedio (€)"
                )
                if grafico_cat:
                    doc.add_picture(grafico_cat, width=Inches(6.5))

        # SECCIÓN 4: Tabla Resumen
        doc.add_page_break()
        doc.add_heading('4. Tabla Resumen - Datos Generales', 1)

        if stats_sb and stats_sbc:
            tabla = doc.add_table(rows=4, cols=4)
            tabla.style = 'Light Grid Accent 1'

            # Encabezados
            headers = ['Concepto', 'Mujeres', 'Hombres', 'Brecha %']
            for idx, header in enumerate(headers):
                cell = tabla.rows[0].cells[idx]
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True

            # Fila 1: N empleados
            tabla.rows[1].cells[0].text = 'N° Empleados'
            tabla.rows[1].cells[1].text = str(stats_sb.get('n_M', 0))
            tabla.rows[1].cells[2].text = str(stats_sb.get('n_H', 0))
            tabla.rows[1].cells[3].text = '-'

            # Fila 2: Salario Base
            tabla.rows[2].cells[0].text = 'Salario Base Promedio'
            tabla.rows[2].cells[1].text = formato_numero_es(stats_sb['M']) + ' €'
            tabla.rows[2].cells[2].text = formato_numero_es(stats_sb['H']) + ' €'
            tabla.rows[2].cells[3].text = formato_numero_es(stats_sb.get('brecha', 0)) + '%'

            # Fila 3: SB + Complementos
            tabla.rows[3].cells[0].text = 'SB + Complementos Promedio'
            tabla.rows[3].cells[1].text = formato_numero_es(stats_sbc['M']) + ' €'
            tabla.rows[3].cells[2].text = formato_numero_es(stats_sbc['H']) + ' €'
            tabla.rows[3].cells[3].text = formato_numero_es(stats_sbc.get('brecha', 0)) + '%'

        # Generar documento en memoria
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)

        return output

    def _calcular_estadistico(self, df, columna_salario, metodo='media', grupo_col=None, grupo_valor=None):
        """Calcula estadístico (media o mediana) por género"""
        df_filtrado = df[df[columna_salario] > 0].copy()

        if grupo_col and grupo_valor is not None:
            df_filtrado = df_filtrado[df_filtrado[grupo_col].astype(str) == str(grupo_valor)]

        if df_filtrado.empty:
            return None

        resultado = {}
        for sexo in ['M', 'H']:
            df_sexo = df_filtrado[df_filtrado[COLS['sexo']] == sexo]
            resultado[f'n_{sexo}'] = len(df_sexo)

            if len(df_sexo) > 0:
                if metodo == 'media':
                    resultado[sexo] = df_sexo[columna_salario].mean()
                else:
                    resultado[sexo] = df_sexo[columna_salario].median()
            else:
                resultado[sexo] = 0

        # Calcular brecha
        if resultado['H'] > 0:
            resultado['brecha'] = calcular_brecha(resultado['H'], resultado['M'])
        else:
            resultado['brecha'] = 0

        return resultado

    def _crear_grafico_donut(self, datos, titulo):
        """Crea un gráfico donut y retorna BytesIO"""
        valor_m = datos['M']
        valor_h = datos['H']
        brecha = datos.get('brecha', 0)

        # Filtro de privacidad
        n_m = datos.get('n_M', 0)
        n_h = datos.get('n_H', 0)
        if n_m == 1 or n_h == 1:
            return None

        if valor_m == 0 and valor_h == 0:
            return None

        fig, ax = plt.subplots(figsize=(8, 8))
        fig.patch.set_facecolor('white')

        valores = [valor_m, valor_h]
        colores = [COLORES['mujer'], COLORES['hombre']]

        wedges, _ = ax.pie(
            valores,
            labels=['', ''],
            colors=colores,
            startangle=90,
            wedgeprops=dict(width=0.5, edgecolor='white', linewidth=2)
        )

        # Añadir valores en los segmentos
        for wedge, valor in zip(wedges, valores):
            ang = (wedge.theta2 - wedge.theta1) / 2 + wedge.theta1
            radio_centro = 0.80
            x_centro = np.cos(np.deg2rad(ang)) * radio_centro
            y_centro = np.sin(np.deg2rad(ang)) * radio_centro

            ax.text(x_centro, y_centro, formato_numero_es(valor, 2) + ' €',
                    ha='center', va='center',
                    fontsize=14, fontweight='bold', color='white')

        # Hueco central
        centro = plt.Circle((0, 0), 0.55, fc='white')
        ax.add_artist(centro)

        # Brecha en el centro
        color_brecha = '#E74C3C' if brecha > 0 else '#3498DB' if brecha < 0 else '#cc0000'
        texto_brecha = f"Brecha Salarial\n{formato_numero_es(abs(brecha), 2)}%"
        ax.text(0, 0, texto_brecha, ha='center', va='center',
                fontsize=22, fontweight='bold', color=color_brecha)

        # Título
        ax.set_title(titulo, fontsize=18, fontweight='bold', color='#cc0000', pad=20)

        # Leyenda
        from matplotlib.patches import Patch
        legend_elements = [
            Patch(facecolor=COLORES['mujer'], label='Mujeres'),
            Patch(facecolor=COLORES['hombre'], label='Hombres')
        ]
        ax.legend(handles=legend_elements, loc='lower right', frameon=False, fontsize=14)

        plt.tight_layout()

        # Guardar en BytesIO
        output = io.BytesIO()
        plt.savefig(output, dpi=300, bbox_inches='tight', facecolor='white', format='png')
        plt.close()
        output.seek(0)

        return output

    def _crear_grafico_barras(self, datos_lista, titulo, tipo_valor):
        """Crea un gráfico de barras y retorna BytesIO"""
        if not datos_lista:
            return None

        # Aplicar filtro de privacidad
        datos_filtrados = [d for d in datos_lista if d['n_M'] > 1 and d['n_H'] > 1]

        if not datos_filtrados:
            return None

        categorias = [reformatear_etiqueta_escala(d['categoria']) for d in datos_filtrados]
        valores_m = [d['M'] for d in datos_filtrados]
        valores_h = [d['H'] for d in datos_filtrados]

        alto = max(6, len(categorias) * 0.8)
        fig, ax = plt.subplots(figsize=(12, alto))
        fig.patch.set_facecolor('white')

        y_pos = np.arange(len(categorias))
        bar_height = 0.35

        # Barras
        bars_m = ax.barh(y_pos - bar_height/2, valores_m, bar_height,
                         label='Mujeres', color=COLORES['mujer'], alpha=0.9)
        bars_h = ax.barh(y_pos + bar_height/2, valores_h, bar_height,
                         label='Hombres', color=COLORES['hombre'], alpha=0.9)

        # Etiquetas de valor
        for bar in bars_m:
            width = bar.get_width()
            if width > 0:
                ax.text(width + max(valores_m + valores_h) * 0.01, bar.get_y() + bar.get_height()/2,
                        formato_numero_es(width, 0) + ' €',
                        ha='left', va='center', fontsize=11, fontweight='bold')

        for bar in bars_h:
            width = bar.get_width()
            if width > 0:
                ax.text(width + max(valores_m + valores_h) * 0.01, bar.get_y() + bar.get_height()/2,
                        formato_numero_es(width, 0) + ' €',
                        ha='left', va='center', fontsize=11, fontweight='bold')

        ax.set_yticks(y_pos)
        ax.set_yticklabels(categorias)
        ax.set_xlabel(tipo_valor, fontsize=14, fontweight='bold')
        ax.set_title(titulo, fontsize=18, fontweight='bold', color='#cc0000', pad=20)
        ax.legend(fontsize=12, loc='lower right')

        plt.tight_layout()

        # Guardar en BytesIO
        output = io.BytesIO()
        plt.savefig(output, dpi=300, bbox_inches='tight', facecolor='white', format='png')
        plt.close()
        output.seek(0)

        return output
