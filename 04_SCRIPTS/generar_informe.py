# generar_informe.py
# -*- coding: utf-8 -*-
"""
Sistema completo para automatizar la creación de informes con visualizaciones
"""

import sys
import os

# Configurar codificación UTF-8 para evitar errores con emojis
if sys.platform == "win32":
    try:
        import locale
        locale.setlocale(locale.LC_ALL, 'Spanish_Spain.1252')
    except:
        pass
    
    # Configurar salida UTF-8 en Windows
    try:
        sys.stdout.reconfigure(encoding='utf-8')
        sys.stderr.reconfigure(encoding='utf-8')
    except:
        pass

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import yaml
import json
from datetime import datetime
from pathlib import Path

class AutomatedReportSystem:
    def __init__(self, config_file="report_config.yaml"):
        """
        Sistema automatizado de generación de reportes
        """
        self.config = self.load_config(config_file)
        self.data = None
        self.charts_created = {}
        
        # ===== CONFIGURACIÓN GLOBAL DE GRÁFICOS =====
        self.config_graficos = {
            # Tamaños estandarizados
            'ancho_min': 10,
            'ancho_max': 16, 
            'alto_estandar': 8,
            'alto_donut': 10,
            
            # Configuración de títulos
            'titulo': {
                'fontsize': 16,
                'fontweight': 'bold',
                'color': '#cc0000',
                'pad': 20
            },
            
            # Configuración de etiquetas
            'etiquetas': {
                'fontsize': 12,
                'fontweight': 'normal'
            },
            
            # Configuración de valores en gráficos
            'valores': {
                'fontsize': 10,
                'fontweight': 'bold',
                'offset_factor': 0.02  # Factor para separar valores del borde
            },
            
            # Ancho de inserción en documento (inches)
            'ancho_documento': 6.5,
            
            # Resolución
            'dpi': 300
        }
        
        # Configuración de visualización matplotlib
        plt.rcParams['font.size'] = self.config_graficos['etiquetas']['fontsize']
        sns.set_style("whitegrid")
        
        # Paleta de colores para género
        self.colores_genero = {
            'H': '#ea5d41',  # Rojo para hombres
            'M': '#1e4389',  # Azul para mujeres
        }

    def calcular_tamaño_grafico(self, num_elementos=1, tipo_grafico='barra'):
        """
        Calcula el tamaño óptimo para un gráfico según el número de elementos
        """
        config = self.config_graficos
        
        if tipo_grafico == 'donut':
            return (config['ancho_max'], config['alto_donut'])
        elif tipo_grafico == 'barra_horizontal':
            # Para barras horizontales, ajustar altura según elementos
            ancho = min(config['ancho_max'], max(config['ancho_min'], num_elementos * 1.2))
            alto = max(config['alto_estandar'], min(20, num_elementos * 0.6))
            return (ancho, alto)
        elif tipo_grafico == 'barra_vertical':
            # Para barras verticales, ajustar ancho según elementos
            ancho = min(config['ancho_max'], max(config['ancho_min'], num_elementos * 1.8))
            alto = config['alto_estandar']
            return (ancho, alto)
        else:
            # Tamaño estándar
            return (config['ancho_max'], config['alto_estandar'])

    def aplicar_estilo_titulo(self, ax, titulo):
        """
        Aplica el estilo estandarizado a los títulos de gráficos
        """
        config = self.config_graficos['titulo']
        ax.set_title(titulo, 
                    fontsize=config['fontsize'],
                    fontweight=config['fontweight'], 
                    color=config['color'],
                    pad=config['pad'])

    def ajustar_posicion_valores(self, ax, valores_y, margen_extra=0):
        """
        Ajusta la posición de valores para evitar que se salgan del gráfico
        """
        config = self.config_graficos['valores']
        
        # Calcular límites del gráfico
        y_min, y_max = ax.get_ylim()
        rango_y = y_max - y_min
        offset = rango_y * config['offset_factor'] + margen_extra
        
        # Ajustar límite superior si es necesario
        max_valor_y = max(valores_y) if valores_y else 0
        if max_valor_y + offset > y_max:
            ax.set_ylim(y_min, max_valor_y + offset * 1.5)
        
        return offset

    def configurar_grafico_base(self, fig, ax, titulo=""):
        """
        Aplica configuraciones base comunes a todos los gráficos
        """
        # Fondo blanco
        fig.patch.set_facecolor('white')
        
        # Título si se proporciona
        if titulo:
            self.aplicar_estilo_titulo(ax, titulo)
        
        # Grid
        ax.grid(True, axis='y', alpha=0.3, linestyle='-', linewidth=0.5)
        ax.set_axisbelow(True)
    
    def insertar_imagen_estandarizada(self, doc, filename):
        """
        Inserta una imagen en el documento con formato estandarizado
        """
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(filename, width=Inches(self.config_graficos['ancho_documento']))
        return paragraph
    
    def load_config(self, config_file):
        """Carga la configuración desde un archivo YAML"""
        default_config = {
            'excel_file': '',  # Se determina automáticamente
            'template_word': 'plantilla_informe.docx',  # Opcional
            'output_file': '05_INFORMES/registro_retributivo.docx',  # Se generará timestamp dinámico
            'charts': {
                'salario_base_efectivo': {
                    'type': 'donut',
                    'data_columns': ['Salario base efectivo Total'],
                    'metodo': 'efectivos_sb',
                    'title': 'Comparación Salario Base Efectivo Total por Género',
                    'subtitulo': 'Análisis de igualdad retributiva - Salario base efectivamente percibido (solo SB > 0)',
                    'marker': '{grafico_sb_efectivo}'
                },
                'sb_complementos_efectivo': {
                    'type': 'donut',
                    'data_columns': ['Salario base anual + complementos Total'],
                    'metodo': 'efectivos_sb_complementos',
                    'title': 'Salario Base + Complementos Salariales Efectivos por Género',
                    'subtitulo': 'Incluye salario base y complementos salariales efectivamente percibidos (todas las personas)',
                    'marker': '{grafico_sb_comp_efectivo}'
                },
                'sb_total_efectivo': {
                    'type': 'donut',
                    'data_columns': ['Salario base anual + complementos + Extrasalariales Total'],
                    'metodo': 'efectivos_sb_complementos',
                    'title': 'SB + Complementos + Extrasalariales Efectivos por Género',
                    'subtitulo': 'Retribución total efectiva incluyendo todos los conceptos (todas las personas)',
                    'marker': '{grafico_sb_total_efectivo}'
                },
                'salario_base_equiparado': {
                    'type': 'donut',
                    'data_columns': ['salario_base_equiparado'],
                    'metodo': 'equiparados_sb',
                    'title': 'Comparación Salario Base Equiparado por Género',
                    'subtitulo': 'Salario base normalizado a jornada completa y 12 meses (solo SB > 0)',
                    'marker': '{grafico_sb_equiparado}'
                },
                'sb_complementos_equiparado': {
                    'type': 'donut',
                    'data_columns': ['sb_mas_comp_salariales_equiparado'],
                    'metodo': 'equiparados_sb_complementos',
                    'title': 'Salario Base + Complementos Salariales Equiparados por Género',
                    'subtitulo': 'SB + complementos salariales normalizados a jornada completa y 12 meses (todas las personas)',
                    'marker': '{grafico_sb_comp_equiparado}'
                },
                'sb_total_equiparado': {
                    'type': 'donut',
                    'data_columns': ['sb_mas_comp_total_equiparado'],
                    'metodo': 'equiparados_sb_complementos',
                    'title': 'SB + Complementos + Extrasalariales Equiparados por Género',
                    'subtitulo': 'Retribución total equiparada: SB + complementos salariales y extrasalariales (todas las personas)',
                    'marker': '{grafico_sb_total_equiparado}'
                }
            }
        }
        
        try:
            if os.path.exists(config_file):
                with open(config_file, 'r', encoding='utf-8') as f:
                    user_config = yaml.safe_load(f)
                # Combinar configuraciones
                default_config.update(user_config)
        except Exception as e:
            print(f"Error cargando configuración: {e}")
            print("Usando configuración por defecto")
        
        return default_config
    
    def calcular_brecha_entre_valores(self, valor_h, valor_m):
        """
        Calcula la brecha salarial entre dos valores (hombres y mujeres)
        Devuelve el porcentaje de diferencia
        """
        if valor_m is None or valor_m == 0 or valor_h is None or valor_h == 0:
            return None
        
        # Brecha = ((Hombres - Mujeres) / Mujeres) * 100
        brecha = ((valor_h - valor_m) / valor_m) * 100
        return brecha

    def calcular_brecha_salarial(self):
        """Calcula la brecha salarial por grupo profesional"""
        if 'Sexo' not in self.data.columns or 'Salario base anual efectivo' not in self.data.columns:
            print("No se pueden calcular brechas: faltan columnas necesarias")
            return None
        
        # Calcular brecha por grupo profesional si existe
        if 'Grupo profesional' in self.data.columns:
            grupos = self.data['Grupo profesional'].unique()
            brechas = []
            
            for grupo in grupos:
                data_grupo = self.data[self.data['Grupo profesional'] == grupo]
                
                # Calcular salarios promedio por género
                salarios_genero = data_grupo.groupby('Sexo')['Salario base anual efectivo'].mean()
                
                if 'H' in salarios_genero.index and 'M' in salarios_genero.index:
                    salario_h = salarios_genero['H']
                    salario_m = salarios_genero['M']
                    
                    # Calcular brecha porcentual (diferencia respecto al salario mayor)
                    brecha = ((salario_h - salario_m) / max(salario_h, salario_m)) * 100
                    brechas.append({'Grupo profesional': grupo, 'brecha_porcentual': brecha})
                    print(f"Grupo {grupo}: Brecha salarial = {brecha:.2f}%")
            
            if brechas:
                df_brechas = pd.DataFrame(brechas)
                # Añadir las brechas al dataframe principal
                self.data = self.data.merge(df_brechas, on='Grupo profesional', how='left')
                print("Brechas salariales calculadas y añadidas al dataset")
                
                # Añadir gráfico de brecha si no existe en la configuración
                if 'brecha_por_grupo' not in self.config['charts']:
                    self.config['charts']['brecha_por_grupo'] = {
                        'type': 'bar',
                        'data_columns': ['Grupo profesional', 'brecha_porcentual'],
                        'title': 'Brecha Salarial por Grupo Profesional (%)',
                        'marker': '{grafico_brecha_grupo}'
                    }
                    print("Gráfico de brecha salarial añadido a la configuración")
                
                return df_brechas
        
        return None
    
    def formato_numero_es(self, numero, decimales=2):
        """Formatea un número con estilo español (punto como separador de miles)"""
        if pd.isna(numero):
            return "0,00"
        
        # Formatear el número con decimales
        numero_formateado = f"{numero:,.{decimales}f}"
        
        # Cambiar punto por coma para decimales y coma por punto para miles (estilo español)
        numero_formateado = numero_formateado.replace(',', 'X').replace('.', ',').replace('X', '.')
        
        return numero_formateado
    
    def formato_brecha_es(self, brecha, decimales=2):
        """Formatea una brecha siempre en valor absoluto con estilo español"""
        if pd.isna(brecha):
            return "0,00"
        
        # Usar valor absoluto para que siempre sea positivo
        brecha_absoluta = abs(brecha)
        return self.formato_numero_es(brecha_absoluta, decimales)
    
    def generar_resumen_estadisticas(self):
        """Genera un resumen completo de estadísticas para las 3 tablas principales"""
        print("📊 Calculando estadísticas resumidas...")
        
        # Contar totales por género
        df_actual = self.data[self.data.iloc[:, 0] != 'Ex'].copy()
        conteos = df_actual['SEXO'].value_counts()
        total_mujeres = conteos.get('M', 0)
        total_hombres = conteos.get('H', 0)
        
        resumen = {
            'conteos': {
                'mujeres': total_mujeres,
                'hombres': total_hombres
            },
            'tabla1_sb': {},
            'tabla2_sb_comp': {},
            'tabla3_sb_comp_extra': {}
        }
        
        # Tabla 1: Salario Base (SB)
        # Efectivo
        datos_sb_efectivo = self.calcular_promedios_efectivos_sb(self.data, 'Salario base efectivo Total')
        brecha_sb_efectivo = ((datos_sb_efectivo['H'] - datos_sb_efectivo['M']) / datos_sb_efectivo['M']) * 100 if datos_sb_efectivo['M'] > 0 else 0
        
        # Equiparado
        datos_sb_equiparado = self.calcular_promedios_equiparados_sb(self.data, 'salario_base_equiparado')
        brecha_sb_equiparado = ((datos_sb_equiparado['H'] - datos_sb_equiparado['M']) / datos_sb_equiparado['M']) * 100 if datos_sb_equiparado['M'] > 0 else 0
        
        resumen['tabla1_sb'] = {
            'efectivo': {
                'mujeres': datos_sb_efectivo['M'],
                'hombres': datos_sb_efectivo['H'],
                'brecha': brecha_sb_efectivo
            },
            'equiparado': {
                'mujeres': datos_sb_equiparado['M'],
                'hombres': datos_sb_equiparado['H'],
                'brecha': brecha_sb_equiparado
            }
        }
        
        # Tabla 2: SB + Complementos (SB+C)
        # Efectivo
        datos_sbc_efectivo = self.calcular_promedios_efectivos_sb_complementos(self.data, 'Salario base anual + complementos Total')
        brecha_sbc_efectivo = ((datos_sbc_efectivo['H'] - datos_sbc_efectivo['M']) / datos_sbc_efectivo['M']) * 100 if datos_sbc_efectivo['M'] > 0 else 0
        
        # Equiparado
        datos_sbc_equiparado = self.calcular_promedios_equiparados_sb_complementos(self.data, 'sb_mas_comp_salariales_equiparado')
        brecha_sbc_equiparado = ((datos_sbc_equiparado['H'] - datos_sbc_equiparado['M']) / datos_sbc_equiparado['M']) * 100 if datos_sbc_equiparado['M'] > 0 else 0
        
        resumen['tabla2_sb_comp'] = {
            'efectivo': {
                'mujeres': datos_sbc_efectivo['M'],
                'hombres': datos_sbc_efectivo['H'],
                'brecha': brecha_sbc_efectivo
            },
            'equiparado': {
                'mujeres': datos_sbc_equiparado['M'],
                'hombres': datos_sbc_equiparado['H'],
                'brecha': brecha_sbc_equiparado
            }
        }
        
        # Tabla 3: SB + Complementos + Extrasalariales (SB+C+ES)
        # Efectivo
        datos_total_efectivo = self.calcular_promedios_efectivos_sb_complementos(self.data, 'Salario base anual + complementos + Extrasalariales Total')
        brecha_total_efectivo = ((datos_total_efectivo['H'] - datos_total_efectivo['M']) / datos_total_efectivo['M']) * 100 if datos_total_efectivo['M'] > 0 else 0
        
        # Equiparado
        datos_total_equiparado = self.calcular_promedios_equiparados_sb_complementos(self.data, 'sb_mas_comp_total_equiparado')
        brecha_total_equiparado = ((datos_total_equiparado['H'] - datos_total_equiparado['M']) / datos_total_equiparado['M'] ) * 100 if datos_total_equiparado['M'] > 0 else 0
        
        resumen['tabla3_sb_comp_extra'] = {
            'efectivo': {
                'mujeres': datos_total_efectivo['M'],
                'hombres': datos_total_efectivo['H'],
                'brecha': brecha_total_efectivo
            },
            'equiparado': {
                'mujeres': datos_total_equiparado['M'],
                'hombres': datos_total_equiparado['H'],
                'brecha': brecha_total_equiparado
            }
        }
        
        return resumen
    
    def generar_datos_grupo_profesional(self):
        """Genera los datos para la tabla de retribución por grupo profesional"""
        print("📊 Calculando datos por grupo profesional...")
        
        # Filtrar datos actuales (sin "Ex" en primera columna)
        df_actual = self.data[self.data.iloc[:, 0] != 'Ex'].copy()
        
        # Obtener grupos profesionales únicos
        if 'Grupo profesional' not in df_actual.columns:
            print("⚠️ Columna 'Grupo profesional' no encontrada")
            return []
        
        grupos = df_actual['Grupo profesional'].unique()
        # Convertir a string para evitar problemas de ordenación con tipos mixtos
        grupos = [str(g) for g in grupos if pd.notna(g)]
        datos_grupos = []
        
        for grupo in sorted(grupos):
            # Filtrar datos del grupo (convertir también la columna original a string para comparación)
            df_grupo = df_actual[df_actual['Grupo profesional'].astype(str) == grupo]
            
            # Contar por género
            conteos = df_grupo['SEXO'].value_counts()
            n_mujeres = conteos.get('M', 0)
            n_hombres = conteos.get('H', 0)
            
            if n_mujeres == 0 and n_hombres == 0:
                continue
            
            # Calcular retribución SIN complementos (solo salario base efectivo)
            col_sin_comp = 'Salario base efectivo Total'
            if col_sin_comp in df_grupo.columns:
                # Filtrar solo registros con salario > 0
                df_sin_comp = df_grupo[(df_grupo[col_sin_comp].notna()) & (df_grupo[col_sin_comp] > 0)]
                if len(df_sin_comp) > 0:
                    promedios_sin_comp = df_sin_comp.groupby('SEXO')[col_sin_comp].mean()
                    sin_comp_m = promedios_sin_comp.get('M', 0)
                    sin_comp_h = promedios_sin_comp.get('H', 0)
                    brecha_sin_comp = ((sin_comp_h - sin_comp_m) / sin_comp_m * 100) if sin_comp_m > 0 else 0
                else:
                    sin_comp_m = sin_comp_h = brecha_sin_comp = 0
            else:
                sin_comp_m = sin_comp_h = brecha_sin_comp = 0
            
            # Calcular retribución CON complementos
            col_con_comp = 'Salario base anual + complementos Total'
            if col_con_comp in df_grupo.columns:
                # Incluir todos los registros (incluir SB = 0)
                df_con_comp = df_grupo[df_grupo[col_con_comp].notna()]
                if len(df_con_comp) > 0:
                    promedios_con_comp = df_con_comp.groupby('SEXO')[col_con_comp].mean()
                    con_comp_m = promedios_con_comp.get('M', 0)
                    con_comp_h = promedios_con_comp.get('H', 0)
                    brecha_con_comp = ((con_comp_h - con_comp_m) / con_comp_m * 100) if con_comp_m > 0 else 0
                else:
                    con_comp_m = con_comp_h = brecha_con_comp = 0
            else:
                con_comp_m = con_comp_h = brecha_con_comp = 0
            
            # Calcular retribución CON complementos + extrasalariales
            col_con_extra = 'Salario base anual + complementos + Extrasalariales Total'
            if col_con_extra in df_grupo.columns:
                df_con_extra = df_grupo[df_grupo[col_con_extra].notna()]
                if len(df_con_extra) > 0:
                    promedios_con_extra = df_con_extra.groupby('SEXO')[col_con_extra].mean()
                    con_extra_m = promedios_con_extra.get('M', 0)
                    con_extra_h = promedios_con_extra.get('H', 0)
                    brecha_con_extra = ((con_extra_h - con_extra_m) / con_extra_m * 100) if con_extra_m > 0 else 0
                else:
                    con_extra_m = con_extra_h = brecha_con_extra = 0
            else:
                con_extra_m = con_extra_h = brecha_con_extra = 0
            
            datos_grupos.append({
                'grupo': grupo,
                'n_mujeres': n_mujeres,
                'n_hombres': n_hombres,
                'sin_comp_m': sin_comp_m,
                'sin_comp_h': sin_comp_h,
                'brecha_sin_comp': brecha_sin_comp,
                'con_comp_m': con_comp_m,
                'con_comp_h': con_comp_h,
                'brecha_con_comp': brecha_con_comp,
                'con_extra_m': con_extra_m,
                'con_extra_h': con_extra_h,
                'brecha_con_extra': brecha_con_extra
            })
        
        # Calcular totales
        if datos_grupos:
            total_m = sum([d['n_mujeres'] for d in datos_grupos])
            total_h = sum([d['n_hombres'] for d in datos_grupos])
            
            # Totales para sin complementos
            datos_totales_sin = self.calcular_promedios_efectivos_sb(df_actual, col_sin_comp)
            brecha_total_sin = ((datos_totales_sin['H'] - datos_totales_sin['M']) / datos_totales_sin['M'] * 100) if datos_totales_sin['M'] > 0 else 0
            
            # Totales para con complementos
            datos_totales_con = self.calcular_promedios_efectivos_sb_complementos(df_actual, col_con_comp)
            brecha_total_con = ((datos_totales_con['H'] - datos_totales_con['M']) / datos_totales_con['M'] * 100) if datos_totales_con['M'] > 0 else 0
            
            # Totales para con extrasalariales
            datos_totales_extra = self.calcular_promedios_efectivos_sb_complementos(df_actual, col_con_extra)
            brecha_total_extra = ((datos_totales_extra['H'] - datos_totales_extra['M']) / datos_totales_extra['M'] * 100) if datos_totales_extra['M'] > 0 else 0
            
            # Insertar totales al principio
            datos_grupos.insert(0, {
                'grupo': 'Totales',
                'n_mujeres': total_m,
                'n_hombres': total_h,
                'sin_comp_m': datos_totales_sin['M'],
                'sin_comp_h': datos_totales_sin['H'],
                'brecha_sin_comp': brecha_total_sin,
                'con_comp_m': datos_totales_con['M'],
                'con_comp_h': datos_totales_con['H'],
                'brecha_con_comp': brecha_total_con,
                'con_extra_m': datos_totales_extra['M'],
                'con_extra_h': datos_totales_extra['H'],
                'brecha_con_extra': brecha_total_extra
            })
        
        return datos_grupos
    
    def generar_datos_grupo_profesional_equiparado(self):
        """Genera los datos para la tabla de retribución por grupo profesional EQUIPARADO"""
        print("📊 Calculando datos equiparados por grupo profesional...")
        
        # Filtrar datos actuales (sin "Ex" en primera columna)
        df_actual = self.data[self.data.iloc[:, 0] != 'Ex'].copy()
        
        # Obtener grupos profesionales únicos
        if 'Grupo profesional' not in df_actual.columns:
            print("❌ Error: No se encontró la columna 'Grupo profesional'")
            return []
        
        grupos = df_actual['Grupo profesional'].unique()
        # Convertir a string para evitar problemas de ordenación con tipos mixtos
        grupos = [str(g) for g in grupos if pd.notna(g)]
        datos_grupos = []
        
        # Columnas para valores equiparados
        col_sin_comp = 'salario_base_equiparado'  # SB equiparado
        col_con_comp = 'sb_mas_comp_salariales_equiparado'  # SB + Complementos equiparado
        col_con_extra = 'sb_mas_comp_total_equiparado'  # SB + Complementos + Extrasalariales equiparado
        
        for grupo in sorted(grupos):
            # Manejar tipos mixtos en la columna 'Grupo profesional' (string, int, float)
            # Usar conversión a string para hacer comparaciones consistentes
            df_grupo = df_actual[df_actual['Grupo profesional'].astype(str) == str(grupo)]
            
            # Conteos por género
            conteos = df_grupo['SEXO'].value_counts()
            n_mujeres = conteos.get('M', 0)
            n_hombres = conteos.get('H', 0)
            
            # SIN complementos (solo SB equiparado) - solo personas con SB > 0
            datos_sin_comp = self.calcular_promedios_equiparados_sb(df_grupo, col_sin_comp)
            sin_comp_m = datos_sin_comp['M']
            sin_comp_h = datos_sin_comp['H']
            brecha_sin_comp = ((sin_comp_h - sin_comp_m) / sin_comp_m * 100) if sin_comp_m > 0 else 0
            
            # CON complementos equiparados - todas las personas
            datos_con_comp = self.calcular_promedios_equiparados_sb_complementos(df_grupo, col_con_comp)
            con_comp_m = datos_con_comp['M']
            con_comp_h = datos_con_comp['H']
            brecha_con_comp = ((con_comp_h - con_comp_m) / con_comp_m * 100) if con_comp_m > 0 else 0
            
            # CON extrasalariales equiparados - todas las personas
            datos_con_extra = self.calcular_promedios_equiparados_sb_complementos(df_grupo, col_con_extra)
            con_extra_m = datos_con_extra['M']
            con_extra_h = datos_con_extra['H']
            brecha_con_extra = ((con_extra_h - con_extra_m) / con_extra_m * 100) if con_extra_m > 0 else 0
            
            datos_grupos.append({
                'grupo': grupo,
                'n_mujeres': n_mujeres,
                'n_hombres': n_hombres,
                'sin_comp_m': sin_comp_m,
                'sin_comp_h': sin_comp_h,
                'brecha_sin_comp': brecha_sin_comp,
                'con_comp_m': con_comp_m,
                'con_comp_h': con_comp_h,
                'brecha_con_comp': brecha_con_comp,
                'con_extra_m': con_extra_m,
                'con_extra_h': con_extra_h,
                'brecha_con_extra': brecha_con_extra
            })
        
        # Calcular totales equiparados
        if datos_grupos:
            total_m = sum([d['n_mujeres'] for d in datos_grupos])
            total_h = sum([d['n_hombres'] for d in datos_grupos])
            
            # Totales para sin complementos equiparados
            datos_totales_sin = self.calcular_promedios_equiparados_sb(df_actual, col_sin_comp)
            brecha_total_sin = ((datos_totales_sin['H'] - datos_totales_sin['M']) / datos_totales_sin['M'] * 100) if datos_totales_sin['M'] > 0 else 0
            
            # Totales para con complementos equiparados
            datos_totales_con = self.calcular_promedios_equiparados_sb_complementos(df_actual, col_con_comp)
            brecha_total_con = ((datos_totales_con['H'] - datos_totales_con['M']) / datos_totales_con['M'] * 100) if datos_totales_con['M'] > 0 else 0
            
            # Totales para con extrasalariales equiparados
            datos_totales_extra = self.calcular_promedios_equiparados_sb_complementos(df_actual, col_con_extra)
            brecha_total_extra = ((datos_totales_extra['H'] - datos_totales_extra['M']) / datos_totales_extra['M'] * 100) if datos_totales_extra['M'] > 0 else 0
            
            # Insertar totales al principio
            datos_grupos.insert(0, {
                'grupo': 'Totales',
                'n_mujeres': total_m,
                'n_hombres': total_h,
                'sin_comp_m': datos_totales_sin['M'],
                'sin_comp_h': datos_totales_sin['H'],
                'brecha_sin_comp': brecha_total_sin,
                'con_comp_m': datos_totales_con['M'],
                'con_comp_h': datos_totales_con['H'],
                'brecha_con_comp': brecha_total_con,
                'con_extra_m': datos_totales_extra['M'],
                'con_extra_h': datos_totales_extra['H'],
                'brecha_con_extra': brecha_total_extra
            })
        
        return datos_grupos
    
    def crear_tabla_grupo_profesional(self, doc, datos_grupos):
        """Crea la tabla de retribución por grupo profesional"""
        from docx.shared import Cm, Pt
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.shared import OxmlElement, qn
        
        if not datos_grupos:
            doc.add_paragraph("No hay datos disponibles para generar la tabla por grupo profesional.")
            return
        
        # Crear tabla con cabeceras según la imagen
        num_filas = len(datos_grupos) + 1  # +1 para cabecera
        table = doc.add_table(rows=num_filas, cols=12)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        
        # Definir cabeceras según la imagen proporcionada
        headers = [
            'Grupo Profesional', 'Nº M', 'Nº H',
            'Retribución Promedio SIN Complementos (Mujeres)', 'Retribución Promedio SIN Complementos (Hombres)', 'Brecha Salarial SIN Complementos',
            'Retribución Promedio CON Complementos (Mujeres)', 'Retribución Promedio CON Complementos (Hombres)', 'Brecha Salarial CON Complementos',
            'Retribución Promedio CON Complementos ES (Mujeres)', 'Retribución Promedio CON Complementos ES (Hombres)', 'Brecha Salarial CON Complementos ES'
        ]
        
        # Configurar cabeceras
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            if i < len(header_row.cells):
                cell = header_row.cells[i]
                cell.text = header
                # Aplicar formato de cabecera
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(8)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Llenar datos
        for row_idx, datos in enumerate(datos_grupos, 1):
            if row_idx < len(table.rows):
                data_row = table.rows[row_idx]
                
                # Preparar valores
                values = [
                    str(datos['grupo']),
                    str(datos['n_mujeres']),
                    str(datos['n_hombres']),
                    f"{self.formato_numero_es(datos['sin_comp_m'], 2)} €" if datos['sin_comp_m'] > 0 else "-",
                    f"{self.formato_numero_es(datos['sin_comp_h'], 2)} €" if datos['sin_comp_h'] > 0 else "-",
                    f"{self.formato_brecha_es(datos['brecha_sin_comp'], 2)}%" if datos['sin_comp_m'] > 0 and datos['sin_comp_h'] > 0 else "-",
                    f"{self.formato_numero_es(datos['con_comp_m'], 2)} €" if datos['con_comp_m'] > 0 else "-",
                    f"{self.formato_numero_es(datos['con_comp_h'], 2)} €" if datos['con_comp_h'] > 0 else "-",
                    f"{self.formato_brecha_es(datos['brecha_con_comp'], 2)}%" if datos['con_comp_m'] > 0 and datos['con_comp_h'] > 0 else "-",
                    f"{self.formato_numero_es(datos['con_extra_m'], 2)} €" if datos['con_extra_m'] > 0 else "-",
                    f"{self.formato_numero_es(datos['con_extra_h'], 2)} €" if datos['con_extra_h'] > 0 else "-",
                    f"{self.formato_brecha_es(datos['brecha_con_extra'], 2)}%" if datos['con_extra_m'] > 0 and datos['con_extra_h'] > 0 else "-"
                ]
                
                # Llenar celdas y aplicar colores
                for i, value in enumerate(values):
                    if i < len(data_row.cells):
                        cell = data_row.cells[i]
                        cell.text = value
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph.runs:
                                run.font.size = Pt(8)
                        
                        # Aplicar colores según las reglas de negocio
                        es_total = (datos['grupo'] == 'Totales')
                        self.aplicar_color_celda_grupo_profesional(cell, i, datos, es_total)
        
        # Agregar espacio después de la tabla
        doc.add_paragraph()
        
        return table
    
    def aplicar_color_celda_grupo_profesional(self, cell, columna_index, datos, es_total):
        """Aplica colores al TEXTO de las celdas de la tabla de grupo profesional"""
        # Colores para el texto - mismo esquema que tabla principal
        ROJO = RGBColor(234, 93, 65)   # Cuando favorece a hombres (#ea5d41)
        AZUL = RGBColor(30, 67, 137)   # Cuando favorece a mujeres (#1e4389)
        NEGRO = RGBColor(0, 0, 0)      # Para fila de totales
        
        # Destacar fila de totales con texto negro normal
        if es_total:
            self.colorear_texto_celda(cell, NEGRO)
            return
        
        # Aplicar colores a las columnas de salarios y brechas
        if columna_index == 3:  # Sin complementos - Mujeres
            if datos['sin_comp_m'] < datos['sin_comp_h'] and datos['sin_comp_m'] > 0 and datos['sin_comp_h'] > 0:
                self.colorear_texto_celda(cell, ROJO)  # Desfavorable para mujeres
            elif datos['sin_comp_m'] > datos['sin_comp_h'] and datos['sin_comp_m'] > 0 and datos['sin_comp_h'] > 0:
                self.colorear_texto_celda(cell, AZUL)   # Favorable para mujeres
                
        elif columna_index == 4:  # Sin complementos - Hombres
            if datos['sin_comp_h'] > datos['sin_comp_m'] and datos['sin_comp_m'] > 0 and datos['sin_comp_h'] > 0:
                self.colorear_texto_celda(cell, AZUL)   # Favorable para hombres
            elif datos['sin_comp_h'] < datos['sin_comp_m'] and datos['sin_comp_m'] > 0 and datos['sin_comp_h'] > 0:
                self.colorear_texto_celda(cell, ROJO)   # Desfavorable para hombres
                
        elif columna_index == 5:  # Brecha sin complementos
            if datos['sin_comp_m'] > 0 and datos['sin_comp_h'] > 0:
                if datos['brecha_sin_comp'] > 0:  # Favorable a hombres
                    self.colorear_texto_celda(cell, ROJO)
                elif datos['brecha_sin_comp'] < 0:  # Favorable a mujeres
                    self.colorear_texto_celda(cell, AZUL)
                    
        elif columna_index == 6:  # Con complementos - Mujeres
            if datos['con_comp_m'] < datos['con_comp_h'] and datos['con_comp_m'] > 0 and datos['con_comp_h'] > 0:
                self.colorear_texto_celda(cell, ROJO)  # Desfavorable para mujeres
            elif datos['con_comp_m'] > datos['con_comp_h'] and datos['con_comp_m'] > 0 and datos['con_comp_h'] > 0:
                self.colorear_texto_celda(cell, AZUL)   # Favorable para mujeres
                
        elif columna_index == 7:  # Con complementos - Hombres
            if datos['con_comp_h'] > datos['con_comp_m'] and datos['con_comp_m'] > 0 and datos['con_comp_h'] > 0:
                self.colorear_texto_celda(cell, AZUL)   # Favorable para hombres
            elif datos['con_comp_h'] < datos['con_comp_m'] and datos['con_comp_m'] > 0 and datos['con_comp_h'] > 0:
                self.colorear_texto_celda(cell, ROJO)   # Desfavorable para hombres
                
        elif columna_index == 8:  # Brecha con complementos
            if datos['con_comp_m'] > 0 and datos['con_comp_h'] > 0:
                if datos['brecha_con_comp'] > 0:  # Favorable a hombres
                    self.colorear_texto_celda(cell, ROJO)
                elif datos['brecha_con_comp'] < 0:  # Favorable a mujeres
                    self.colorear_texto_celda(cell, AZUL)
                    
        elif columna_index == 9:  # Con extrasalariales - Mujeres
            if datos['con_extra_m'] < datos['con_extra_h'] and datos['con_extra_m'] > 0 and datos['con_extra_h'] > 0:
                self.colorear_texto_celda(cell, ROJO)  # Desfavorable para mujeres
            elif datos['con_extra_m'] > datos['con_extra_h'] and datos['con_extra_m'] > 0 and datos['con_extra_h'] > 0:
                self.colorear_texto_celda(cell, AZUL)   # Favorable para mujeres
                
        elif columna_index == 10:  # Con extrasalariales - Hombres
            if datos['con_extra_h'] > datos['con_extra_m'] and datos['con_extra_m'] > 0 and datos['con_extra_h'] > 0:
                self.colorear_texto_celda(cell, AZUL)   # Favorable para hombres
            elif datos['con_extra_h'] < datos['con_extra_m'] and datos['con_extra_m'] > 0 and datos['con_extra_h'] > 0:
                self.colorear_texto_celda(cell, ROJO)   # Desfavorable para hombres
                
        elif columna_index == 11:  # Brecha con extrasalariales
            if datos['con_extra_m'] > 0 and datos['con_extra_h'] > 0:
                if datos['brecha_con_extra'] > 0:  # Favorable a hombres
                    self.colorear_texto_celda(cell, ROJO)
                elif datos['brecha_con_extra'] < 0:  # Favorable a mujeres
                    self.colorear_texto_celda(cell, AZUL)
    
    def crear_tabla_word(self, doc, titulo, resumen, tipo_tabla):
        """Crea una tabla de resumen en el documento Word siguiendo el formato exacto proporcionado"""
        from docx.shared import Cm
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.shared import OxmlElement, qn
        
        # Agregar título de la tabla
        heading = doc.add_heading(titulo, level=3)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Obtener datos según el tipo de tabla
        if tipo_tabla == 'SB':
            datos = resumen['tabla1_sb']
        elif tipo_tabla == 'SB+C':
            datos = resumen['tabla2_sb_comp']
        elif tipo_tabla == 'SB+C+ES':
            datos = resumen['tabla3_sb_comp_extra']
        
        # Crear tabla con el formato exacto: 1 fila cabecera + 1 fila datos, 8 columnas
        table = doc.add_table(rows=2, cols=8)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        
        # Definir cabeceras exactas según el formato proporcionado
        if tipo_tabla == 'SB':
            headers = [
                'Mujeres', 'Hombres',
                'Promedio (SB) Efectivo [Mujeres]', 'Promedio (SB) Efectivo [Hombres]', 'Brecha (SB) Efectivo',
                'Promedio (SB) Equiparado [Mujeres]', 'Promedio (SB) Equiparado [Hombres]', 'Brecha (SB) Equiparado'
            ]
        elif tipo_tabla == 'SB+C':
            headers = [
                'Mujeres', 'Hombres',
                'Promedio (SB+C) Efectivo [Mujeres]', 'Promedio (SB+C) Efectivo [Hombres]', 'Brecha (SB+C) Efectivo',
                'Promedio (SB+C) Equiparado [Mujeres]', 'Promedio (SB+C) Equiparado [Hombres]', 'Brecha (SB+C) Equiparado'
            ]
        else:  # SB+C+ES
            headers = [
                'Mujeres', 'Hombres',
                'Promedio (SB+C+ES) [Mujeres] Efectivo', 'Promedio (SB+C+ES) [Hombres] Efectivo', 'Brecha (SB+C+ES) Efectivo',
                'Promedio (SB+C+ES) Equiparado [Mujeres]', 'Promedio (SB+C+ES) Equiparado [Hombres]', 'Brecha (SB+C+ES) Equiparado'
            ]
        
        # Configurar cabeceras
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            if i < len(header_row.cells):
                cell = header_row.cells[i]
                cell.text = header
                # Aplicar formato de cabecera
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(8)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Llenar la fila de datos con TODAS las 8 columnas
        data_row = table.rows[1]
        values = [
            str(resumen['conteos']['mujeres']),
            str(resumen['conteos']['hombres']),
            f"{self.formato_numero_es(datos['efectivo']['mujeres'], 2)} €",
            f"{self.formato_numero_es(datos['efectivo']['hombres'], 2)} €", 
            f"{self.formato_brecha_es(datos['efectivo']['brecha'], 2)}%",
            f"{self.formato_numero_es(datos['equiparado']['mujeres'], 2)} €",
            f"{self.formato_numero_es(datos['equiparado']['hombres'], 2)} €",
            f"{self.formato_brecha_es(datos['equiparado']['brecha'], 2)}%"
        ]
        
        # Llenar todas las celdas de datos con colores según las reglas
        for i, value in enumerate(values):
            if i < len(data_row.cells):
                cell = data_row.cells[i]
                cell.text = value
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                
                # Aplicar colores según las reglas de negocio
                self.aplicar_color_celda(cell, i, datos)
        
        # Agregar espacio después de la tabla
        doc.add_paragraph()
        
        return table
    
    def aplicar_color_celda(self, cell, columna_index, datos):
        """Aplica colores al TEXTO según las reglas: rojo cuando favorece a hombres, azul cuando favorece a mujeres"""
        # Colores para el texto
        ROJO = RGBColor(234, 93, 65)   # Cuando favorece a hombres (#ea5d41)
        AZUL = RGBColor(30, 67, 137)   # Cuando favorece a mujeres (#1e4389)
        
        # Columna 2: Promedio Efectivo Mujeres
        if columna_index == 2:
            if datos['efectivo']['mujeres'] < datos['efectivo']['hombres']:
                self.colorear_texto_celda(cell, ROJO)  # Desfavorable para mujeres
            else:
                self.colorear_texto_celda(cell, AZUL)  # Favorable para mujeres
                
        # Columna 3: Promedio Efectivo Hombres
        elif columna_index == 3:
            if datos['efectivo']['hombres'] > datos['efectivo']['mujeres']:
                self.colorear_texto_celda(cell, AZUL)  # Favorable para hombres
            else:
                self.colorear_texto_celda(cell, ROJO)  # Desfavorable para hombres
                
        # Columna 4: Brecha Efectivo
        elif columna_index == 4:
            if datos['efectivo']['brecha'] > 0:  # Favorable a hombres
                self.colorear_texto_celda(cell, ROJO)
            else:  # Favorable a mujeres
                self.colorear_texto_celda(cell, AZUL)
                
        # Columna 5: Promedio Equiparado Mujeres
        elif columna_index == 5:
            if datos['equiparado']['mujeres'] < datos['equiparado']['hombres']:
                self.colorear_texto_celda(cell, ROJO)  # Desfavorable para mujeres
            else:
                self.colorear_texto_celda(cell, AZUL)  # Favorable para mujeres
                
        # Columna 6: Promedio Equiparado Hombres
        elif columna_index == 6:
            if datos['equiparado']['hombres'] > datos['equiparado']['mujeres']:
                self.colorear_texto_celda(cell, AZUL)  # Favorable para hombres
            else:
                self.colorear_texto_celda(cell, ROJO)  # Desfavorable para hombres
                
        # Columna 7: Brecha Equiparado
        elif columna_index == 7:
            if datos['equiparado']['brecha'] > 0:  # Favorable a hombres
                self.colorear_texto_celda(cell, ROJO)
            else:  # Favorable a mujeres
                self.colorear_texto_celda(cell, AZUL)
    
    def colorear_texto_celda(self, cell, color_rgb):
        """Aplica un color al texto de una celda"""
        try:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = color_rgb
        except Exception as e:
            print(f"Error al colorear texto de celda: {e}")
            # Continuamos sin colorear si hay error
    
    def colorear_celda_simple(self, cell, color_hex):
        """Aplica un color de fondo simple a una celda"""
        try:
            from docx.oxml.shared import OxmlElement, qn
            
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), color_hex)
            tcPr.append(shd)
        except Exception as e:
            print(f"Error al colorear celda: {e}")
            # Continuamos sin colorear si hay error
    
    def colorear_celda_brecha(self, cell, brecha):
        """Colorea una celda según el valor de la brecha"""
        # Determinar color según la brecha
        if abs(brecha) <= 2:
            color = "D4EDDA"  # Verde claro - brecha baja
        elif abs(brecha) <= 5:
            color = "FFF3CD"  # Amarillo claro - brecha media
        else:
            color = "F8D7DA"  # Rojo claro - brecha alta
        
        # Aplicar color de fondo usando método simplificado
        self.colorear_celda_simple(cell, color)
    
    def calcular_promedios_efectivos_sb(self, df, columna_salario):
        """
        Calcula promedios para SALARIO BASE EFECTIVOS siguiendo las reglas:
        - Solo registros actuales (excluir 'Ex')
        - Solo donde SB > 0
        - Usar columnas 'Total efectivo'
        """
        # Filtrar registros actuales (sin "Ex" en primera columna)
        df_actual = df[df.iloc[:, 0] != 'Ex'].copy()
        
        # Filtrar datos válidos (SB > 0) y calcular promedio solo de estos registros
        df_sb_mayor_0 = df_actual[(df_actual[columna_salario].notna()) & (df_actual[columna_salario] > 0)]
        
        # Calcular promedios por género solo de personas con SB > 0
        promedios = df_sb_mayor_0.groupby('SEXO')[columna_salario].mean()
        
        return {
            'H': promedios.get('H', 0),
            'M': promedios.get('M', 0)
        }
    
    def calcular_promedios_efectivos_sb_complementos(self, df, columna_salario):
        """
        Calcula promedios para SB + COMPLEMENTOS EFECTIVOS siguiendo las reglas:
        - Solo registros actuales (excluir 'Ex')
        - TODAS las personas (incluir SB = 0)
        - Usar columnas 'Total efectivo'
        """
        # Filtrar registros actuales (sin "Ex" en primera columna)
        df_actual = df[df.iloc[:, 0] != 'Ex'].copy()
        
        # Filtrar datos válidos (solo quitar nulos, incluir SB = 0)
        df_valido = df_actual[df_actual[columna_salario].notna()]
        
        # Calcular promedios por género (incluir todas las personas)
        promedios = df_valido.groupby('SEXO')[columna_salario].mean()
        
        return {
            'H': promedios.get('H', 0),
            'M': promedios.get('M', 0)
        }
    
    def calcular_promedios_equiparados_sb(self, df, columna_salario):
        """
        Calcula promedios para SALARIO BASE EQUIPARADOS siguiendo las reglas:
        - Solo registros actuales (excluir 'Ex')
        - Solo donde SB base > 0
        - Usar columnas equiparadas basadas en situación contractual actual
        """
        # Filtrar registros actuales (sin "Ex" en primera columna)
        df_actual = df[df.iloc[:, 0] != 'Ex'].copy()
        
        # Filtrar datos válidos (SB base > 0) - usar la columna de salario base efectivo como referencia
        col_sb_efectivo = "Salario base anual efectivo"
        if col_sb_efectivo in df.columns:
            df_sb_mayor_0 = df_actual[(df_actual[col_sb_efectivo].notna()) & (df_actual[col_sb_efectivo] > 0)]
        else:
            df_sb_mayor_0 = df_actual[df_actual[columna_salario].notna() & (df_actual[columna_salario] > 0)]
        
        # Calcular promedios por género
        promedios = df_sb_mayor_0.groupby('SEXO')[columna_salario].mean()
        
        return {
            'H': promedios.get('H', 0),
            'M': promedios.get('M', 0)
        }
    
    def calcular_promedios_equiparados_sb_complementos(self, df, columna_salario):
        """
        Calcula promedios para SB + COMPLEMENTOS EQUIPARADOS siguiendo las reglas:
        - Solo registros actuales (excluir 'Ex')
        - TODAS las personas (incluir SB = 0)
        - Usar columnas equiparadas basadas en situación contractual actual
        """
        # Filtrar registros actuales (sin "Ex" en primera columna)
        df_actual = df[df.iloc[:, 0] != 'Ex'].copy()
        
        # Filtrar datos válidos (solo quitar nulos, incluir SB = 0)
        df_valido = df_actual[df_actual[columna_salario].notna()]
        
        # Calcular promedios por género
        promedios = df_valido.groupby('SEXO')[columna_salario].mean()
        
        return {
            'H': promedios.get('H', 0),
            'M': promedios.get('M', 0)
        }
    
    # ==================== FUNCIONES PARA ANÁLISIS CON MEDIANA ====================
    
    def calcular_medianas_efectivos_sb(self, df, columna_salario):
        """
        Calcula medianas para SALARIO BASE EFECTIVOS siguiendo las reglas:
        - Solo registros actuales (excluir 'Ex')
        - Solo donde SB > 0
        - Usar columnas 'Total efectivo'
        """
        # Filtrar registros actuales (sin "Ex" en primera columna)
        df_actual = df[df.iloc[:, 0] != 'Ex'].copy()
        
        # Filtrar datos válidos (SB > 0) y calcular mediana solo de estos registros
        df_sb_mayor_0 = df_actual[(df_actual[columna_salario].notna()) & (df_actual[columna_salario] > 0)]
        
        # Calcular medianas por género solo de personas con SB > 0
        medianas = df_sb_mayor_0.groupby('SEXO')[columna_salario].median()
        
        return {
            'H': medianas.get('H', 0),
            'M': medianas.get('M', 0)
        }
    
    def calcular_medianas_efectivos_sb_complementos(self, df, columna_salario):
        """
        Calcula medianas para SB + COMPLEMENTOS EFECTIVOS siguiendo las reglas:
        - Solo registros actuales (excluir 'Ex')
        - TODAS las personas (incluir SB = 0)
        - Usar columnas 'Total efectivo'
        """
        # Filtrar registros actuales (sin "Ex" en primera columna)
        df_actual = df[df.iloc[:, 0] != 'Ex'].copy()
        
        # Filtrar datos válidos (solo quitar nulos, incluir SB = 0)
        df_valido = df_actual[df_actual[columna_salario].notna()]
        
        # Calcular medianas por género (incluir todas las personas)
        medianas = df_valido.groupby('SEXO')[columna_salario].median()
        
        return {
            'H': medianas.get('H', 0),
            'M': medianas.get('M', 0)
        }
    
    def calcular_medianas_equiparados_sb(self, df, columna_salario):
        """
        Calcula medianas para SALARIO BASE EQUIPARADOS siguiendo las reglas:
        - Solo registros actuales (excluir 'Ex')
        - Solo donde SB base > 0
        - Usar columnas equiparadas basadas en situación contractual actual
        """
        # Filtrar registros actuales (sin "Ex" en primera columna)
        df_actual = df[df.iloc[:, 0] != 'Ex'].copy()
        
        # Filtrar datos válidos (SB base > 0) - usar la columna de salario base efectivo como referencia
        col_sb_efectivo = "Salario base anual efectivo"
        if col_sb_efectivo in df.columns:
            df_sb_mayor_0 = df_actual[(df_actual[col_sb_efectivo].notna()) & (df_actual[col_sb_efectivo] > 0)]
        else:
            df_sb_mayor_0 = df_actual[df_actual[columna_salario].notna() & (df_actual[columna_salario] > 0)]
        
        # Calcular medianas por género
        medianas = df_sb_mayor_0.groupby('SEXO')[columna_salario].median()
        
        return {
            'H': medianas.get('H', 0),
            'M': medianas.get('M', 0)
        }
    
    def calcular_medianas_equiparados_sb_complementos(self, df, columna_salario):
        """
        Calcula medianas para SB + COMPLEMENTOS EQUIPARADOS siguiendo las reglas:
        - Solo registros actuales (excluir 'Ex')
        - TODAS las personas (incluir SB = 0)
        - Usar columnas equiparadas basadas en situación contractual actual
        """
        # Filtrar registros actuales (sin "Ex" en primera columna)
        df_actual = df[df.iloc[:, 0] != 'Ex'].copy()
        
        # Filtrar datos válidos (solo quitar nulos, incluir SB = 0)
        df_valido = df_actual[df_actual[columna_salario].notna()]
        
        # Calcular medianas por género
        medianas = df_valido.groupby('SEXO')[columna_salario].median()
        
        return {
            'H': medianas.get('H', 0),
            'M': medianas.get('M', 0)
        }
    
    def generar_resumen_estadisticas_mediana(self):
        """Genera un resumen completo de estadísticas CON MEDIANA para las 3 tablas principales"""
        print("📊 Calculando estadísticas resumidas con MEDIANA...")
        
        # Contar totales por género
        df_actual = self.data[self.data.iloc[:, 0] != 'Ex'].copy()
        conteos = df_actual['SEXO'].value_counts()
        total_mujeres = conteos.get('M', 0)
        total_hombres = conteos.get('H', 0)
        
        resumen = {
            'conteos': {
                'mujeres': total_mujeres,
                'hombres': total_hombres
            },
            'tabla1_sb': {},
            'tabla2_sb_comp': {},
            'tabla3_sb_comp_extra': {}
        }
        
        # Tabla 1: Salario Base (SB)
        # Efectivo
        datos_sb_efectivo = self.calcular_medianas_efectivos_sb(self.data, 'Salario base efectivo Total')
        brecha_sb_efectivo = ((datos_sb_efectivo['H'] - datos_sb_efectivo['M']) / datos_sb_efectivo['M']) * 100 if datos_sb_efectivo['M'] > 0 else 0
        
        # Equiparado
        datos_sb_equiparado = self.calcular_medianas_equiparados_sb(self.data, 'salario_base_equiparado')
        brecha_sb_equiparado = ((datos_sb_equiparado['H'] - datos_sb_equiparado['M']) / datos_sb_equiparado['M']) * 100 if datos_sb_equiparado['M'] > 0 else 0
        
        resumen['tabla1_sb'] = {
            'efectivo': {
                'mujeres': datos_sb_efectivo['M'],
                'hombres': datos_sb_efectivo['H'],
                'brecha': brecha_sb_efectivo
            },
            'equiparado': {
                'mujeres': datos_sb_equiparado['M'],
                'hombres': datos_sb_equiparado['H'],
                'brecha': brecha_sb_equiparado
            }
        }
        
        # Tabla 2: SB + Complementos (SB+C)
        # Efectivo
        datos_sbc_efectivo = self.calcular_medianas_efectivos_sb_complementos(self.data, 'Salario base anual + complementos Total')
        brecha_sbc_efectivo = ((datos_sbc_efectivo['H'] - datos_sbc_efectivo['M']) / datos_sbc_efectivo['M']) * 100 if datos_sbc_efectivo['M'] > 0 else 0
        
        # Equiparado
        datos_sbc_equiparado = self.calcular_medianas_equiparados_sb_complementos(self.data, 'sb_mas_comp_salariales_equiparado')
        brecha_sbc_equiparado = ((datos_sbc_equiparado['H'] - datos_sbc_equiparado['M']) / datos_sbc_equiparado['M']) * 100 if datos_sbc_equiparado['M'] > 0 else 0
        
        resumen['tabla2_sb_comp'] = {
            'efectivo': {
                'mujeres': datos_sbc_efectivo['M'],
                'hombres': datos_sbc_efectivo['H'],
                'brecha': brecha_sbc_efectivo
            },
            'equiparado': {
                'mujeres': datos_sbc_equiparado['M'],
                'hombres': datos_sbc_equiparado['H'],
                'brecha': brecha_sbc_equiparado
            }
        }
        
        # Tabla 3: SB + Complementos + Extrasalariales (SB+C+E)
        # Efectivo
        datos_sbce_efectivo = self.calcular_medianas_efectivos_sb_complementos(self.data, 'Salario base anual + complementos + Extrasalariales Total')
        brecha_sbce_efectivo = ((datos_sbce_efectivo['H'] - datos_sbce_efectivo['M']) / datos_sbce_efectivo['M']) * 100 if datos_sbce_efectivo['M'] > 0 else 0
        
        # Equiparado
        datos_sbce_equiparado = self.calcular_medianas_equiparados_sb_complementos(self.data, 'sb_mas_comp_total_equiparado')
        brecha_sbce_equiparado = ((datos_sbce_equiparado['H'] - datos_sbce_equiparado['M']) / datos_sbce_equiparado['M']) * 100 if datos_sbce_equiparado['M'] > 0 else 0
        
        resumen['tabla3_sb_comp_extra'] = {
            'efectivo': {
                'mujeres': datos_sbce_efectivo['M'],
                'hombres': datos_sbce_efectivo['H'],
                'brecha': brecha_sbce_efectivo
            },
            'equiparado': {
                'mujeres': datos_sbce_equiparado['M'],
                'hombres': datos_sbce_equiparado['H'],
                'brecha': brecha_sbce_equiparado
            }
        }
        
        return resumen
    
    def generar_datos_grupo_profesional_mediana(self):
        """Genera los datos para la tabla de retribución por grupo profesional usando MEDIANA"""
        print("📊 Calculando datos por grupo profesional con MEDIANA...")
        
        # Filtrar datos actuales (sin "Ex" en primera columna)
        df_actual = self.data[self.data.iloc[:, 0] != 'Ex'].copy()
        
        # Obtener grupos profesionales únicos
        if 'Grupo profesional' not in df_actual.columns:
            print("⚠️ Columna 'Grupo profesional' no encontrada")
            return []
        
        grupos = df_actual['Grupo profesional'].unique()
        # Convertir a string para evitar problemas de ordenación con tipos mixtos
        grupos = [str(g) for g in grupos if pd.notna(g)]
        datos_grupos = []
        
        for grupo in sorted(grupos):
            # Filtrar datos del grupo (convertir también la columna original a string para comparación)
            df_grupo = df_actual[df_actual['Grupo profesional'].astype(str) == grupo]
            
            # Contar por género
            conteos = df_grupo['SEXO'].value_counts()
            n_mujeres = conteos.get('M', 0)
            n_hombres = conteos.get('H', 0)
            
            if n_mujeres == 0 and n_hombres == 0:
                continue
            
            # Calcular retribución SIN complementos (solo salario base efectivo) - MEDIANA
            col_sin_comp = 'Salario base efectivo Total'
            if col_sin_comp in df_grupo.columns:
                # Filtrar solo registros con salario > 0
                df_sin_comp = df_grupo[(df_grupo[col_sin_comp].notna()) & (df_grupo[col_sin_comp] > 0)]
                if len(df_sin_comp) > 0:
                    medianas_sin_comp = df_sin_comp.groupby('SEXO')[col_sin_comp].median()
                    sin_comp_m = medianas_sin_comp.get('M', 0)
                    sin_comp_h = medianas_sin_comp.get('H', 0)
                    brecha_sin_comp = ((sin_comp_h - sin_comp_m) / sin_comp_m * 100) if sin_comp_m > 0 else 0
                else:
                    sin_comp_m = sin_comp_h = brecha_sin_comp = 0
            else:
                sin_comp_m = sin_comp_h = brecha_sin_comp = 0
            
            # Calcular retribución CON complementos - MEDIANA
            col_con_comp = 'Salario base anual + complementos Total'
            if col_con_comp in df_grupo.columns:
                # Incluir todos los registros (incluir SB = 0)
                df_con_comp = df_grupo[df_grupo[col_con_comp].notna()]
                if len(df_con_comp) > 0:
                    medianas_con_comp = df_con_comp.groupby('SEXO')[col_con_comp].median()
                    con_comp_m = medianas_con_comp.get('M', 0)
                    con_comp_h = medianas_con_comp.get('H', 0)
                    brecha_con_comp = ((con_comp_h - con_comp_m) / con_comp_m * 100) if con_comp_m > 0 else 0
                else:
                    con_comp_m = con_comp_h = brecha_con_comp = 0
            else:
                con_comp_m = con_comp_h = brecha_con_comp = 0
            
            # Calcular retribución CON complementos + extrasalariales - MEDIANA
            col_con_extra = 'Salario base anual + complementos + Extrasalariales Total'
            if col_con_extra in df_grupo.columns:
                df_con_extra = df_grupo[df_grupo[col_con_extra].notna()]
                if len(df_con_extra) > 0:
                    medianas_con_extra = df_con_extra.groupby('SEXO')[col_con_extra].median()
                    con_extra_m = medianas_con_extra.get('M', 0)
                    con_extra_h = medianas_con_extra.get('H', 0)
                    brecha_con_extra = ((con_extra_h - con_extra_m) / con_extra_m * 100) if con_extra_m > 0 else 0
                else:
                    con_extra_m = con_extra_h = brecha_con_extra = 0
            else:
                con_extra_m = con_extra_h = brecha_con_extra = 0
            
            datos_grupos.append({
                'grupo': grupo,
                'n_mujeres': n_mujeres,
                'n_hombres': n_hombres,
                'sin_comp_m': sin_comp_m,
                'sin_comp_h': sin_comp_h,
                'brecha_sin_comp': brecha_sin_comp,
                'con_comp_m': con_comp_m,
                'con_comp_h': con_comp_h,
                'brecha_con_comp': brecha_con_comp,
                'con_extra_m': con_extra_m,
                'con_extra_h': con_extra_h,
                'brecha_con_extra': brecha_con_extra
            })
        
        # Calcular totales usando MEDIANA
        if datos_grupos:
            total_m = sum([d['n_mujeres'] for d in datos_grupos])
            total_h = sum([d['n_hombres'] for d in datos_grupos])
            
            # Totales para sin complementos - MEDIANA
            datos_totales_sin = self.calcular_medianas_efectivos_sb(df_actual, col_sin_comp)
            brecha_total_sin = ((datos_totales_sin['H'] - datos_totales_sin['M']) / datos_totales_sin['M'] * 100) if datos_totales_sin['M'] > 0 else 0
            
            # Totales para con complementos - MEDIANA
            datos_totales_con = self.calcular_medianas_efectivos_sb_complementos(df_actual, col_con_comp)
            brecha_total_con = ((datos_totales_con['H'] - datos_totales_con['M']) / datos_totales_con['M'] * 100) if datos_totales_con['M'] > 0 else 0
            
            # Totales para con extrasalariales - MEDIANA
            datos_totales_extra = self.calcular_medianas_efectivos_sb_complementos(df_actual, col_con_extra)
            brecha_total_extra = ((datos_totales_extra['H'] - datos_totales_extra['M']) / datos_totales_extra['M'] * 100) if datos_totales_extra['M'] > 0 else 0
            
            # Insertar totales al principio
            datos_grupos.insert(0, {
                'grupo': 'Totales',
                'n_mujeres': total_m,
                'n_hombres': total_h,
                'sin_comp_m': datos_totales_sin['M'],
                'sin_comp_h': datos_totales_sin['H'],
                'brecha_sin_comp': brecha_total_sin,
                'con_comp_m': datos_totales_con['M'],
                'con_comp_h': datos_totales_con['H'],
                'brecha_con_comp': brecha_total_con,
                'con_extra_m': datos_totales_extra['M'],
                'con_extra_h': datos_totales_extra['H'],
                'brecha_con_extra': brecha_total_extra
            })
        
        return datos_grupos
    
    def generar_datos_grupo_profesional_equiparado_mediana(self):
        """Genera los datos para la tabla de retribución por grupo profesional EQUIPARADO usando MEDIANA"""
        print("📊 Calculando datos equiparados por grupo profesional con MEDIANA...")
        
        # Filtrar datos actuales (sin "Ex" en primera columna)
        df_actual = self.data[self.data.iloc[:, 0] != 'Ex'].copy()
        
        # Obtener grupos profesionales únicos
        if 'Grupo profesional' not in df_actual.columns:
            print("⚠️ Columna 'Grupo profesional' no encontrada")
            return []
        
        grupos = df_actual['Grupo profesional'].unique()
        # Convertir a string para evitar problemas de ordenación con tipos mixtos
        grupos = [str(g) for g in grupos if pd.notna(g)]
        datos_grupos = []
        
        for grupo in sorted(grupos):
            # Filtrar datos del grupo (convertir también la columna original a string para comparación)
            df_grupo = df_actual[df_actual['Grupo profesional'].astype(str) == grupo]
            
            # Contar por género
            conteos = df_grupo['SEXO'].value_counts()
            n_mujeres = conteos.get('M', 0)
            n_hombres = conteos.get('H', 0)
            
            if n_mujeres == 0 and n_hombres == 0:
                continue
            
            # Calcular retribución SIN complementos (solo salario base equiparado) - MEDIANA
            col_sin_comp = 'salario_base_equiparado'
            if col_sin_comp in df_grupo.columns:
                # Filtrar solo registros con salario base efectivo > 0 como referencia
                col_sb_efectivo = "Salario base anual efectivo"
                if col_sb_efectivo in df_grupo.columns:
                    df_sin_comp = df_grupo[(df_grupo[col_sb_efectivo].notna()) & (df_grupo[col_sb_efectivo] > 0)]
                else:
                    df_sin_comp = df_grupo[df_grupo[col_sin_comp].notna() & (df_grupo[col_sin_comp] > 0)]
                
                if len(df_sin_comp) > 0:
                    medianas_sin_comp = df_sin_comp.groupby('SEXO')[col_sin_comp].median()
                    sin_comp_m = medianas_sin_comp.get('M', 0)
                    sin_comp_h = medianas_sin_comp.get('H', 0)
                    brecha_sin_comp = ((sin_comp_h - sin_comp_m) / sin_comp_m * 100) if sin_comp_m > 0 else 0
                else:
                    sin_comp_m = sin_comp_h = brecha_sin_comp = 0
            else:
                sin_comp_m = sin_comp_h = brecha_sin_comp = 0
            
            # Calcular retribución CON complementos equiparados - MEDIANA
            col_con_comp = 'sb_mas_comp_salariales_equiparado'
            if col_con_comp in df_grupo.columns:
                # Incluir todos los registros (incluir SB = 0)
                df_con_comp = df_grupo[df_grupo[col_con_comp].notna()]
                if len(df_con_comp) > 0:
                    medianas_con_comp = df_con_comp.groupby('SEXO')[col_con_comp].median()
                    con_comp_m = medianas_con_comp.get('M', 0)
                    con_comp_h = medianas_con_comp.get('H', 0)
                    brecha_con_comp = ((con_comp_h - con_comp_m) / con_comp_m * 100) if con_comp_m > 0 else 0
                else:
                    con_comp_m = con_comp_h = brecha_con_comp = 0
            else:
                con_comp_m = con_comp_h = brecha_con_comp = 0
            
            # Calcular retribución CON complementos + extrasalariales equiparados - MEDIANA
            col_con_extra = 'sb_mas_comp_total_equiparado'
            if col_con_extra in df_grupo.columns:
                df_con_extra = df_grupo[df_grupo[col_con_extra].notna()]
                if len(df_con_extra) > 0:
                    medianas_con_extra = df_con_extra.groupby('SEXO')[col_con_extra].median()
                    con_extra_m = medianas_con_extra.get('M', 0)
                    con_extra_h = medianas_con_extra.get('H', 0)
                    brecha_con_extra = ((con_extra_h - con_extra_m) / con_extra_m * 100) if con_extra_m > 0 else 0
                else:
                    con_extra_m = con_extra_h = brecha_con_extra = 0
            else:
                con_extra_m = con_extra_h = brecha_con_extra = 0
            
            datos_grupos.append({
                'grupo': grupo,
                'n_mujeres': n_mujeres,
                'n_hombres': n_hombres,
                'sin_comp_m': sin_comp_m,
                'sin_comp_h': sin_comp_h,
                'brecha_sin_comp': brecha_sin_comp,
                'con_comp_m': con_comp_m,
                'con_comp_h': con_comp_h,
                'brecha_con_comp': brecha_con_comp,
                'con_extra_m': con_extra_m,
                'con_extra_h': con_extra_h,
                'brecha_con_extra': brecha_con_extra
            })
        
        # Calcular totales usando MEDIANA para equiparados
        if datos_grupos:
            total_m = sum([d['n_mujeres'] for d in datos_grupos])
            total_h = sum([d['n_hombres'] for d in datos_grupos])
            
            # Totales para sin complementos equiparados - MEDIANA
            datos_totales_sin = self.calcular_medianas_equiparados_sb(df_actual, col_sin_comp)
            brecha_total_sin = ((datos_totales_sin['H'] - datos_totales_sin['M']) / datos_totales_sin['M'] * 100) if datos_totales_sin['M'] > 0 else 0
            
            # Totales para con complementos equiparados - MEDIANA
            datos_totales_con = self.calcular_medianas_equiparados_sb_complementos(df_actual, col_con_comp)
            brecha_total_con = ((datos_totales_con['H'] - datos_totales_con['M']) / datos_totales_con['M'] * 100) if datos_totales_con['M'] > 0 else 0
            
            # Totales para con extrasalariales equiparados - MEDIANA
            datos_totales_extra = self.calcular_medianas_equiparados_sb_complementos(df_actual, col_con_extra)
            brecha_total_extra = ((datos_totales_extra['H'] - datos_totales_extra['M']) / datos_totales_extra['M'] * 100) if datos_totales_extra['M'] > 0 else 0
            
            # Insertar totales al principio
            datos_grupos.insert(0, {
                'grupo': 'Totales',
                'n_mujeres': total_m,
                'n_hombres': total_h,
                'sin_comp_m': datos_totales_sin['M'],
                'sin_comp_h': datos_totales_sin['H'],
                'brecha_sin_comp': brecha_total_sin,
                'con_comp_m': datos_totales_con['M'],
                'con_comp_h': datos_totales_con['H'],
                'brecha_con_comp': brecha_total_con,
                'con_extra_m': datos_totales_extra['M'],
                'con_extra_h': datos_totales_extra['H'],
                'brecha_con_extra': brecha_total_extra
            })
        
        return datos_grupos
    
    def generar_datos_svpt_puesto_efectivo_mediana(self):
        """Genera los datos para la tabla de retribución por SVPT + Puesto de trabajo EFECTIVO combinados usando MEDIANA"""
        print("📊 Calculando datos efectivos por SVPT + Puesto de trabajo con MEDIANA...")
        
        # Filtrar datos actuales (sin "Ex" en primera columna)
        df_actual = self.data[self.data.iloc[:, 0] != 'Ex'].copy()
        
        # Verificar que ambas columnas existen
        if 'Nivel SVPT' not in df_actual.columns:
            print("❌ Error: No se encontró la columna 'Nivel SVPT'")
            return []
        
        if 'Puesto de trabajo' not in df_actual.columns:
            print("❌ Error: No se encontró la columna 'Puesto de trabajo'")
            return []
        
        # Crear una columna combinada de SVPT + Puesto de trabajo
        df_actual['SVPT_Puesto'] = df_actual['Nivel SVPT'].astype(str) + ' - ' + df_actual['Puesto de trabajo'].astype(str)
        
        # Obtener combinaciones únicas
        combinaciones = df_actual['SVPT_Puesto'].unique()
        combinaciones = [str(c) for c in combinaciones if pd.notna(c) and str(c) != 'nan - nan']
        datos_combinaciones = []
        
        # Columnas para valores efectivos
        col_sin_comp = 'Salario base efectivo Total'  # SB efectivo
        col_con_comp = 'Salario base anual + complementos Total'  # SB + Complementos efectivo
        col_con_extra = 'Salario base anual + complementos + Extrasalariales Total'  # SB + Complementos + Extrasalariales efectivo
        
        for combinacion in sorted(combinaciones):
            # Filtrar por la combinación específica
            df_comb = df_actual[df_actual['SVPT_Puesto'].astype(str) == str(combinacion)]
            
            # Conteos por género
            conteos = df_comb['SEXO'].value_counts()
            n_mujeres = conteos.get('M', 0)
            n_hombres = conteos.get('H', 0)
            
            # SIN complementos (solo SB efectivo) - solo personas con SB > 0 - MEDIANA
            datos_sin_comp = self.calcular_medianas_efectivos_sb(df_comb, col_sin_comp)
            sin_comp_m = datos_sin_comp['M']
            sin_comp_h = datos_sin_comp['H']
            brecha_sin_comp = ((sin_comp_h - sin_comp_m) / sin_comp_m * 100) if sin_comp_m > 0 else 0
            
            # CON complementos efectivos - todas las personas - MEDIANA
            datos_con_comp = self.calcular_medianas_efectivos_sb_complementos(df_comb, col_con_comp)
            con_comp_m = datos_con_comp['M']
            con_comp_h = datos_con_comp['H']
            brecha_con_comp = ((con_comp_h - con_comp_m) / con_comp_m * 100) if con_comp_m > 0 else 0
            
            # CON extrasalariales efectivos - todas las personas - MEDIANA
            datos_con_extra = self.calcular_medianas_efectivos_sb_complementos(df_comb, col_con_extra)
            con_extra_m = datos_con_extra['M']
            con_extra_h = datos_con_extra['H']
            brecha_con_extra = ((con_extra_h - con_extra_m) / con_extra_m * 100) if con_extra_m > 0 else 0
            
            datos_combinaciones.append({
                'svpt_puesto': combinacion,
                'n_mujeres': n_mujeres,
                'n_hombres': n_hombres,
                'sin_comp_m': sin_comp_m,
                'sin_comp_h': sin_comp_h,
                'brecha_sin_comp': brecha_sin_comp,
                'con_comp_m': con_comp_m,
                'con_comp_h': con_comp_h,
                'brecha_con_comp': brecha_con_comp,
                'con_extra_m': con_extra_m,
                'con_extra_h': con_extra_h,
                'brecha_con_extra': brecha_con_extra
            })
        
        # Calcular totales efectivos con MEDIANA
        if datos_combinaciones:
            total_m = sum([d['n_mujeres'] for d in datos_combinaciones])
            total_h = sum([d['n_hombres'] for d in datos_combinaciones])
            
            # Totales para sin complementos efectivos - MEDIANA
            datos_totales_sin = self.calcular_medianas_efectivos_sb(df_actual, col_sin_comp)
            brecha_total_sin = ((datos_totales_sin['H'] - datos_totales_sin['M']) / datos_totales_sin['M'] * 100) if datos_totales_sin['M'] > 0 else 0
            
            # Totales para con complementos efectivos - MEDIANA
            datos_totales_con = self.calcular_medianas_efectivos_sb_complementos(df_actual, col_con_comp)
            brecha_total_con = ((datos_totales_con['H'] - datos_totales_con['M']) / datos_totales_con['M'] * 100) if datos_totales_con['M'] > 0 else 0
            
            # Totales para con extrasalariales efectivos - MEDIANA
            datos_totales_extra = self.calcular_medianas_efectivos_sb_complementos(df_actual, col_con_extra)
            brecha_total_extra = ((datos_totales_extra['H'] - datos_totales_extra['M']) / datos_totales_extra['M'] * 100) if datos_totales_extra['M'] > 0 else 0
            
            # Insertar totales al principio
            datos_combinaciones.insert(0, {
                'svpt_puesto': 'Totales',
                'n_mujeres': total_m,
                'n_hombres': total_h,
                'sin_comp_m': datos_totales_sin['M'],
                'sin_comp_h': datos_totales_sin['H'],
                'brecha_sin_comp': brecha_total_sin,
                'con_comp_m': datos_totales_con['M'],
                'con_comp_h': datos_totales_con['H'],
                'brecha_con_comp': brecha_total_con,
                'con_extra_m': datos_totales_extra['M'],
                'con_extra_h': datos_totales_extra['H'],
                'brecha_con_extra': brecha_total_extra
            })
        
        return datos_combinaciones
    
    def generar_datos_svpt_puesto_equiparado_mediana(self):
        """Genera los datos para la tabla de retribución por SVPT + Puesto de trabajo EQUIPARADO combinados usando MEDIANA"""
        print("📊 Calculando datos equiparados por SVPT + Puesto de trabajo con MEDIANA...")
        
        # Filtrar datos actuales (sin "Ex" en primera columna)
        df_actual = self.data[self.data.iloc[:, 0] != 'Ex'].copy()
        
        # Verificar que ambas columnas existen
        if 'Nivel SVPT' not in df_actual.columns:
            print("❌ Error: No se encontró la columna 'Nivel SVPT'")
            return []
        
        if 'Puesto de trabajo' not in df_actual.columns:
            print("❌ Error: No se encontró la columna 'Puesto de trabajo'")
            return []
        
        # Crear una columna combinada de SVPT + Puesto de trabajo
        df_actual['SVPT_Puesto'] = df_actual['Nivel SVPT'].astype(str) + ' - ' + df_actual['Puesto de trabajo'].astype(str)
        
        # Obtener combinaciones únicas
        combinaciones = df_actual['SVPT_Puesto'].unique()
        combinaciones = [str(c) for c in combinaciones if pd.notna(c) and str(c) != 'nan - nan']
        datos_combinaciones = []
        
        # Columnas para valores equiparados
        col_sin_comp = 'salario_base_equiparado'  # SB equiparado
        col_con_comp = 'sb_mas_comp_salariales_equiparado'  # SB + Complementos equiparado
        col_con_extra = 'sb_mas_comp_total_equiparado'  # SB + Complementos + Extrasalariales equiparado
        
        for combinacion in sorted(combinaciones):
            # Filtrar por la combinación específica
            df_comb = df_actual[df_actual['SVPT_Puesto'].astype(str) == str(combinacion)]
            
            # Conteos por género
            conteos = df_comb['SEXO'].value_counts()
            n_mujeres = conteos.get('M', 0)
            n_hombres = conteos.get('H', 0)
            
            # SIN complementos (solo SB equiparado) - solo personas con SB > 0 - MEDIANA
            datos_sin_comp = self.calcular_medianas_equiparados_sb(df_comb, col_sin_comp)
            sin_comp_m = datos_sin_comp['M']
            sin_comp_h = datos_sin_comp['H']
            brecha_sin_comp = ((sin_comp_h - sin_comp_m) / sin_comp_m * 100) if sin_comp_m > 0 else 0
            
            # CON complementos equiparados - todas las personas - MEDIANA
            datos_con_comp = self.calcular_medianas_equiparados_sb_complementos(df_comb, col_con_comp)
            con_comp_m = datos_con_comp['M']
            con_comp_h = datos_con_comp['H']
            brecha_con_comp = ((con_comp_h - con_comp_m) / con_comp_m * 100) if con_comp_m > 0 else 0
            
            # CON extrasalariales equiparados - todas las personas - MEDIANA
            datos_con_extra = self.calcular_medianas_equiparados_sb_complementos(df_comb, col_con_extra)
            con_extra_m = datos_con_extra['M']
            con_extra_h = datos_con_extra['H']
            brecha_con_extra = ((con_extra_h - con_extra_m) / con_extra_m * 100) if con_extra_m > 0 else 0
            
            datos_combinaciones.append({
                'svpt_puesto': combinacion,
                'n_mujeres': n_mujeres,
                'n_hombres': n_hombres,
                'sin_comp_m': sin_comp_m,
                'sin_comp_h': sin_comp_h,
                'brecha_sin_comp': brecha_sin_comp,
                'con_comp_m': con_comp_m,
                'con_comp_h': con_comp_h,
                'brecha_con_comp': brecha_con_comp,
                'con_extra_m': con_extra_m,
                'con_extra_h': con_extra_h,
                'brecha_con_extra': brecha_con_extra
            })
        
        # Calcular totales equiparados con MEDIANA
        if datos_combinaciones:
            total_m = sum([d['n_mujeres'] for d in datos_combinaciones])
            total_h = sum([d['n_hombres'] for d in datos_combinaciones])
            
            # Totales para sin complementos equiparados - MEDIANA
            datos_totales_sin = self.calcular_medianas_equiparados_sb(df_actual, col_sin_comp)
            brecha_total_sin = ((datos_totales_sin['H'] - datos_totales_sin['M']) / datos_totales_sin['M'] * 100) if datos_totales_sin['M'] > 0 else 0
            
            # Totales para con complementos equiparados - MEDIANA
            datos_totales_con = self.calcular_medianas_equiparados_sb_complementos(df_actual, col_con_comp)
            brecha_total_con = ((datos_totales_con['H'] - datos_totales_con['M']) / datos_totales_con['M'] * 100) if datos_totales_con['M'] > 0 else 0
            
            # Totales para con extrasalariales equiparados - MEDIANA
            datos_totales_extra = self.calcular_medianas_equiparados_sb_complementos(df_actual, col_con_extra)
            brecha_total_extra = ((datos_totales_extra['H'] - datos_totales_extra['M']) / datos_totales_extra['M'] * 100) if datos_totales_extra['M'] > 0 else 0
            
            # Insertar totales al principio
            datos_combinaciones.insert(0, {
                'svpt_puesto': 'Totales',
                'n_mujeres': total_m,
                'n_hombres': total_h,
                'sin_comp_m': datos_totales_sin['M'],
                'sin_comp_h': datos_totales_sin['H'],
                'brecha_sin_comp': brecha_total_sin,
                'con_comp_m': datos_totales_con['M'],
                'con_comp_h': datos_totales_con['H'],
                'brecha_con_comp': brecha_total_con,
                'con_extra_m': datos_totales_extra['M'],
                'con_extra_h': datos_totales_extra['H'],
                'brecha_con_extra': brecha_total_extra
            })
        
        return datos_combinaciones
    
    def generar_datos_por_escala_svpt_mediana(self, escala, tipo="Efectivo"):
        """
        Genera los datos para una escala SVPT específica (E1, E2, etc.) usando MEDIANA
        """
        # Filtrar datos actuales (sin "Ex" en primera columna)
        df_actual = self.data[self.data.iloc[:, 0] != 'Ex'].copy()
        
        # Verificar que existan las columnas necesarias
        if 'Nivel SVPT' not in df_actual.columns:
            print(f"❌ Error: No se encontró la columna 'Nivel SVPT' para escala {escala}")
            return []
        
        if 'Puesto de trabajo' not in df_actual.columns:
            print(f"❌ Error: No se encontró la columna 'Puesto de trabajo' para escala {escala}")
            return []
        
        # Filtrar por la escala específica
        df_escala = df_actual[df_actual['Nivel SVPT'].astype(str) == str(escala)]
        
        if len(df_escala) == 0:
            print(f"⚠️ No hay datos para la escala {escala}")
            return []
        
        # Obtener puestos únicos dentro de esta escala
        puestos = df_escala['Puesto de trabajo'].unique()
        puestos = [str(p) for p in puestos if pd.notna(p)]
        datos_puestos = []
        
        # Columnas según el tipo
        if tipo == "Efectivo":
            col_sin_comp = 'Salario base efectivo Total'
            col_con_comp = 'Salario base anual + complementos Total'
            col_con_extra = 'Salario base anual + complementos + Extrasalariales Total'
        else:  # Equiparado
            col_sin_comp = 'salario_base_equiparado'
            col_con_comp = 'sb_mas_comp_salariales_equiparado'
            col_con_extra = 'sb_mas_comp_total_equiparado'
        
        for puesto in sorted(puestos):
            # Filtrar datos del puesto dentro de la escala
            df_puesto = df_escala[df_escala['Puesto de trabajo'].astype(str) == puesto]
            
            # Contar por género
            conteos = df_puesto['SEXO'].value_counts()
            n_mujeres = conteos.get('M', 0)
            n_hombres = conteos.get('H', 0)
            
            if n_mujeres == 0 and n_hombres == 0:
                continue
            
            # Calcular retribución CON complementos + extrasalariales usando MEDIANA
            if tipo == "Efectivo":
                # Para efectivo, incluir todas las personas
                df_calc = df_puesto[df_puesto[col_con_extra].notna()]
            else:
                # Para equiparado, incluir todas las personas
                df_calc = df_puesto[df_puesto[col_con_extra].notna()]
            
            if len(df_calc) > 0:
                medianas = df_calc.groupby('SEXO')[col_con_extra].median()  # USAR MEDIANA en lugar de mean
                con_extra_m = medianas.get('M', 0)
                con_extra_h = medianas.get('H', 0)
            else:
                con_extra_m = con_extra_h = 0
            
            # Calcular mediana de puntos para este puesto
            puntos_validos = df_puesto['Puntos'].dropna() if 'Puntos' in df_puesto.columns else []
            mediana_puntos = puntos_validos.median() if len(puntos_validos) > 0 else 0  # USAR MEDIANA en lugar de mean
            
            datos_puestos.append({
                'puesto': puesto,
                'escala': escala,
                'n_mujeres': n_mujeres,
                'n_hombres': n_hombres,
                'con_extra_m': con_extra_m,
                'con_extra_h': con_extra_h,
                'puntos': mediana_puntos
            })
        
        return datos_puestos
    
    def generar_datos_por_nivel_mediana(self, tipo="Efectivo"):
        """
        Genera los datos de retribución por nivel usando MEDIANA
        """
        # Filtrar datos actuales (sin "Ex" en primera columna)
        df_actual = self.data[self.data.iloc[:, 0] != 'Ex'].copy()
        
        # Buscar la columna "Nivel Convenio Colectivo" (puede tener espacios)
        columna_nivel = None
        posibles_nombres = ["Nivel Convenio Colectivo", "Nivel Convenio Colectivo "]
        
        for nombre in posibles_nombres:
            if nombre in df_actual.columns:
                columna_nivel = nombre
                break
        
        if columna_nivel is None:
            print(f"❌ Error: No se encontró la columna 'Nivel Convenio Colectivo'")
            print(f"📋 Columnas disponibles que contienen 'Nivel': {[col for col in df_actual.columns if 'Nivel' in col]}")
            return []
        
        print(f"📋 Usando columna de nivel para MEDIANA: {columna_nivel}")
        
        # Obtener niveles únicos
        niveles = df_actual[columna_nivel].dropna().unique()
        niveles_validos = []
        
        for nivel in niveles:
            nivel_str = str(nivel).strip()
            if nivel_str and nivel_str != 'nan' and len(nivel_str) > 0:
                niveles_validos.append(nivel_str)
        
        print(f"📊 Niveles encontrados para MEDIANA: {niveles_validos}")
        
        # Ordenar niveles de forma natural
        try:
            def ordenar_nivel(nivel_str):
                # Extraer números del nivel para ordenamiento natural
                import re
                match = re.search(r'(\d+)', nivel_str)
                if match:
                    return int(match.group(1))
                return float('inf')
            
            niveles_ordenados = sorted(niveles_validos, key=ordenar_nivel)
        except:
            niveles_ordenados = sorted(niveles_validos)
        
        datos_niveles = []
        
        # Columnas según el tipo
        if tipo == "Efectivo":
            col_sin_comp = 'Salario base efectivo Total'
            col_con_comp = 'Salario base anual + complementos Total'
            col_con_extra = 'Salario base anual + complementos + Extrasalariales Total'
        else:  # Equiparado
            col_sin_comp = 'salario_base_equiparado'
            col_con_comp = 'sb_mas_comp_salariales_equiparado'
            col_con_extra = 'sb_mas_comp_total_equiparado'
        
        # Procesar cada nivel
        for nivel in niveles_ordenados:
            # Filtrar datos del nivel
            df_nivel = df_actual[df_actual[columna_nivel].astype(str) == nivel]
            
            if len(df_nivel) == 0:
                continue
            
            # Contar por género
            conteos = df_nivel['SEXO'].value_counts()
            n_mujeres = conteos.get('M', 0)
            n_hombres = conteos.get('H', 0)
            
            # Calcular MEDIANAS de retribución (sin complementos)
            if tipo == "Efectivo":
                df_calc_sin = df_nivel[df_nivel[col_sin_comp].notna() & (df_nivel[col_sin_comp] > 0)]
            else:
                df_calc_sin = df_nivel[df_nivel[col_sin_comp].notna()]
            
            if len(df_calc_sin) > 0:
                medianas_sin = df_calc_sin.groupby('SEXO')[col_sin_comp].median()  # USAR MEDIANA
                sin_comp_m = medianas_sin.get('M', 0)
                sin_comp_h = medianas_sin.get('H', 0)
            else:
                sin_comp_m = sin_comp_h = 0
            
            # Calcular MEDIANAS CON complementos
            if tipo == "Efectivo":
                df_calc_con = df_nivel[df_nivel[col_con_comp].notna() & (df_nivel[col_con_comp] > 0)]
            else:
                df_calc_con = df_nivel[df_nivel[col_con_comp].notna()]
            
            if len(df_calc_con) > 0:
                medianas_con = df_calc_con.groupby('SEXO')[col_con_comp].median()  # USAR MEDIANA
                con_comp_m = medianas_con.get('M', 0)
                con_comp_h = medianas_con.get('H', 0)
            else:
                con_comp_m = con_comp_h = 0
            
            # Calcular MEDIANAS CON complementos + extrasalariales
            if tipo == "Efectivo":
                df_calc_extra = df_nivel[df_nivel[col_con_extra].notna() & (df_nivel[col_con_extra] > 0)]
            else:
                df_calc_extra = df_nivel[df_nivel[col_con_extra].notna()]
            
            if len(df_calc_extra) > 0:
                medianas_extra = df_calc_extra.groupby('SEXO')[col_con_extra].median()  # USAR MEDIANA
                con_extra_m = medianas_extra.get('M', 0)
                con_extra_h = medianas_extra.get('H', 0)
            else:
                con_extra_m = con_extra_h = 0
            
            datos_niveles.append({
                'nivel': nivel,
                'n_mujeres': n_mujeres,
                'n_hombres': n_hombres,
                'sin_comp_m': sin_comp_m,
                'sin_comp_h': sin_comp_h,
                'con_comp_m': con_comp_m,
                'con_comp_h': con_comp_h,
                'con_extra_m': con_extra_m,
                'con_extra_h': con_extra_h
            })
        
        # Calcular totales usando MEDIANA del dataset completo
        if datos_niveles:
            total_m = sum(d['n_mujeres'] for d in datos_niveles)
            total_h = sum(d['n_hombres'] for d in datos_niveles)
            
            # Calcular medianas totales usando las funciones de mediana
            if tipo == "Efectivo":
                datos_totales_sin = self.calcular_medianas_efectivos_sb(df_actual, col_sin_comp)
                datos_totales_con = self.calcular_medianas_efectivos_sb_complementos(df_actual, col_con_comp)
                datos_totales_extra = self.calcular_medianas_efectivos_sb_complementos(df_actual, col_con_extra)
            else:
                datos_totales_sin = self.calcular_medianas_equiparados_sb(df_actual, col_sin_comp)
                datos_totales_con = self.calcular_medianas_equiparados_sb_complementos(df_actual, col_con_comp)
                datos_totales_extra = self.calcular_medianas_equiparados_sb_complementos(df_actual, col_con_extra)
            
            datos_niveles.append({
                'nivel': 'Totales',
                'n_mujeres': total_m,
                'n_hombres': total_h,
                'sin_comp_m': datos_totales_sin['M'],
                'sin_comp_h': datos_totales_sin['H'],
                'con_comp_m': datos_totales_con['M'],
                'con_comp_h': datos_totales_con['H'],
                'con_extra_m': datos_totales_extra['M'],
                'con_extra_h': datos_totales_extra['H']
            })
        
        return datos_niveles

    # ==================== FIN FUNCIONES PARA ANÁLISIS CON MEDIANA ====================
    
    def crear_grafico_barras_grupo_profesional(self, datos_grupos, tipo="Efectivo"):
        """
        Crea un gráfico de barras horizontales por grupo profesional
        Muestra salarios CON complementos + extrasalariales por género
        tipo: "Efectivo" o "Equiparado" para personalizar el título
        """
        import matplotlib.pyplot as plt
        import numpy as np
        
        # Filtrar datos válidos (excluir totales y grupos sin datos)
        datos_validos = [d for d in datos_grupos if d['grupo'] != 'Totales' and d['con_extra_m'] > 0 and d['con_extra_h'] > 0]
        
        if not datos_validos:
            print("No hay datos válidos para el gráfico de barras")
            return None
        
        # Preparar datos
        grupos = [d['grupo'] for d in datos_validos]
        salarios_m = [d['con_extra_m'] for d in datos_validos]
        salarios_h = [d['con_extra_h'] for d in datos_validos]
        
        # Crear figura
        fig, ax = plt.subplots(figsize=(12, 8))
        fig.patch.set_facecolor('white')
        
        # Configurar posiciones de las barras
        y_pos = np.arange(len(grupos))
        bar_height = 0.35
        
        # Crear barras horizontales con los mismos colores que los gráficos donut
        bars_m = ax.barh(y_pos - bar_height/2, salarios_m, bar_height, 
                        label='Mujeres', color=self.colores_genero['M'], alpha=0.8)  # Azul para mujeres
        bars_h = ax.barh(y_pos + bar_height/2, salarios_h, bar_height, 
                        label='Hombres', color=self.colores_genero['H'], alpha=0.8)   # Rojo para hombres
        
        # Agregar valores en las barras
        for i, (bar_m, bar_h, sal_m, sal_h) in enumerate(zip(bars_m, bars_h, salarios_m, salarios_h)):
            # Valor para mujeres
            ax.text(bar_m.get_width() + max(salarios_m + salarios_h) * 0.01, 
                   bar_m.get_y() + bar_m.get_height()/2, 
                   f'{self.formato_numero_es(sal_m, 2)} €', 
                   va='center', ha='left', fontsize=9, fontweight='bold')
            
            # Valor para hombres  
            ax.text(bar_h.get_width() + max(salarios_m + salarios_h) * 0.01, 
                   bar_h.get_y() + bar_h.get_height()/2, 
                   f'{self.formato_numero_es(sal_h, 2)} €', 
                   va='center', ha='left', fontsize=9, fontweight='bold')
        
        # Configurar ejes
        ax.set_yticks(y_pos)
        ax.set_yticklabels(grupos)
        ax.invert_yaxis()  # Para mostrar el primer grupo arriba
        
        # Configurar etiquetas y título (personalizado según el tipo)
        ax.set_xlabel('Salario (€)', fontsize=12, fontweight='bold')
        ax.set_ylabel('Grupos profesionales', fontsize=12, fontweight='bold')
        ax.set_title(f'Salarios Medios CON Complementos + ES por Grupo Profesional {tipo}', 
                    fontsize=14, fontweight='bold', pad=20)
        
        # Configurar leyenda
        ax.legend(loc='lower right', fontsize=11)
        
        # Configurar grid
        ax.grid(True, axis='x', alpha=0.3)
        ax.set_axisbelow(True)
        
        # Ajustar márgenes
        plt.tight_layout()
        
        return fig

    def crear_grafico_donut(self, datos_genero, titulo, subtitulo="", formato_moneda=True):
        """
        Crea un gráfico de donut compacto con la brecha salarial en el centro
        Optimizado para mostrar dos gráficos por fila sin solapamientos
        """
        # Preparar datos
        valores = [datos_genero['H'], datos_genero['M']]
        etiquetas = ['Hombres', 'Mujeres']
        colores = [self.colores_genero['H'], self.colores_genero['M']]
        
        # Calcular la brecha salarial
        if datos_genero['M'] > 0:
            brecha = ((datos_genero['H'] - datos_genero['M']) / datos_genero['M']) * 100
        else:
            brecha = 0
            
        # Configurar la figura más pequeña y con mejor spacing
        # Crear figura compacta con configuración base
        fig, ax = plt.subplots(figsize=(5.0, 4.0))
        self.configurar_grafico_base(fig, ax)
        
        # Crear el gráfico de donut con etiquetas más limpias
        wedges, texts, autotexts = ax.pie(valores, labels=None, autopct='%1.1f%%',  # Sin labels en el pie
                                          colors=colores, startangle=90, 
                                          wedgeprops=dict(width=0.4, edgecolor='white', linewidth=2),
                                          pctdistance=0.82,  # Porcentajes más lejos del centro
                                          textprops={'fontsize': 10, 'fontweight': 'bold'})
        
        # Mejorar la apariencia de los porcentajes
        for autotext in autotexts:
            autotext.set_color('white')
            autotext.set_fontweight('bold')
            autotext.set_fontsize(10)
            # Fondo semitransparente para mejor legibilidad
            autotext.set_bbox(dict(boxstyle="round,pad=0.15", facecolor='black', alpha=0.6))
        
        # Añadir texto de brecha en el centro MÁS LIMPIO
        color_brecha = '#e74c3c' if brecha > 0 else '#27ae60' if brecha < 0 else '#95a5a6'
        
        ax.text(0, 0.05, 'Brecha Salarial',
                horizontalalignment='center', verticalalignment='center',
                fontsize=9, fontweight='bold', color='#2c3e50')
        
        # Mostrar brecha en valor absoluto (el color ya indica la dirección)
        porcentaje_texto = f'{self.formato_brecha_es(brecha, 2)}%'
        ax.text(0, -0.10, porcentaje_texto, 
                horizontalalignment='center', verticalalignment='center',
                fontsize=14, fontweight='bold', color=color_brecha)
        
        # Aplicar título estandarizado (ajustado para donut compacto)
        ax.set_title(titulo, 
                    fontsize=self.config_graficos['titulo']['fontsize'] - 4,  # Más pequeño para donut
                    fontweight=self.config_graficos['titulo']['fontweight'], 
                    color=self.config_graficos['titulo']['color'],
                    pad=10)
        # Eliminamos el subtítulo para mejor legibilidad
        
        # Leyenda EXTERNA para evitar solapamientos
        if formato_moneda:
            leyenda_labels = [f'{etiqueta}: {self.formato_numero_es(valor, 0)}€' for etiqueta, valor in zip(etiquetas, valores)]
        else:
            leyenda_labels = [f'{etiqueta}: {self.formato_numero_es(valor, 2)}' for etiqueta, valor in zip(etiquetas, valores)]
        
        # Leyenda fuera del gráfico para evitar solapamientos
        ax.legend(wedges, leyenda_labels, 
                  loc="center left", bbox_to_anchor=(1.05, 0.5),  # Más alejada
                  fontsize=8)
        
        # Ajustar el aspecto
        ax.axis('equal')
        plt.tight_layout()
        
        return fig

    def generar_datos_por_nivel(self, tipo="Efectivo"):
        """
        Genera los datos de retribución promedio por nivel
        """
        # Filtrar datos actuales (sin "Ex" en primera columna)
        df_actual = self.data[self.data.iloc[:, 0] != 'Ex'].copy()
        
        # Buscar la columna "Nivel Convenio Colectivo" (puede tener espacios)
        columna_nivel = None
        posibles_nombres = ["Nivel Convenio Colectivo", "Nivel Convenio Colectivo "]
        
        for nombre in posibles_nombres:
            if nombre in df_actual.columns:
                columna_nivel = nombre
                break
        
        if columna_nivel is None:
            print(f"❌ Error: No se encontró la columna 'Nivel Convenio Colectivo'")
            print(f"📋 Columnas disponibles que contienen 'Nivel': {[col for col in df_actual.columns if 'Nivel' in col]}")
            return []
        
        print(f"📋 Usando columna de nivel: {columna_nivel}")
        
        # Obtener niveles únicos
        niveles = df_actual[columna_nivel].dropna().unique()
        niveles_validos = []
        
        for nivel in niveles:
            nivel_str = str(nivel).strip()
            if nivel_str and nivel_str != 'nan' and len(nivel_str) > 0:
                niveles_validos.append(nivel_str)
        
        print(f"📊 Niveles encontrados: {niveles_validos}")
        
        # Ordenar niveles de forma natural
        try:
            def ordenar_nivel(nivel_str):
                # Extraer números del nivel para ordenamiento natural
                import re
                match = re.search(r'(\d+)', nivel_str)
                if match:
                    return int(match.group(1))
                return float('inf')
            
            niveles_ordenados = sorted(niveles_validos, key=ordenar_nivel)
        except:
            niveles_ordenados = sorted(niveles_validos)
        
        datos_niveles = []
        
        # Columnas según el tipo
        if tipo == "Efectivo":
            col_sin_comp = 'Salario base efectivo Total'
            col_con_comp = 'Salario base anual + complementos Total'
            col_con_extra = 'Salario base anual + complementos + Extrasalariales Total'
        else:  # Equiparado
            col_sin_comp = 'salario_base_equiparado'
            col_con_comp = 'sb_mas_comp_salariales_equiparado'
            col_con_extra = 'sb_mas_comp_total_equiparado'
        
        # Procesar cada nivel
        for nivel in niveles_ordenados:
            # Filtrar datos del nivel
            df_nivel = df_actual[df_actual[columna_nivel].astype(str) == nivel]
            
            if len(df_nivel) == 0:
                continue
            
            # Contar por género
            conteos = df_nivel['SEXO'].value_counts()
            n_mujeres = conteos.get('M', 0)
            n_hombres = conteos.get('H', 0)
            
            # Calcular promedios de retribución (sin complementos)
            if tipo == "Efectivo":
                df_calc_sin = df_nivel[df_nivel[col_sin_comp].notna() & (df_nivel[col_sin_comp] > 0)]
            else:
                df_calc_sin = df_nivel[df_nivel[col_sin_comp].notna()]
            
            if len(df_calc_sin) > 0:
                promedios_sin = df_calc_sin.groupby('SEXO')[col_sin_comp].mean()
                sin_comp_m = promedios_sin.get('M', 0)
                sin_comp_h = promedios_sin.get('H', 0)
            else:
                sin_comp_m = sin_comp_h = 0
            
            # Calcular promedios CON complementos
            if tipo == "Efectivo":
                df_calc_con = df_nivel[df_nivel[col_con_comp].notna() & (df_nivel[col_con_comp] > 0)]
            else:
                df_calc_con = df_nivel[df_nivel[col_con_comp].notna()]
            
            if len(df_calc_con) > 0:
                promedios_con = df_calc_con.groupby('SEXO')[col_con_comp].mean()
                con_comp_m = promedios_con.get('M', 0)
                con_comp_h = promedios_con.get('H', 0)
            else:
                con_comp_m = con_comp_h = 0
            
            # Calcular promedios CON complementos + extrasalariales
            if tipo == "Efectivo":
                df_calc_extra = df_nivel[df_nivel[col_con_extra].notna() & (df_nivel[col_con_extra] > 0)]
            else:
                df_calc_extra = df_nivel[df_nivel[col_con_extra].notna()]
            
            if len(df_calc_extra) > 0:
                promedios_extra = df_calc_extra.groupby('SEXO')[col_con_extra].mean()
                con_extra_m = promedios_extra.get('M', 0)
                con_extra_h = promedios_extra.get('H', 0)
            else:
                con_extra_m = con_extra_h = 0
            
            datos_niveles.append({
                'nivel': nivel,
                'n_mujeres': n_mujeres,
                'n_hombres': n_hombres,
                'sin_comp_m': sin_comp_m,
                'sin_comp_h': sin_comp_h,
                'con_comp_m': con_comp_m,
                'con_comp_h': con_comp_h,
                'con_extra_m': con_extra_m,
                'con_extra_h': con_extra_h
            })
        
        # Calcular totales
        if datos_niveles:
            total_m = sum(d['n_mujeres'] for d in datos_niveles)
            total_h = sum(d['n_hombres'] for d in datos_niveles)
            
            # Promedios ponderados para totales
            if total_m > 0:
                total_sin_comp_m = sum(d['sin_comp_m'] * d['n_mujeres'] for d in datos_niveles) / total_m if total_m > 0 else 0
                total_con_comp_m = sum(d['con_comp_m'] * d['n_mujeres'] for d in datos_niveles) / total_m if total_m > 0 else 0
                total_con_extra_m = sum(d['con_extra_m'] * d['n_mujeres'] for d in datos_niveles) / total_m if total_m > 0 else 0
            else:
                total_sin_comp_m = total_con_comp_m = total_con_extra_m = 0
            
            if total_h > 0:
                total_sin_comp_h = sum(d['sin_comp_h'] * d['n_hombres'] for d in datos_niveles) / total_h if total_h > 0 else 0
                total_con_comp_h = sum(d['con_comp_h'] * d['n_hombres'] for d in datos_niveles) / total_h if total_h > 0 else 0
                total_con_extra_h = sum(d['con_extra_h'] * d['n_hombres'] for d in datos_niveles) / total_h if total_h > 0 else 0
            else:
                total_sin_comp_h = total_con_comp_h = total_con_extra_h = 0
            
            datos_niveles.append({
                'nivel': 'Totales',
                'n_mujeres': total_m,
                'n_hombres': total_h,
                'sin_comp_m': total_sin_comp_m,
                'sin_comp_h': total_sin_comp_h,
                'con_comp_m': total_con_comp_m,
                'con_comp_h': total_con_comp_h,
                'con_extra_m': total_con_extra_m,
                'con_extra_h': total_con_extra_h
            })
        
        return datos_niveles

    def crear_tabla_por_nivel(self, doc, datos_niveles):
        """Crea la tabla de retribución por nivel siguiendo el formato estándar"""
        from docx.shared import Cm, Pt
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.shared import OxmlElement, qn
        
        if not datos_niveles:
            doc.add_paragraph("No hay datos disponibles para generar la tabla por nivel.")
            return
        
        # Crear tabla con cabeceras según el formato estándar
        num_filas = len(datos_niveles) + 1  # +1 para cabecera
        table = doc.add_table(rows=num_filas, cols=12)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        
        # Definir cabeceras siguiendo el formato de grupo profesional
        headers = [
            'Nivel', 'Nº M', 'Nº H',
            'Retribución Promedio SIN Complementos (Mujeres)', 'Retribución Promedio SIN Complementos (Hombres)', 'Brecha Salarial SIN Complementos',
            'Retribución Promedio CON Complementos (Mujeres)', 'Retribución Promedio CON Complementos (Hombres)', 'Brecha Salarial CON Complementos',
            'Retribución Promedio CON Complementos ES (Mujeres)', 'Retribución Promedio CON Complementos ES (Hombres)', 'Brecha Salarial CON Complementos ES'
        ]
        
        # Configurar cabeceras
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            if i < len(header_row.cells):
                cell = header_row.cells[i]
                cell.text = header
                # Aplicar formato de cabecera
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(8)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Llenar datos
        for row_idx, datos in enumerate(datos_niveles, 1):
            if row_idx < len(table.rows):
                data_row = table.rows[row_idx]
                
                # Calcular brechas
                brecha_sin = self.calcular_brecha_entre_valores(datos['sin_comp_h'], datos['sin_comp_m'])
                brecha_con = self.calcular_brecha_entre_valores(datos['con_comp_h'], datos['con_comp_m'])
                brecha_extra = self.calcular_brecha_entre_valores(datos['con_extra_h'], datos['con_extra_m'])
                
                # Preparar valores
                values = [
                    str(datos['nivel']),
                    str(datos['n_mujeres']),
                    str(datos['n_hombres']),
                    f"{self.formato_numero_es(datos['sin_comp_m'], 2)} €" if datos['sin_comp_m'] > 0 else "-",
                    f"{self.formato_numero_es(datos['sin_comp_h'], 2)} €" if datos['sin_comp_h'] > 0 else "-",
                    f"{self.formato_brecha_es(brecha_sin, 2)}%" if datos['sin_comp_m'] > 0 and datos['sin_comp_h'] > 0 else "-",
                    f"{self.formato_numero_es(datos['con_comp_m'], 2)} €" if datos['con_comp_m'] > 0 else "-",
                    f"{self.formato_numero_es(datos['con_comp_h'], 2)} €" if datos['con_comp_h'] > 0 else "-",
                    f"{self.formato_brecha_es(brecha_con, 2)}%" if datos['con_comp_m'] > 0 and datos['con_comp_h'] > 0 else "-",
                    f"{self.formato_numero_es(datos['con_extra_m'], 2)} €" if datos['con_extra_m'] > 0 else "-",
                    f"{self.formato_numero_es(datos['con_extra_h'], 2)} €" if datos['con_extra_h'] > 0 else "-",
                    f"{self.formato_brecha_es(brecha_extra, 2)}%" if datos['con_extra_m'] > 0 and datos['con_extra_h'] > 0 else "-"
                ]
                
                # Llenar celdas y aplicar colores
                for i, value in enumerate(values):
                    if i < len(data_row.cells):
                        cell = data_row.cells[i]
                        cell.text = value
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph.runs:
                                run.font.size = Pt(8)
                        
                        # Aplicar colores según las reglas de negocio
                        es_total = (datos['nivel'] == 'Totales')
                        self.aplicar_color_celda_nivel(cell, i, datos, es_total)
        
        # Agregar espacio después de la tabla
        doc.add_paragraph()
        
        return table

    def aplicar_color_celda_nivel(self, cell, columna_index, datos, es_total):
        """Aplica colores al TEXTO de las celdas de la tabla por nivel siguiendo el patrón estándar"""
        # Colores para el texto - mismo esquema que tabla principal
        ROJO = RGBColor(234, 93, 65)   # Cuando favorece a hombres (#ea5d41)
        AZUL = RGBColor(30, 67, 137)   # Cuando favorece a mujeres (#1e4389)
        NEGRO = RGBColor(0, 0, 0)      # Para fila de totales
        
        # Destacar fila de totales con texto negro normal
        if es_total:
            self.colorear_texto_celda(cell, NEGRO)
            return
        
        # Aplicar colores a las columnas de salarios y brechas
        if columna_index == 3:  # Sin complementos - Mujeres
            if datos['sin_comp_m'] < datos['sin_comp_h'] and datos['sin_comp_m'] > 0 and datos['sin_comp_h'] > 0:
                self.colorear_texto_celda(cell, ROJO)
            elif datos['sin_comp_m'] > 0:
                self.colorear_texto_celda(cell, AZUL)
                
        elif columna_index == 4:  # Sin complementos - Hombres  
            if datos['sin_comp_h'] > datos['sin_comp_m'] and datos['sin_comp_m'] > 0 and datos['sin_comp_h'] > 0:
                self.colorear_texto_celda(cell, AZUL)
            elif datos['sin_comp_h'] > 0:
                self.colorear_texto_celda(cell, ROJO)
                
        elif columna_index == 5:  # Brecha Sin complementos
            brecha_sin = self.calcular_brecha_entre_valores(datos['sin_comp_h'], datos['sin_comp_m'])
            if brecha_sin is not None:
                if brecha_sin > 0:  # Favorable a hombres
                    self.colorear_texto_celda(cell, ROJO)
                else:  # Favorable a mujeres
                    self.colorear_texto_celda(cell, AZUL)
                    
        elif columna_index == 6:  # Con complementos - Mujeres
            if datos['con_comp_m'] < datos['con_comp_h'] and datos['con_comp_m'] > 0 and datos['con_comp_h'] > 0:
                self.colorear_texto_celda(cell, ROJO)
            elif datos['con_comp_m'] > 0:
                self.colorear_texto_celda(cell, AZUL)
                
        elif columna_index == 7:  # Con complementos - Hombres
            if datos['con_comp_h'] > datos['con_comp_m'] and datos['con_comp_m'] > 0 and datos['con_comp_h'] > 0:
                self.colorear_texto_celda(cell, AZUL)
            elif datos['con_comp_h'] > 0:
                self.colorear_texto_celda(cell, ROJO)
                
        elif columna_index == 8:  # Brecha Con complementos
            brecha_con = self.calcular_brecha_entre_valores(datos['con_comp_h'], datos['con_comp_m'])
            if brecha_con is not None:
                if brecha_con > 0:  # Favorable a hombres
                    self.colorear_texto_celda(cell, ROJO)
                else:  # Favorable a mujeres
                    self.colorear_texto_celda(cell, AZUL)
                    
        elif columna_index == 9:  # Con complementos + ES - Mujeres
            if datos['con_extra_m'] < datos['con_extra_h'] and datos['con_extra_m'] > 0 and datos['con_extra_h'] > 0:
                self.colorear_texto_celda(cell, ROJO)
            elif datos['con_extra_m'] > 0:
                self.colorear_texto_celda(cell, AZUL)
                
        elif columna_index == 10:  # Con complementos + ES - Hombres
            if datos['con_extra_h'] > datos['con_extra_m'] and datos['con_extra_m'] > 0 and datos['con_extra_h'] > 0:
                self.colorear_texto_celda(cell, AZUL)
            elif datos['con_extra_h'] > 0:
                self.colorear_texto_celda(cell, ROJO)
                
        elif columna_index == 11:  # Brecha Con complementos + ES
            brecha_extra = self.calcular_brecha_entre_valores(datos['con_extra_h'], datos['con_extra_m'])
            if brecha_extra is not None:
                if brecha_extra > 0:  # Favorable a hombres
                    self.colorear_texto_celda(cell, ROJO)
                else:  # Favorable a mujeres
                    self.colorear_texto_celda(cell, AZUL)

    def crear_grafico_barras_por_nivel(self, datos_niveles, tipo="Efectivo"):
        """
        Crea un gráfico de barras horizontales por nivel
        Muestra salarios CON complementos + extrasalariales por género
        """
        # Filtrar solo los niveles (no totales)
        datos_validos = [d for d in datos_niveles if d['nivel'] != 'Totales']
        
        if not datos_validos:
            print("No hay datos válidos para el gráfico de barras por nivel")
            return None
        
        # Preparar datos - usar CON complementos + extrasalariales
        niveles = [d['nivel'] for d in datos_validos]
        salarios_m = [d['con_extra_m'] if d['con_extra_m'] > 0 else 0 for d in datos_validos]
        salarios_h = [d['con_extra_h'] if d['con_extra_h'] > 0 else 0 for d in datos_validos]
        
        # Crear figura con tamaño estandarizado
        figsize = self.calcular_tamaño_grafico(len(niveles), 'barra_horizontal')
        fig, ax = plt.subplots(figsize=figsize)
        self.configurar_grafico_base(fig, ax)
        
        # Configurar posiciones de las barras
        y_pos = np.arange(len(niveles))
        bar_height = 0.35
        
        # Crear barras horizontales
        bars_m = ax.barh(y_pos - bar_height/2, salarios_m, bar_height, 
                        label='Mujeres', color=self.colores_genero['M'], alpha=0.8)
        bars_h = ax.barh(y_pos + bar_height/2, salarios_h, bar_height, 
                        label='Hombres', color=self.colores_genero['H'], alpha=0.8)
        
        # Agregar valores en las barras
        for i, (bar_m, bar_h, sal_m, sal_h) in enumerate(zip(bars_m, bars_h, salarios_m, salarios_h)):
            if sal_m > 0:
                ax.text(bar_m.get_width() + max(salarios_m + salarios_h) * 0.01, 
                       bar_m.get_y() + bar_m.get_height()/2, 
                       f'{self.formato_numero_es(sal_m, 0)} €', 
                       ha='left', va='center', fontsize=10, fontweight='bold')
            
            if sal_h > 0:
                ax.text(bar_h.get_width() + max(salarios_m + salarios_h) * 0.01, 
                       bar_h.get_y() + bar_h.get_height()/2, 
                       f'{self.formato_numero_es(sal_h, 0)} €', 
                       ha='left', va='center', fontsize=10, fontweight='bold')
        
        # Configurar ejes
        ax.set_yticks(y_pos)
        ax.set_yticklabels(niveles, fontsize=12)
        ax.set_xlabel('Salario Medio CON Complementos + ES (€)', fontsize=14, fontweight='bold')
        ax.set_ylabel('Nivel', fontsize=14, fontweight='bold')
        
        # Aplicar título estandarizado
        titulo = f'Salarios Medios CON Complementos + ES por Nivel {tipo}'
        self.aplicar_estilo_titulo(ax, titulo)
        
        # Configurar leyenda
        ax.legend(loc='lower right', fontsize=11)
        
        # Configurar grid
        ax.grid(True, axis='x', alpha=0.3)
        ax.set_axisbelow(True)
        
        # Ajustar márgenes
        plt.tight_layout()
        
        return fig

    def load_data(self):
        """Carga los datos desde Excel usando la lógica de registro retributivo"""
        try:
            # Obtener la ruta base del proyecto
            ruta_base = Path.cwd().parent if Path.cwd().name == '04_SCRIPTS' else Path.cwd()
            
            # Buscar el archivo más reciente en la carpeta de resultados
            carpeta_resultados = ruta_base / '02_RESULTADOS'
            print(f"Buscando archivos en: {carpeta_resultados}")
            
            # Buscar archivos que empiecen con REPORTE_
            archivos_reporte = list(carpeta_resultados.glob('REPORTE_*.xlsx'))
            if archivos_reporte:
                # Tomar el archivo más reciente
                archivo_mas_reciente = max(archivos_reporte, key=os.path.getctime)
                ruta_datos = archivo_mas_reciente
                print(f"Archivo más reciente encontrado: {ruta_datos.name}")
            else:
                # Fallback al archivo específico
                ruta_datos = carpeta_resultados / 'REPORTE_DATOS registro retributivo_20250902_235323.xlsx'
                print(f"Usando archivo específico: {ruta_datos}")

            print(f"Cargando datos desde: {ruta_datos}")

            # Leer todas las hojas del archivo procesado
            datos_procesados = {}
            archivo_excel = pd.ExcelFile(ruta_datos)
            for hoja in archivo_excel.sheet_names:
                datos_procesados[hoja] = pd.read_excel(archivo_excel, sheet_name=hoja)
                print(f"Cargada hoja '{hoja}': {len(datos_procesados[hoja])} registros")

            print("\nHojas disponibles en el archivo procesado:")
            for hoja in datos_procesados.keys():
                print(f"  - {hoja}")

            # Verificar si existe la hoja con datos procesados (actualizada)
            if 'DATOS_PROCESADOS' in datos_procesados:
                self.data = datos_procesados['DATOS_PROCESADOS']
                print(f"\nDataset principal cargado: {len(self.data)} registros")
                print(f"Columnas principales: {list(self.data.columns[:5])}...")
                
                # Verificar columnas clave
                columnas_clave = ['Sexo', 'Salario base anual efectivo', 'salario_base_equiparado']
                for col in columnas_clave:
                    if col in self.data.columns:
                        print(f"Columna '{col}' encontrada")
                    else:
                        print(f"Columna '{col}' NO encontrada")
                
                # Mostrar todas las columnas disponibles para verificar
                print(f"\nTodas las columnas disponibles ({len(self.data.columns)}):")
                for i, col in enumerate(self.data.columns, 1):
                    print(f"   {i:2d}. {col}")
                
                # Almacenar todas las hojas para uso posterior
                self.datos_procesados = datos_procesados
                
                # Configurar columna SEXO para compatibilidad
                if 'Sexo' in self.data.columns:
                    self.data['SEXO'] = self.data['Sexo'].map({'Hombres': 'H', 'Mujeres': 'M'})
                    print("✅ Columna SEXO configurada correctamente")
                    distribucion = self.data['SEXO'].value_counts()
                    print(f"   Distribución: {distribucion.to_dict()}")
                
                return True
            else:
                print("No se encontró la hoja 'DATOS_PROCESADOS'")
                if datos_procesados:
                    print("Hojas disponibles:")
                    for hoja in datos_procesados.keys():
                        print(f"  - {hoja}")
                return False
                
        except Exception as e:
            print(f"Error cargando datos: {e}")
            return False
    
    def prepare_chart_data(self, chart_config):
        """Prepara los datos para un gráfico específico"""
        columns = chart_config['data_columns']
        chart_type = chart_config['type']
        
        if len(columns) == 1:
            # Datos con una columna (para conteos o distribuciones)
            column = columns[0]
            if chart_type == 'pie':
                data = self.data[column].value_counts()
            else:
                data = self.data[column].dropna()
            
            if 'limit' in chart_config:
                data = data.head(chart_config['limit'])
            
            return data
            
        elif len(columns) == 2:
            # Datos con dos columnas (x, y)
            x_col, y_col = columns
            
            if chart_type == 'box':
                # Para boxplot, devolver los datos sin procesar
                return self.data[[x_col, y_col]].dropna()
            
            elif chart_type == 'scatter':
                # Para scatter, devolver los datos sin procesar
                return self.data[[x_col, y_col]].dropna()
            
            elif chart_type in ['bar', 'line']:
                # Para bar/line, agrupar y calcular promedio o suma según el contexto
                if 'brecha' in y_col.lower() or 'porcentual' in y_col.lower():
                    # Para brechas, calcular promedio
                    data = self.data.groupby(x_col)[y_col].mean()
                elif 'salario' in y_col.lower():
                    # Para salarios, calcular promedio
                    data = self.data.groupby(x_col)[y_col].mean()
                else:
                    # Por defecto, sumar
                    data = self.data.groupby(x_col)[y_col].sum()
                
                # Aplicar límite si existe
                if 'limit' in chart_config:
                    if chart_type == 'bar':
                        data = data.nlargest(chart_config['limit'])
                    else:
                        data = data.head(chart_config['limit'])
                
                return data
        
        # Fallback: devolver datos sin procesar
        return self.data[columns].dropna()
    
    def create_chart(self, chart_id, chart_config):
        """Crea un gráfico individual - especializado para gráficos de donut del registro retributivo"""
        chart_type = chart_config['type']
        title = chart_config['title']
        columns = chart_config['data_columns']
        
        if chart_type == 'donut':
            # Gráfico de donut específico del registro retributivo
            columna_datos = columns[0]  # La columna de salario a analizar
            metodo_calculo = chart_config.get('metodo', 'efectivos_sb')
            subtitulo = chart_config.get('subtitulo', '')
            
            # Seleccionar el método de cálculo apropiado
            if metodo_calculo == 'efectivos_sb':
                datos_genero = self.calcular_promedios_efectivos_sb(self.data, columna_datos)
            elif metodo_calculo == 'efectivos_sb_complementos':
                datos_genero = self.calcular_promedios_efectivos_sb_complementos(self.data, columna_datos)
            elif metodo_calculo == 'equiparados_sb':
                datos_genero = self.calcular_promedios_equiparados_sb(self.data, columna_datos)
            elif metodo_calculo == 'equiparados_sb_complementos':
                datos_genero = self.calcular_promedios_equiparados_sb_complementos(self.data, columna_datos)
            else:
                print(f"Método de cálculo no reconocido: {metodo_calculo}")
                return None
            
            # Crear el gráfico de donut
            fig = self.crear_grafico_donut(datos_genero, title, subtitulo)
            
            # Mostrar estadísticas
            diferencia = datos_genero['H'] - datos_genero['M']
            porcentaje_diferencia = (diferencia / datos_genero['M']) * 100 if datos_genero['M'] > 0 else 0
            
            print(f"📊 {title}")
            print(f"   Hombres: {self.formato_numero_es(datos_genero['H'], 2)}€")
            print(f"   Mujeres: {self.formato_numero_es(datos_genero['M'], 2)}€")
            print(f"   Brecha: {self.formato_numero_es(porcentaje_diferencia, 2)}%")
            print("-" * 50)
            
            # Guardar el gráfico
            chart_filename = f"temp_chart_{chart_id}.png"
            fig.savefig(chart_filename, dpi=300, bbox_inches='tight', 
                       facecolor='white', edgecolor='none')
            plt.close(fig)
            
            self.charts_created[chart_id] = {
                'filename': chart_filename,
                'marker': chart_config['marker'],
                'title': title
            }
            
            return chart_filename
        
        else:
            # Mantener funcionalidad para otros tipos de gráficos
            # Configurar tamaño de figura
            plt.figure(figsize=(12, 8))
            
            if chart_type == 'bar':
                data = self.prepare_chart_data(chart_config)
                bars = plt.bar(range(len(data)), data.values, color=sns.color_palette("husl", len(data)))
                plt.xticks(range(len(data)), data.index, rotation=45, ha='right')
                
                # Añadir valores encima de las barras
                for bar, value in zip(bars, data.values):
                    plt.text(bar.get_x() + bar.get_width()/2, bar.get_height() + max(data.values)*0.01,
                            f'{value:,.1f}%' if 'porcentual' in columns[1] else f'{value:,.0f}', 
                            ha='center', va='bottom', fontweight='bold')
            
            elif chart_type == 'pie':
                if len(columns) == 1:
                    # Para una sola columna, usar value_counts
                    data = self.data[columns[0]].value_counts()
                    colors = [self.colores_genero.get(idx, '#7C7C7C') for idx in data.index] if columns[0] == 'Sexo' else sns.color_palette("husl", len(data))
                else:
                    data = self.prepare_chart_data(chart_config)
                    colors = sns.color_palette("husl", len(data))
                
                wedges, texts, autotexts = plt.pie(data.values, labels=data.index, autopct='%1.1f%%', 
                                                  colors=colors, startangle=90)
                
                # Mejorar la apariencia del texto
                for autotext in autotexts:
                    autotext.set_color('white')
                    autotext.set_fontweight('bold')
            
            # Configurar el gráfico
            plt.title(title, fontsize=14, fontweight='bold', pad=20)
            plt.tight_layout()
            
            # Guardar el gráfico
            chart_filename = f"temp_chart_{chart_id}.png"
            plt.savefig(chart_filename, dpi=300, bbox_inches='tight', 
                       facecolor='white', edgecolor='none')
            plt.close()
            
            self.charts_created[chart_id] = {
                'filename': chart_filename,
                'marker': chart_config['marker'],
                'title': title
            }
            
            return chart_filename
    
    def create_all_charts(self):
        """Crea todos los gráficos definidos en la configuración"""
        if self.data is None:
            print("Error: No hay datos cargados")
            return False
        
        for chart_id, chart_config in self.config['charts'].items():
            try:
                print(f"Creando gráfico: {chart_id}")
                self.create_chart(chart_id, chart_config)
            except Exception as e:
                print(f"Error creando gráfico {chart_id}: {e}")
        
        return True
    
    def create_word_document(self):
        """Crea o actualiza el documento Word"""
        template_path = self.config.get('template_word')
        
        # Cargar plantilla o crear documento nuevo
        if template_path and os.path.exists(template_path):
            doc = Document(template_path)
            print(f"Usando plantilla: {template_path}")
        else:
            doc = Document()
            print("Creando documento nuevo")
            
            # Añadir contenido básico si no hay plantilla
            doc.add_heading('Informe de Análisis de Registro Retributivo', 0)
            doc.add_paragraph(f'Generado automáticamente el {datetime.now().strftime("%d/%m/%Y %H:%M")}')
            doc.add_page_break()
            
            # 1. PRIMERA SECCIÓN: Valores promedios de los salarios
            print("📊 Generando sección de valores promedios...")
            doc.add_heading('Valores promedios de los salarios', level=1)
            doc.add_paragraph('A continuación se presentan los resúmenes estadísticos de los diferentes conceptos retributivos analizados, comparando los valores efectivos (realmente percibidos) con los equiparados (normalizados a jornada completa y 12 meses).')
            doc.add_paragraph()
            
            # Generar resumen estadístico
            resumen = self.generar_resumen_estadisticas()
            
            # Crear las tres tablas de resumen
            self.crear_tabla_word(doc, '1. Análisis Salario Base (SB)', resumen, 'SB')
            self.crear_tabla_word(doc, '2. Análisis Salario Base + Complementos Salariales (SB+C)', resumen, 'SB+C')
            self.crear_tabla_word(doc, '3. Análisis Salario Base + Complementos + Extrasalariales (SB+C+ES)', resumen, 'SB+C+ES')
            
            doc.add_page_break()
            
            # 2. SEGUNDA SECCIÓN: Gráficos donut
            doc.add_heading('Análisis Gráfico Detallado', level=1)
            doc.add_paragraph('Los siguientes gráficos de donut permiten visualizar de forma clara las diferencias retributivas entre géneros y las brechas salariales correspondientes.')
            doc.add_paragraph()
            
            # Salario Base - Título seguido inmediatamente de sus gráficos
            doc.add_heading('1. Análisis Salario Base (SB)', level=2)
            doc.add_paragraph('Comparación entre valores efectivos (realmente percibidos) y equiparados (normalizados):')
            doc.add_paragraph('{grafico_sb_efectivo}')  # Se insertarán ambos gráficos aquí
            
            # Salario Base + Complementos - Título seguido inmediatamente de sus gráficos
            doc.add_heading('2. Análisis Salario Base + Complementos Salariales (SB+C)', level=2)
            doc.add_paragraph('Comparación entre valores efectivos y equiparados:')
            doc.add_paragraph('{grafico_sb_comp_efectivo}')  # Se insertarán ambos gráficos aquí
            
            # Salario Base + Complementos + Extrasalariales - Título seguido inmediatamente de sus gráficos
            doc.add_heading('3. Análisis Salario Base + Complementos + Extrasalariales (SB+C+ES)', level=2)
            doc.add_paragraph('Comparación entre valores efectivos y equiparados:')
            doc.add_paragraph('{grafico_sb_total_efectivo}')  # Se insertarán ambos gráficos aquí
            
            doc.add_page_break()
            
            # 3. TERCERA SECCIÓN: Tabla por Grupo Profesional
            print("📊 Generando tabla por Grupo Profesional...")
            doc.add_heading('Retribución Promedio por Grupo Profesional Efectivo', level=1)
            doc.add_paragraph('A continuación se presenta el análisis detallado de las retribuciones por grupo profesional, mostrando las diferencias entre valores efectivos (sin complementos y con complementos) incluyendo las respectivas brechas salariales.')
            doc.add_paragraph()
            
            # Generar y crear tabla por grupo profesional
            datos_grupo_profesional = self.generar_datos_grupo_profesional()
            self.crear_tabla_grupo_profesional(doc, datos_grupo_profesional)
            
            # 4. CUARTA SECCIÓN: Gráfico de barras por grupo profesional
            print("📊 Generando gráfico de barras por grupo profesional...")
            doc.add_heading('Análisis Visual por Grupo Profesional', level=2)
            doc.add_paragraph('El siguiente gráfico de barras muestra la comparación visual de los salarios CON complementos + extrasalariales por grupo profesional y género:')
            doc.add_paragraph()
            
            # Crear y guardar el gráfico de barras
            fig_barras = self.crear_grafico_barras_grupo_profesional(datos_grupo_profesional, "Efectivo")
            if fig_barras:
                barras_filename = "temp_chart_barras_grupo_profesional.png"
                fig_barras.savefig(barras_filename, dpi=300, bbox_inches='tight', 
                                 facecolor='white', edgecolor='none')
                plt.close(fig_barras)
                
                # Insertar el gráfico en el documento
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(barras_filename, width=Inches(6.5))
                
                # Limpiar archivo temporal
                if os.path.exists(barras_filename):
                    os.remove(barras_filename)
            
            # 5. QUINTA SECCIÓN: Tabla por Grupo Profesional EQUIPARADO
            print("📊 Generando tabla por Grupo Profesional EQUIPARADO...")
            doc.add_heading('Retribución Promedio por Grupo Profesional Equiparado', level=1)
            doc.add_paragraph('A continuación se presenta el análisis detallado de las retribuciones equiparadas por grupo profesional, normalizadas a jornada completa y 12 meses, mostrando las diferencias entre valores sin complementos y con complementos incluyendo las respectivas brechas salariales.')
            doc.add_paragraph()
            
            # Generar y crear tabla por grupo profesional equiparado
            datos_grupo_profesional_equiparado = self.generar_datos_grupo_profesional_equiparado()
            self.crear_tabla_grupo_profesional(doc, datos_grupo_profesional_equiparado)
            
            # 6. SEXTA SECCIÓN: Gráfico de barras por grupo profesional equiparado
            print("📊 Generando gráfico de barras por grupo profesional equiparado...")
            doc.add_heading('Análisis Visual por Grupo Profesional Equiparado', level=2)
            doc.add_paragraph('El siguiente gráfico de barras muestra la comparación visual de los salarios equiparados CON complementos + extrasalariales por grupo profesional y género:')
            doc.add_paragraph()
            
            # Crear y guardar el gráfico de barras equiparado
            fig_barras_equiparado = self.crear_grafico_barras_grupo_profesional(datos_grupo_profesional_equiparado, "Equiparado")
            if fig_barras_equiparado:
                barras_equiparado_filename = "temp_chart_barras_grupo_profesional_equiparado.png"
                fig_barras_equiparado.savefig(barras_equiparado_filename, dpi=300, bbox_inches='tight', 
                                 facecolor='white', edgecolor='none')
                plt.close(fig_barras_equiparado)
                
                # Insertar el gráfico en el documento
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(barras_equiparado_filename, width=Inches(6.5))
                
                # Limpiar archivo temporal
                if os.path.exists(barras_equiparado_filename):
                    os.remove(barras_equiparado_filename)
        
            # === ANÁLISIS POR SVPT + PUESTO DE TRABAJO EFECTIVO COMBINADO ===
            print("📊 Generando análisis por SVPT + Puesto de Trabajo Efectivo...")
            
            # Título principal para la sección combinada
            doc.add_heading('Retribución Promedio por SVPT y Puesto de Trabajo Efectivo', level=1)
            doc.add_paragraph('A continuación se presenta el análisis detallado de las retribuciones combinando el nivel SVPT (Sistema de Valoración de Puestos de Trabajo) con el puesto de trabajo específico, mostrando las diferencias salariales por género según esta categorización cruzada.')
            doc.add_paragraph()
            
            # Generar datos para la combinación SVPT + Puesto de Trabajo
            datos_combinacion = self.generar_datos_svpt_puesto_efectivo()
            
            if datos_combinacion:
                # Crear tabla combinada
                self.crear_tabla_svpt_puesto(doc, datos_combinacion)
                
                # Crear gráfico de barras horizontales para la combinación
                print("📊 Generando gráfico de barras por SVPT + Puesto de Trabajo...")
                doc.add_heading('Análisis Visual por SVPT y Puesto de Trabajo Efectivo', level=2)
                doc.add_paragraph('El siguiente gráfico de barras muestra la comparación visual de los salarios CON complementos + extrasalariales por combinación de SVPT y puesto de trabajo y género:')
                doc.add_paragraph()
                
                fig_barras_combinacion = self.crear_grafico_barras_svpt_puesto(datos_combinacion, "Efectivo")
                if fig_barras_combinacion:
                    # Guardar gráfico
                    barras_combinacion_filename = "temp_chart_barras_svpt_puesto.png"
                    fig_barras_combinacion.savefig(barras_combinacion_filename, dpi=300, bbox_inches='tight', 
                                                  facecolor='white', edgecolor='none')
                    plt.close(fig_barras_combinacion)
                    
                    # Insertar gráfico en el documento
                    paragraph = doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run()
                    run.add_picture(barras_combinacion_filename, width=Inches(6.5))
                    
                    # Limpiar archivo temporal
                    if os.path.exists(barras_combinacion_filename):
                        os.remove(barras_combinacion_filename)
                
                # Crear gráfico vertical adicional con puntos
                doc.add_heading('Análisis Detallado con Puntos por SVPT y Puesto de Trabajo Efectivo', level=3)
                doc.add_paragraph('El siguiente gráfico complementario muestra la relación entre salarios y puntos por combinación de SVPT y puesto de trabajo:')
                doc.add_paragraph()
                
                fig_vertical_combinacion = self.crear_grafico_barras_vertical_svpt_puesto(datos_combinacion, "Efectivo")
                if fig_vertical_combinacion:
                    # Guardar gráfico vertical
                    vertical_combinacion_filename = "temp_chart_vertical_svpt_puesto.png"
                    fig_vertical_combinacion.savefig(vertical_combinacion_filename, dpi=300, bbox_inches='tight', 
                                                    facecolor='white', edgecolor='none')
                    plt.close(fig_vertical_combinacion)
                    
                    # Insertar gráfico vertical en el documento
                    paragraph = doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run()
                    run.add_picture(vertical_combinacion_filename, width=Inches(self.config_graficos['ancho_documento']))
                    
                    # Limpiar archivo temporal
                    if os.path.exists(vertical_combinacion_filename):
                        os.remove(vertical_combinacion_filename)
            else:
                doc.add_paragraph("No hay datos suficientes para generar el análisis por SVPT y Puesto de Trabajo Efectivo.")
        
            # === ANÁLISIS POR SVPT + PUESTO DE TRABAJO EQUIPARADO COMBINADO ===
            print("📊 Generando análisis por SVPT + Puesto de Trabajo Equiparado...")
            
            # Título principal para la sección equiparada
            doc.add_heading('Retribución Promedio por SVPT y Puesto de Trabajo Equiparado', level=1)
            doc.add_paragraph('A continuación se presenta el análisis detallado de las retribuciones equiparadas combinando el nivel SVPT (Sistema de Valoración de Puestos de Trabajo) con el puesto de trabajo específico, normalizadas a jornada completa y 12 meses, mostrando las diferencias salariales por género según esta categorización cruzada.')
            doc.add_paragraph()
            
            # Generar datos equiparados para la combinación SVPT + Puesto de Trabajo
            datos_combinacion_equiparado = self.generar_datos_svpt_puesto_equiparado()
            
            if datos_combinacion_equiparado:
                # Crear tabla combinada equiparada
                self.crear_tabla_svpt_puesto(doc, datos_combinacion_equiparado)
                
                # Crear gráfico de barras horizontales para la combinación equiparada
                print("📊 Generando gráfico de barras por SVPT + Puesto de Trabajo Equiparado...")
                doc.add_heading('Análisis Visual por SVPT y Puesto de Trabajo Equiparado', level=2)
                doc.add_paragraph('El siguiente gráfico de barras muestra la comparación visual de los salarios equiparados CON complementos + extrasalariales por combinación de SVPT y puesto de trabajo y género:')
                doc.add_paragraph()
                
                fig_barras_combinacion_equiparado = self.crear_grafico_barras_svpt_puesto(datos_combinacion_equiparado, "Equiparado")
                if fig_barras_combinacion_equiparado:
                    # Guardar gráfico
                    barras_combinacion_equiparado_filename = "temp_chart_barras_svpt_puesto_equiparado.png"
                    fig_barras_combinacion_equiparado.savefig(barras_combinacion_equiparado_filename, dpi=300, bbox_inches='tight', 
                                                             facecolor='white', edgecolor='none')
                    plt.close(fig_barras_combinacion_equiparado)
                    
                    # Insertar gráfico en el documento
                    paragraph = doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run()
                    run.add_picture(barras_combinacion_equiparado_filename, width=Inches(6.5))
                    
                    # Limpiar archivo temporal
                    if os.path.exists(barras_combinacion_equiparado_filename):
                        os.remove(barras_combinacion_equiparado_filename)
                
                # Crear gráfico vertical adicional con puntos para equiparado
                doc.add_heading('Análisis Detallado con Puntos por SVPT y Puesto de Trabajo Equiparado', level=3)
                doc.add_paragraph('El siguiente gráfico complementario muestra la relación entre salarios equiparados y puntos por combinación de SVPT y puesto de trabajo:')
                doc.add_paragraph()
                
                fig_vertical_combinacion_equiparado = self.crear_grafico_barras_vertical_svpt_puesto(datos_combinacion_equiparado, "Equiparado")
                if fig_vertical_combinacion_equiparado:
                    # Guardar gráfico vertical equiparado
                    vertical_combinacion_equiparado_filename = "temp_chart_vertical_svpt_puesto_equiparado.png"
                    fig_vertical_combinacion_equiparado.savefig(vertical_combinacion_equiparado_filename, dpi=300, bbox_inches='tight', 
                                                               facecolor='white', edgecolor='none')
                    plt.close(fig_vertical_combinacion_equiparado)
                    
                    # Insertar gráfico vertical equiparado en el documento
                    paragraph = doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run()
                    run.add_picture(vertical_combinacion_equiparado_filename, width=Inches(self.config_graficos['ancho_documento']))
                    
                    # Limpiar archivo temporal
                    if os.path.exists(vertical_combinacion_equiparado_filename):
                        os.remove(vertical_combinacion_equiparado_filename)
            else:
                doc.add_paragraph("No hay datos suficientes para generar el análisis por SVPT y Puesto de Trabajo Equiparado.")

        # === ANÁLISIS DETALLADO POR ESCALA SVPT ===
        print("📊 Generando análisis por Escalas SVPT individuales...")
        
        # Obtener escalas disponibles
        escalas_disponibles = self.obtener_escalas_svpt()
        
        if escalas_disponibles:
            doc.add_heading('Análisis Detallado por Agrupación/Escala (Nivel SVPT)', level=1)
            doc.add_paragraph('A continuación se presenta el análisis detallado de las retribuciones segmentado por cada Agrupación/Escala específica del Sistema de Valoración de Puestos de Trabajo (SVPT), mostrando tanto los datos efectivos como los equiparados para cada escala por separado.')
            doc.add_paragraph()
            
            # Generar gráficos para cada escala
            for escala in escalas_disponibles:
                print(f"📈 Procesando escala {escala}...")
                
                # Análisis Efectivo por Escala
                doc.add_heading(f'Retribución Promedio por Agrupación/Escala {escala} (Nivel SVPT) y Puesto de trabajo', level=2)
                
                # Gráfico Efectivo
                doc.add_heading(f'Análisis Efectivo - Escala {escala}', level=3)
                doc.add_paragraph(f'Salarios medios CON complementos para la Agrupación/Escala {escala}, mostrando la relación entre puestos de trabajo y puntos SVPT:')
                doc.add_paragraph()
                
                datos_escala_efectivo = self.generar_datos_por_escala_svpt(escala, "Efectivo")
                if datos_escala_efectivo:
                    fig_escala_efectivo = self.crear_grafico_barras_por_escala(datos_escala_efectivo, escala, "Efectivo")
                    if fig_escala_efectivo:
                        # Guardar gráfico
                        escala_efectivo_filename = f"temp_chart_escala_{escala}_efectivo.png"
                        fig_escala_efectivo.savefig(escala_efectivo_filename, dpi=300, bbox_inches='tight', 
                                                   facecolor='white', edgecolor='none')
                        plt.close(fig_escala_efectivo)
                        
                        # Insertar gráfico en el documento
                        paragraph = doc.add_paragraph()
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = paragraph.add_run()
                        run.add_picture(escala_efectivo_filename, width=Inches(6.5))
                        
                        # Limpiar archivo temporal
                        if os.path.exists(escala_efectivo_filename):
                            os.remove(escala_efectivo_filename)
                else:
                    doc.add_paragraph(f"No hay datos suficientes para generar el gráfico efectivo de la escala {escala}.")
                
                # Gráfico Equiparado
                doc.add_heading(f'Análisis Equiparado - Escala {escala}', level=3)
                doc.add_paragraph(f'Salarios medios equiparados CON complementos para la Agrupación/Escala {escala}, normalizados a jornada completa:')
                doc.add_paragraph()
                
                datos_escala_equiparado = self.generar_datos_por_escala_svpt(escala, "Equiparado")
                if datos_escala_equiparado:
                    fig_escala_equiparado = self.crear_grafico_barras_por_escala(datos_escala_equiparado, escala, "Equiparado")
                    if fig_escala_equiparado:
                        # Guardar gráfico
                        escala_equiparado_filename = f"temp_chart_escala_{escala}_equiparado.png"
                        fig_escala_equiparado.savefig(escala_equiparado_filename, dpi=300, bbox_inches='tight', 
                                                     facecolor='white', edgecolor='none')
                        plt.close(fig_escala_equiparado)
                        
                        # Insertar gráfico en el documento
                        paragraph = doc.add_paragraph()
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = paragraph.add_run()
                        run.add_picture(escala_equiparado_filename, width=Inches(6.5))
                        
                        # Limpiar archivo temporal
                        if os.path.exists(escala_equiparado_filename):
                            os.remove(escala_equiparado_filename)
                else:
                    doc.add_paragraph(f"No hay datos suficientes para generar el gráfico equiparado de la escala {escala}.")
                
                # Separador entre escalas
                doc.add_paragraph()
        else:
            doc.add_paragraph("No se encontraron escalas SVPT válidas en los datos.")

        # === ANÁLISIS POR NIVEL ===
        print("📊 Generando análisis por Nivel...")
        
        # === ANÁLISIS POR NIVEL EFECTIVO ===
        doc.add_heading('Retribución Promedio por Nivel Efectivo', level=1)
        doc.add_paragraph('A continuación se presenta el análisis detallado de las retribuciones efectivas por nivel, mostrando las diferencias salariales por género según esta categorización.')
        doc.add_paragraph()
        
        # Generar datos efectivos por Nivel
        datos_nivel_efectivo = self.generar_datos_por_nivel("Efectivo")
        
        if datos_nivel_efectivo:
            # Crear tabla por nivel efectivo
            self.crear_tabla_por_nivel(doc, datos_nivel_efectivo)
            
            # Crear gráfico de barras por nivel efectivo
            print("📊 Generando gráfico de barras por Nivel Efectivo...")
            doc.add_heading('Análisis Visual por Nivel Efectivo', level=2)
            doc.add_paragraph('El siguiente gráfico de barras muestra la comparación visual de los salarios CON complementos + extrasalariales por nivel y género:')
            doc.add_paragraph()
            
            fig_nivel_efectivo = self.crear_grafico_barras_por_nivel(datos_nivel_efectivo, "Efectivo")
            if fig_nivel_efectivo:
                # Guardar gráfico
                nivel_efectivo_filename = "temp_chart_nivel_efectivo.png"
                fig_nivel_efectivo.savefig(nivel_efectivo_filename, dpi=self.config_graficos['dpi'], bbox_inches='tight', 
                                          facecolor='white', edgecolor='none')
                plt.close(fig_nivel_efectivo)
                
                # Insertar gráfico en el documento
                self.insertar_imagen_estandarizada(doc, nivel_efectivo_filename)
                
                # Limpiar archivo temporal
                if os.path.exists(nivel_efectivo_filename):
                    os.remove(nivel_efectivo_filename)
        else:
            doc.add_paragraph("No hay datos suficientes para generar el análisis por Nivel Efectivo.")

        # === ANÁLISIS POR NIVEL EQUIPARADO ===
        doc.add_heading('Retribución Promedio por Nivel Equiparado', level=1)
        doc.add_paragraph('A continuación se presenta el análisis detallado de las retribuciones equiparadas por nivel, normalizadas a jornada completa y 12 meses, mostrando las diferencias salariales por género según esta categorización.')
        doc.add_paragraph()
        
        # Generar datos equiparados por Nivel
        datos_nivel_equiparado = self.generar_datos_por_nivel("Equiparado")
        
        if datos_nivel_equiparado:
            # Crear tabla por nivel equiparado
            self.crear_tabla_por_nivel(doc, datos_nivel_equiparado)
            
            # Crear gráfico de barras por nivel equiparado
            print("📊 Generando gráfico de barras por Nivel Equiparado...")
            doc.add_heading('Análisis Visual por Nivel Equiparado', level=2)
            doc.add_paragraph('El siguiente gráfico de barras muestra la comparación visual de los salarios equiparados CON complementos + extrasalariales por nivel y género:')
            doc.add_paragraph()
            
            fig_nivel_equiparado = self.crear_grafico_barras_por_nivel(datos_nivel_equiparado, "Equiparado")
            if fig_nivel_equiparado:
                # Guardar gráfico
                nivel_equiparado_filename = "temp_chart_nivel_equiparado.png"
                fig_nivel_equiparado.savefig(nivel_equiparado_filename, dpi=self.config_graficos['dpi'], bbox_inches='tight', 
                                            facecolor='white', edgecolor='none')
                plt.close(fig_nivel_equiparado)
                
                # Insertar gráfico en el documento
                self.insertar_imagen_estandarizada(doc, nivel_equiparado_filename)
                
                # Limpiar archivo temporal
                if os.path.exists(nivel_equiparado_filename):
                    os.remove(nivel_equiparado_filename)
        else:
            doc.add_paragraph("No hay datos suficientes para generar el análisis por Nivel Equiparado.")
        
        # ============================================
        # === SECCIÓN COMPLETA CON MEDIANA ===
        # ============================================
        
        doc.add_page_break()
        doc.add_heading('ANÁLISIS COMPLETO CON MEDIANA', level=1)
        doc.add_paragraph('A continuación se presenta la totalidad del análisis anterior pero utilizando la MEDIANA en lugar de la MEDIA como estadístico central, manteniendo exactamente el mismo orden y estructura.')
        doc.add_paragraph()
        
        # === 1. VALORES MEDIANOS DE LOS SALARIOS ===
        print("📊 Generando sección de valores promedios CON MEDIANA...")
        doc.add_heading('Valores MEDIANOS de los salarios', level=1)
        doc.add_paragraph('A continuación se presentan los resúmenes estadísticos utilizando la MEDIANA de los diferentes conceptos retributivos analizados, comparando los valores efectivos (realmente percibidos) con los equiparados (normalizados a jornada completa y 12 meses).')
        doc.add_paragraph()
        
        # Generar resumen estadístico con MEDIANA
        resumen_mediana = self.generar_resumen_estadisticas_mediana()
        
        # Crear las tres tablas de resumen con MEDIANA
        self.crear_tabla_word(doc, '1. Análisis Salario Base (SB) - MEDIANA', resumen_mediana, 'SB')
        self.crear_tabla_word(doc, '2. Análisis Salario Base + Complementos Salariales (SB+C) - MEDIANA', resumen_mediana, 'SB+C')
        self.crear_tabla_word(doc, '3. Análisis Salario Base + Complementos + Extrasalariales (SB+C+ES) - MEDIANA', resumen_mediana, 'SB+C+ES')
        
        # === 2. ANÁLISIS POR GRUPO PROFESIONAL CON MEDIANA ===
        
        # Tabla por Grupo Profesional EFECTIVO con MEDIANA
        print("📊 Generando tabla por Grupo Profesional EFECTIVO con MEDIANA...")
        doc.add_heading('Retribución MEDIANA por Grupo Profesional Efectivo', level=1)
        doc.add_paragraph('A continuación se presenta el análisis detallado de las retribuciones utilizando la MEDIANA por grupo profesional, mostrando las diferencias entre valores efectivos (sin complementos y con complementos) incluyendo las respectivas brechas salariales.')
        doc.add_paragraph()
        
        # Generar y crear tabla por grupo profesional con mediana
        datos_grupo_profesional_mediana = self.generar_datos_grupo_profesional_mediana()
        self.crear_tabla_grupo_profesional(doc, datos_grupo_profesional_mediana)
        
        # Gráfico de barras por grupo profesional con mediana
        print("📊 Generando gráfico de barras por grupo profesional con MEDIANA...")
        doc.add_heading('Análisis Visual por Grupo Profesional Efectivo - MEDIANA', level=2)
        doc.add_paragraph('El siguiente gráfico de barras muestra la comparación visual de los salarios (MEDIANA) CON complementos + extrasalariales por grupo profesional y género:')
        doc.add_paragraph()
        
        # Crear y guardar el gráfico de barras con mediana
        fig_barras_mediana = self.crear_grafico_barras_grupo_profesional(datos_grupo_profesional_mediana, "Efectivo - MEDIANA")
        if fig_barras_mediana:
            barras_mediana_filename = "temp_chart_barras_grupo_profesional_mediana.png"
            fig_barras_mediana.savefig(barras_mediana_filename, dpi=300, bbox_inches='tight', 
                                     facecolor='white', edgecolor='none')
            plt.close(fig_barras_mediana)
            
            # Insertar el gráfico en el documento
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(barras_mediana_filename, width=Inches(6.5))
            
            # Limpiar archivo temporal
            if os.path.exists(barras_mediana_filename):
                os.remove(barras_mediana_filename)
        
        # Tabla por Grupo Profesional EQUIPARADO con MEDIANA
        print("📊 Generando tabla por Grupo Profesional EQUIPARADO con MEDIANA...")
        doc.add_heading('Retribución MEDIANA por Grupo Profesional Equiparado', level=1)
        doc.add_paragraph('A continuación se presenta el análisis detallado de las retribuciones equiparadas utilizando la MEDIANA por grupo profesional, normalizadas a jornada completa y 12 meses, mostrando las diferencias entre valores sin complementos y con complementos incluyendo las respectivas brechas salariales.')
        doc.add_paragraph()
        
        # Generar y crear tabla por grupo profesional equiparado con mediana
        datos_grupo_profesional_equiparado_mediana = self.generar_datos_grupo_profesional_equiparado_mediana()
        self.crear_tabla_grupo_profesional(doc, datos_grupo_profesional_equiparado_mediana)
        
        # Gráfico de barras por grupo profesional equiparado con mediana
        print("📊 Generando gráfico de barras por grupo profesional equiparado con MEDIANA...")
        doc.add_heading('Análisis Visual por Grupo Profesional Equiparado - MEDIANA', level=2)
        doc.add_paragraph('El siguiente gráfico de barras muestra la comparación visual de los salarios equiparados (MEDIANA) CON complementos + extrasalariales por grupo profesional y género:')
        doc.add_paragraph()
        
        # Crear y guardar el gráfico de barras equiparado con mediana
        fig_barras_equiparado_mediana = self.crear_grafico_barras_grupo_profesional(datos_grupo_profesional_equiparado_mediana, "Equiparado - MEDIANA")
        if fig_barras_equiparado_mediana:
            barras_equiparado_mediana_filename = "temp_chart_barras_grupo_profesional_equiparado_mediana.png"
            fig_barras_equiparado_mediana.savefig(barras_equiparado_mediana_filename, dpi=300, bbox_inches='tight', 
                                                 facecolor='white', edgecolor='none')
            plt.close(fig_barras_equiparado_mediana)
            
            # Insertar el gráfico en el documento
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(barras_equiparado_mediana_filename, width=Inches(6.5))
            
            # Limpiar archivo temporal
            if os.path.exists(barras_equiparado_mediana_filename):
                os.remove(barras_equiparado_mediana_filename)

        # === 3. ANÁLISIS POR SVPT + PUESTO DE TRABAJO CON MEDIANA ===
            
        # SVPT + Puesto de Trabajo EFECTIVO con MEDIANA
        print("📊 Generando análisis por SVPT + Puesto de Trabajo Efectivo con MEDIANA...")
        
        doc.add_heading('Retribución MEDIANA por SVPT y Puesto de Trabajo Efectivo', level=1)
        doc.add_paragraph('A continuación se presenta el análisis detallado de las retribuciones utilizando la MEDIANA combinando el nivel SVPT (Sistema de Valoración de Puestos de Trabajo) con el puesto de trabajo específico, mostrando las diferencias salariales por género según esta categorización cruzada.')
        doc.add_paragraph()
        
        # Generar datos para la combinación SVPT + Puesto de Trabajo con MEDIANA
        datos_combinacion_mediana = self.generar_datos_svpt_puesto_efectivo_mediana()
        
        if datos_combinacion_mediana:
            # Crear tabla combinada con mediana
            self.crear_tabla_svpt_puesto(doc, datos_combinacion_mediana)
            
            # Crear gráfico de barras horizontales para la combinación con mediana
            print("📊 Generando gráfico de barras por SVPT + Puesto de Trabajo con MEDIANA...")
            doc.add_heading('Análisis Visual por SVPT y Puesto de Trabajo Efectivo - MEDIANA', level=2)
            doc.add_paragraph('El siguiente gráfico de barras muestra la comparación visual de los salarios (MEDIANA) CON complementos + extrasalariales por combinación de SVPT y puesto de trabajo y género:')
            doc.add_paragraph()
            
            fig_barras_combinacion_mediana = self.crear_grafico_barras_svpt_puesto(datos_combinacion_mediana, "Efectivo - MEDIANA")
            if fig_barras_combinacion_mediana:
                # Guardar gráfico
                barras_combinacion_mediana_filename = "temp_chart_barras_svpt_puesto_mediana.png"
                fig_barras_combinacion_mediana.savefig(barras_combinacion_mediana_filename, dpi=300, bbox_inches='tight', 
                                                      facecolor='white', edgecolor='none')
                plt.close(fig_barras_combinacion_mediana)
                
                # Insertar gráfico en el documento
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(barras_combinacion_mediana_filename, width=Inches(6.5))
                
                # Limpiar archivo temporal
                if os.path.exists(barras_combinacion_mediana_filename):
                    os.remove(barras_combinacion_mediana_filename)
            
            # Crear gráfico vertical adicional con puntos con mediana
            doc.add_heading('Análisis Detallado con Puntos por SVPT y Puesto de Trabajo Efectivo - MEDIANA', level=3)
            doc.add_paragraph('El siguiente gráfico complementario muestra la relación entre salarios (MEDIANA) y puntos por combinación de SVPT y puesto de trabajo:')
            doc.add_paragraph()
            
            fig_vertical_combinacion_mediana = self.crear_grafico_barras_vertical_svpt_puesto(datos_combinacion_mediana, "Efectivo - MEDIANA")
            if fig_vertical_combinacion_mediana:
                # Guardar gráfico vertical
                vertical_combinacion_mediana_filename = "temp_chart_vertical_svpt_puesto_mediana.png"
                fig_vertical_combinacion_mediana.savefig(vertical_combinacion_mediana_filename, dpi=300, bbox_inches='tight', 
                                                        facecolor='white', edgecolor='none')
                plt.close(fig_vertical_combinacion_mediana)
                
                # Insertar gráfico vertical en el documento
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(vertical_combinacion_mediana_filename, width=Inches(self.config_graficos['ancho_documento']))
                
                # Limpiar archivo temporal
                if os.path.exists(vertical_combinacion_mediana_filename):
                    os.remove(vertical_combinacion_mediana_filename)
        else:
            doc.add_paragraph("No hay datos suficientes para generar el análisis por SVPT y Puesto de Trabajo Efectivo con MEDIANA.")
    
        # SVPT + Puesto de Trabajo EQUIPARADO con MEDIANA
        print("📊 Generando análisis por SVPT + Puesto de Trabajo Equiparado con MEDIANA...")
        
        doc.add_heading('Retribución MEDIANA por SVPT y Puesto de Trabajo Equiparado', level=1)
        doc.add_paragraph('A continuación se presenta el análisis detallado de las retribuciones equiparadas utilizando la MEDIANA combinando el nivel SVPT (Sistema de Valoración de Puestos de Trabajo) con el puesto de trabajo específico, normalizadas a jornada completa y 12 meses, mostrando las diferencias salariales por género según esta categorización cruzada.')
        doc.add_paragraph()
        
        # Generar datos equiparados para la combinación SVPT + Puesto de Trabajo con MEDIANA
        datos_combinacion_equiparado_mediana = self.generar_datos_svpt_puesto_equiparado_mediana()
        
        if datos_combinacion_equiparado_mediana:
            # Crear tabla combinada equiparada con mediana
            self.crear_tabla_svpt_puesto(doc, datos_combinacion_equiparado_mediana)
            
            # Crear gráfico de barras horizontales para la combinación equiparada con mediana
            print("📊 Generando gráfico de barras por SVPT + Puesto de Trabajo Equiparado con MEDIANA...")
            doc.add_heading('Análisis Visual por SVPT y Puesto de Trabajo Equiparado - MEDIANA', level=2)
            doc.add_paragraph('El siguiente gráfico de barras muestra la comparación visual de los salarios equiparados (MEDIANA) CON complementos + extrasalariales por combinación de SVPT y puesto de trabajo y género:')
            doc.add_paragraph()
            
            fig_barras_combinacion_equiparado_mediana = self.crear_grafico_barras_svpt_puesto(datos_combinacion_equiparado_mediana, "Equiparado - MEDIANA")
            if fig_barras_combinacion_equiparado_mediana:
                # Guardar gráfico
                barras_combinacion_equiparado_mediana_filename = "temp_chart_barras_svpt_puesto_equiparado_mediana.png"
                fig_barras_combinacion_equiparado_mediana.savefig(barras_combinacion_equiparado_mediana_filename, dpi=300, bbox_inches='tight', 
                                                                 facecolor='white', edgecolor='none')
                plt.close(fig_barras_combinacion_equiparado_mediana)
                
                # Insertar gráfico en el documento
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(barras_combinacion_equiparado_mediana_filename, width=Inches(6.5))
                
                # Limpiar archivo temporal
                if os.path.exists(barras_combinacion_equiparado_mediana_filename):
                    os.remove(barras_combinacion_equiparado_mediana_filename)
            
            # Crear gráfico vertical adicional con puntos para equiparado con mediana
            doc.add_heading('Análisis Detallado con Puntos por SVPT y Puesto de Trabajo Equiparado - MEDIANA', level=3)
            doc.add_paragraph('El siguiente gráfico complementario muestra la relación entre salarios equiparados (MEDIANA) y puntos por combinación de SVPT y puesto de trabajo:')
            doc.add_paragraph()
            
            fig_vertical_combinacion_equiparado_mediana = self.crear_grafico_barras_vertical_svpt_puesto(datos_combinacion_equiparado_mediana, "Equiparado - MEDIANA")
            if fig_vertical_combinacion_equiparado_mediana:
                # Guardar gráfico vertical equiparado con mediana
                vertical_combinacion_equiparado_mediana_filename = "temp_chart_vertical_svpt_puesto_equiparado_mediana.png"
                fig_vertical_combinacion_equiparado_mediana.savefig(vertical_combinacion_equiparado_mediana_filename, dpi=300, bbox_inches='tight', 
                                                                   facecolor='white', edgecolor='none')
                plt.close(fig_vertical_combinacion_equiparado_mediana)
                
                # Insertar gráfico vertical equiparado en el documento
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(vertical_combinacion_equiparado_mediana_filename, width=Inches(self.config_graficos['ancho_documento']))
                
                # Limpiar archivo temporal
                if os.path.exists(vertical_combinacion_equiparado_mediana_filename):
                    os.remove(vertical_combinacion_equiparado_mediana_filename)
        else:
            doc.add_paragraph("No hay datos suficientes para generar el análisis por SVPT y Puesto de Trabajo Equiparado con MEDIANA.")

        # === 4. ANÁLISIS DETALLADO POR ESCALA SVPT CON MEDIANA ===
        print("📊 Generando análisis por Escalas SVPT individuales con MEDIANA...")
        
        # Obtener escalas disponibles
        escalas_disponibles_mediana = self.obtener_escalas_svpt()
        
        if escalas_disponibles_mediana:
            print(f"📋 Escalas SVPT encontradas para MEDIANA: {escalas_disponibles_mediana}")
            
            doc.add_heading('Análisis Detallado por Agrupación/Escala (Nivel SVPT) - MEDIANA', level=1)
            doc.add_paragraph('A continuación se presenta el análisis detallado de las retribuciones utilizando la MEDIANA segmentado por cada Agrupación/Escala específica del Sistema de Valoración de Puestos de Trabajo (SVPT), mostrando tanto los datos efectivos como los equiparados para cada escala por separado.')
            doc.add_paragraph()
            
            # Generar gráficos para cada escala con MEDIANA
            for escala in escalas_disponibles_mediana:
                print(f"📈 Procesando escala {escala} con MEDIANA...")
                
                # Análisis Efectivo por Escala con MEDIANA
                doc.add_heading(f'Retribución MEDIANA por Agrupación/Escala {escala} (Nivel SVPT) y Puesto de trabajo', level=2)
                
                # Gráfico Efectivo con MEDIANA
                doc.add_heading(f'Análisis Efectivo - Escala {escala} (MEDIANA)', level=3)
                doc.add_paragraph(f'Salarios MEDIANOS CON complementos para la Agrupación/Escala {escala}, mostrando la relación entre puestos de trabajo y puntos SVPT:')
                doc.add_paragraph()
                
                datos_escala_efectivo_mediana = self.generar_datos_por_escala_svpt_mediana(escala, "Efectivo")
                if datos_escala_efectivo_mediana:
                    fig_escala_efectivo_mediana = self.crear_grafico_barras_por_escala(datos_escala_efectivo_mediana, escala, "Efectivo - MEDIANA")
                    if fig_escala_efectivo_mediana:
                        # Guardar gráfico
                        escala_efectivo_mediana_filename = f"temp_chart_escala_{escala}_efectivo_mediana.png"
                        fig_escala_efectivo_mediana.savefig(escala_efectivo_mediana_filename, dpi=300, bbox_inches='tight', 
                                                           facecolor='white', edgecolor='none')
                        plt.close(fig_escala_efectivo_mediana)
                        
                        # Insertar gráfico en el documento
                        paragraph = doc.add_paragraph()
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = paragraph.add_run()
                        run.add_picture(escala_efectivo_mediana_filename, width=Inches(6.5))
                        
                        # Limpiar archivo temporal
                        if os.path.exists(escala_efectivo_mediana_filename):
                            os.remove(escala_efectivo_mediana_filename)
                else:
                    doc.add_paragraph(f"No hay datos suficientes para generar el gráfico efectivo con MEDIANA de la escala {escala}.")
                
                # Gráfico Equiparado con MEDIANA
                doc.add_heading(f'Análisis Equiparado - Escala {escala} (MEDIANA)', level=3)
                doc.add_paragraph(f'Salarios MEDIANOS equiparados CON complementos para la Agrupación/Escala {escala}, normalizados a jornada completa:')
                doc.add_paragraph()
                
                datos_escala_equiparado_mediana = self.generar_datos_por_escala_svpt_mediana(escala, "Equiparado")
                if datos_escala_equiparado_mediana:
                    fig_escala_equiparado_mediana = self.crear_grafico_barras_por_escala(datos_escala_equiparado_mediana, escala, "Equiparado - MEDIANA")
                    if fig_escala_equiparado_mediana:
                        # Guardar gráfico
                        escala_equiparado_mediana_filename = f"temp_chart_escala_{escala}_equiparado_mediana.png"
                        fig_escala_equiparado_mediana.savefig(escala_equiparado_mediana_filename, dpi=300, bbox_inches='tight', 
                                                             facecolor='white', edgecolor='none')
                        plt.close(fig_escala_equiparado_mediana)
                        
                        # Insertar gráfico en el documento
                        paragraph = doc.add_paragraph()
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = paragraph.add_run()
                        run.add_picture(escala_equiparado_mediana_filename, width=Inches(6.5))
                        
                        # Limpiar archivo temporal
                        if os.path.exists(escala_equiparado_mediana_filename):
                            os.remove(escala_equiparado_mediana_filename)
                else:
                    doc.add_paragraph(f"No hay datos suficientes para generar el gráfico equiparado con MEDIANA de la escala {escala}.")
                
                # Separador entre escalas
                doc.add_paragraph()
        else:
            doc.add_paragraph("No se encontraron escalas SVPT válidas en los datos para el análisis con MEDIANA.")

        # === 5. ANÁLISIS POR NIVEL CON MEDIANA ===
        print("📊 Generando análisis por Nivel con MEDIANA...")
        
        # === ANÁLISIS POR NIVEL EFECTIVO CON MEDIANA ===
        doc.add_heading('Retribución MEDIANA por Nivel Efectivo', level=1)
        doc.add_paragraph('A continuación se presenta el análisis detallado de las retribuciones efectivas utilizando la MEDIANA por nivel, mostrando las diferencias salariales por género según esta categorización.')
        doc.add_paragraph()
        
        # Generar datos efectivos por Nivel con MEDIANA
        datos_nivel_efectivo_mediana = self.generar_datos_por_nivel_mediana("Efectivo")
        
        if datos_nivel_efectivo_mediana:
            # Crear tabla por nivel efectivo con mediana
            self.crear_tabla_por_nivel(doc, datos_nivel_efectivo_mediana)
            
            # Crear gráfico de barras por nivel efectivo con mediana
            print("📊 Generando gráfico de barras por Nivel Efectivo con MEDIANA...")
            doc.add_heading('Análisis Visual por Nivel Efectivo - MEDIANA', level=2)
            doc.add_paragraph('El siguiente gráfico de barras muestra la comparación visual de los salarios (MEDIANA) CON complementos + extrasalariales por nivel y género:')
            doc.add_paragraph()
            
            fig_nivel_efectivo_mediana = self.crear_grafico_barras_por_nivel(datos_nivel_efectivo_mediana, "Efectivo - MEDIANA")
            if fig_nivel_efectivo_mediana:
                # Guardar gráfico
                nivel_efectivo_mediana_filename = "temp_chart_nivel_efectivo_mediana.png"
                fig_nivel_efectivo_mediana.savefig(nivel_efectivo_mediana_filename, dpi=self.config_graficos['dpi'], bbox_inches='tight', 
                                                  facecolor='white', edgecolor='none')
                plt.close(fig_nivel_efectivo_mediana)
                
                # Insertar gráfico en el documento
                self.insertar_imagen_estandarizada(doc, nivel_efectivo_mediana_filename)
                
                # Limpiar archivo temporal
                if os.path.exists(nivel_efectivo_mediana_filename):
                    os.remove(nivel_efectivo_mediana_filename)
        else:
            doc.add_paragraph("No hay datos suficientes para generar el análisis por Nivel Efectivo con MEDIANA.")

        # === ANÁLISIS POR NIVEL EQUIPARADO CON MEDIANA ===
        doc.add_heading('Retribución MEDIANA por Nivel Equiparado', level=1)
        doc.add_paragraph('A continuación se presenta el análisis detallado de las retribuciones equiparadas utilizando la MEDIANA por nivel, normalizadas a jornada completa y 12 meses, mostrando las diferencias salariales por género según esta categorización.')
        doc.add_paragraph()
        
        # Generar datos equiparados por Nivel con MEDIANA
        datos_nivel_equiparado_mediana = self.generar_datos_por_nivel_mediana("Equiparado")
        
        if datos_nivel_equiparado_mediana:
            # Crear tabla por nivel equiparado con mediana
            self.crear_tabla_por_nivel(doc, datos_nivel_equiparado_mediana)
            
            # Crear gráfico de barras por nivel equiparado con mediana
            print("📊 Generando gráfico de barras por Nivel Equiparado con MEDIANA...")
            doc.add_heading('Análisis Visual por Nivel Equiparado - MEDIANA', level=2)
            doc.add_paragraph('El siguiente gráfico de barras muestra la comparación visual de los salarios equiparados (MEDIANA) CON complementos + extrasalariales por nivel y género:')
            doc.add_paragraph()
            
            fig_nivel_equiparado_mediana = self.crear_grafico_barras_por_nivel(datos_nivel_equiparado_mediana, "Equiparado - MEDIANA")
            if fig_nivel_equiparado_mediana:
                # Guardar gráfico
                nivel_equiparado_mediana_filename = "temp_chart_nivel_equiparado_mediana.png"
                fig_nivel_equiparado_mediana.savefig(nivel_equiparado_mediana_filename, dpi=self.config_graficos['dpi'], bbox_inches='tight', 
                                                    facecolor='white', edgecolor='none')
                plt.close(fig_nivel_equiparado_mediana)
                
                # Insertar gráfico en el documento
                self.insertar_imagen_estandarizada(doc, nivel_equiparado_mediana_filename)
                
                # Limpiar archivo temporal
                if os.path.exists(nivel_equiparado_mediana_filename):
                    os.remove(nivel_equiparado_mediana_filename)
        else:
            doc.add_paragraph("No hay datos suficientes para generar el análisis por Nivel Equiparado con MEDIANA.")
        
        # Reemplazar marcadores con gráficos
        self.replace_markers_with_charts(doc)
        
        # Guardar documento con timestamp único
        base_output = self.config['output_file']
        
        # Generar nombre con timestamp actual
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Limpiar nombre base y agregar timestamp
        if '05_INFORMES/' in base_output:
            # Extraer solo la parte del directorio y nombre base
            parts = base_output.split('05_INFORMES/')
            if len(parts) > 1:
                base_name = parts[1].replace('.docx', '').split('_')[0:2]  # Tomar solo "registro_retributivo"
                output_path = f"05_INFORMES/{'_'.join(base_name)}_{timestamp}.docx"
            else:
                output_path = f"05_INFORMES/registro_retributivo_{timestamp}.docx"
        else:
            if base_output.endswith('.docx'):
                output_path = base_output.replace('.docx', f'_{timestamp}.docx')
            else:
                output_path = f"{base_output}_{timestamp}.docx"
        
        # Crear directorio si no existe
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
            print(f"Directorio creado: {output_dir}")
        
        doc.save(output_path)
        print(f"Documento guardado en: {output_path}")
        
        return output_path
    
    def replace_markers_with_charts(self, doc):
        """Reemplaza los marcadores en el documento con los gráficos organizados en pares"""
        from docx.shared import Inches
        from docx.enum.table import WD_TABLE_ALIGNMENT
        
        # Definir pares de gráficos (efectivo + equiparado)
        pares_graficos = [
            ('{grafico_sb_efectivo}', '{grafico_sb_equiparado}'),
            ('{grafico_sb_comp_efectivo}', '{grafico_sb_comp_equiparado}'),
            ('{grafico_sb_total_efectivo}', '{grafico_sb_total_equiparado}')
        ]
        
        # Necesitamos procesar desde el final hacia el principio para no afectar los índices
        for efectivo_marker, equiparado_marker in reversed(pares_graficos):
            # Buscar el párrafo que contiene el marcador
            for i, paragraph in enumerate(doc.paragraphs):
                if efectivo_marker in paragraph.text:
                    print(f"Encontrado marcador: {efectivo_marker} en párrafo {i}")
                    
                    # Buscar las imágenes correspondientes
                    efectivo_chart = None
                    equiparado_chart = None
                    
                    for chart_id, chart_info in self.charts_created.items():
                        if chart_info['marker'] == efectivo_marker:
                            efectivo_chart = chart_info
                        elif chart_info['marker'] == equiparado_marker:
                            equiparado_chart = chart_info
                    
                    if efectivo_chart and equiparado_chart:
                        print(f"Insertando gráficos: {efectivo_chart['filename']} + {equiparado_chart['filename']}")
                        
                        # Limpiar el párrafo del marcador
                        paragraph.clear()
                        
                        # Insertar directamente en este párrafo como una tabla usando runs
                        # Método más directo: usar el párrafo como contenedor
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Insertar la primera imagen
                        run1 = paragraph.add_run()
                        run1.add_picture(efectivo_chart['filename'], width=Inches(2.8))
                        
                        # Añadir espacio entre imágenes
                        paragraph.add_run("    ")  # Espacios en blanco
                        
                        # Insertar la segunda imagen
                        run2 = paragraph.add_run()
                        run2.add_picture(equiparado_chart['filename'], width=Inches(2.8))
                        
                        # Añadir espacios después
                        # Obtener el elemento padre para insertar párrafos nuevos
                        parent = paragraph._element.getparent()
                        paragraph_index = list(parent).index(paragraph._element)
                        
                        # Crear párrafos vacíos después
                        from docx.oxml import OxmlElement
                        for _ in range(2):
                            new_para = OxmlElement('w:p')
                            parent.insert(paragraph_index + 1, new_para)
                    
                    break  # Salir del bucle de párrafos para este marcador
    
    def cleanup_temp_files(self):
        """Limpia los archivos temporales"""
        for chart_info in self.charts_created.values():
            filename = chart_info['filename']
            if os.path.exists(filename):
                os.remove(filename)
                print(f"Archivo temporal eliminado: {filename}")
    
    def generate_report(self):
        """Ejecuta todo el flujo de generación del reporte"""
        print("*** Iniciando generación automatizada del reporte de registro retributivo...")
        
        # 1. Cargar datos
        if not self.load_data():
            return False
        
        # 2. Calcular brechas salariales
        print("*** Calculando brechas salariales...")
        self.calcular_brecha_salarial()
        
        # 3. Crear todos los gráficos
        if not self.create_all_charts():
            return False
        
        # 4. Crear documento Word
        output_path = self.create_word_document()
        
        # 5. Limpiar archivos temporales
        self.cleanup_temp_files()
        
        print(f"*** Reporte de registro retributivo generado exitosamente: {output_path}")
        return True
    
    def generar_datos_svpt_puesto_efectivo(self):
        """Genera los datos para la tabla de retribución por SVPT + Puesto de trabajo EFECTIVO combinados"""
        print("📊 Calculando datos efectivos por SVPT + Puesto de trabajo...")
        
        # Filtrar datos actuales (sin "Ex" en primera columna)
        df_actual = self.data[self.data.iloc[:, 0] != 'Ex'].copy()
        
        # Verificar que ambas columnas existen
        if 'Nivel SVPT' not in df_actual.columns:
            print("❌ Error: No se encontró la columna 'Nivel SVPT'")
            return []
        
        if 'Puesto de trabajo' not in df_actual.columns:
            print("❌ Error: No se encontró la columna 'Puesto de trabajo'")
            return []
        
        # Crear una columna combinada de SVPT + Puesto de trabajo
        df_actual['SVPT_Puesto'] = df_actual['Nivel SVPT'].astype(str) + ' - ' + df_actual['Puesto de trabajo'].astype(str)
        
        # Obtener combinaciones únicas
        combinaciones = df_actual['SVPT_Puesto'].unique()
        combinaciones = [str(c) for c in combinaciones if pd.notna(c) and str(c) != 'nan - nan']
        datos_combinaciones = []
        
        # Columnas para valores efectivos
        col_sin_comp = 'Salario base efectivo Total'  # SB efectivo
        col_con_comp = 'Salario base anual + complementos Total'  # SB + Complementos efectivo
        col_con_extra = 'Salario base anual + complementos + Extrasalariales Total'  # SB + Complementos + Extrasalariales efectivo
        
        for combinacion in sorted(combinaciones):
            # Filtrar por la combinación específica
            df_comb = df_actual[df_actual['SVPT_Puesto'].astype(str) == str(combinacion)]
            
            # Conteos por género
            conteos = df_comb['SEXO'].value_counts()
            n_mujeres = conteos.get('M', 0)
            n_hombres = conteos.get('H', 0)
            
            # SIN complementos (solo SB efectivo) - solo personas con SB > 0
            datos_sin_comp = self.calcular_promedios_efectivos_sb(df_comb, col_sin_comp)
            sin_comp_m = datos_sin_comp['M']
            sin_comp_h = datos_sin_comp['H']
            brecha_sin_comp = ((sin_comp_h - sin_comp_m) / sin_comp_m * 100) if sin_comp_m > 0 else 0
            
            # CON complementos efectivos - todas las personas
            datos_con_comp = self.calcular_promedios_efectivos_sb_complementos(df_comb, col_con_comp)
            con_comp_m = datos_con_comp['M']
            con_comp_h = datos_con_comp['H']
            brecha_con_comp = ((con_comp_h - con_comp_m) / con_comp_m * 100) if con_comp_m > 0 else 0
            
            # CON extrasalariales efectivos - todas las personas
            datos_con_extra = self.calcular_promedios_efectivos_sb_complementos(df_comb, col_con_extra)
            con_extra_m = datos_con_extra['M']
            con_extra_h = datos_con_extra['H']
            brecha_con_extra = ((con_extra_h - con_extra_m) / con_extra_m * 100) if con_extra_m > 0 else 0
            
            datos_combinaciones.append({
                'svpt_puesto': combinacion,
                'n_mujeres': n_mujeres,
                'n_hombres': n_hombres,
                'sin_comp_m': sin_comp_m,
                'sin_comp_h': sin_comp_h,
                'brecha_sin_comp': brecha_sin_comp,
                'con_comp_m': con_comp_m,
                'con_comp_h': con_comp_h,
                'brecha_con_comp': brecha_con_comp,
                'con_extra_m': con_extra_m,
                'con_extra_h': con_extra_h,
                'brecha_con_extra': brecha_con_extra
            })
        
        # Calcular totales efectivos
        if datos_combinaciones:
            total_m = sum([d['n_mujeres'] for d in datos_combinaciones])
            total_h = sum([d['n_hombres'] for d in datos_combinaciones])
            
            # Totales para sin complementos efectivos
            datos_totales_sin = self.calcular_promedios_efectivos_sb(df_actual, col_sin_comp)
            brecha_total_sin = ((datos_totales_sin['H'] - datos_totales_sin['M']) / datos_totales_sin['M'] * 100) if datos_totales_sin['M'] > 0 else 0
            
            # Totales para con complementos efectivos
            datos_totales_con = self.calcular_promedios_efectivos_sb_complementos(df_actual, col_con_comp)
            brecha_total_con = ((datos_totales_con['H'] - datos_totales_con['M']) / datos_totales_con['M'] * 100) if datos_totales_con['M'] > 0 else 0
            
            # Totales para con extrasalariales efectivos
            datos_totales_extra = self.calcular_promedios_efectivos_sb_complementos(df_actual, col_con_extra)
            brecha_total_extra = ((datos_totales_extra['H'] - datos_totales_extra['M']) / datos_totales_extra['M'] * 100) if datos_totales_extra['M'] > 0 else 0
            
            # Insertar totales al principio
            datos_combinaciones.insert(0, {
                'svpt_puesto': 'Totales',
                'n_mujeres': total_m,
                'n_hombres': total_h,
                'sin_comp_m': datos_totales_sin['M'],
                'sin_comp_h': datos_totales_sin['H'],
                'brecha_sin_comp': brecha_total_sin,
                'con_comp_m': datos_totales_con['M'],
                'con_comp_h': datos_totales_con['H'],
                'brecha_con_comp': brecha_total_con,
                'con_extra_m': datos_totales_extra['M'],
                'con_extra_h': datos_totales_extra['H'],
                'brecha_con_extra': brecha_total_extra
            })
        
        return datos_combinaciones
    
    def generar_datos_svpt_puesto_equiparado(self):
        """Genera los datos para la tabla de retribución por SVPT + Puesto de trabajo EQUIPARADO combinados"""
        print("📊 Calculando datos equiparados por SVPT + Puesto de trabajo...")
        
        # Filtrar datos actuales (sin "Ex" en primera columna)
        df_actual = self.data[self.data.iloc[:, 0] != 'Ex'].copy()
        
        # Verificar que ambas columnas existen
        if 'Nivel SVPT' not in df_actual.columns:
            print("❌ Error: No se encontró la columna 'Nivel SVPT'")
            return []
        
        if 'Puesto de trabajo' not in df_actual.columns:
            print("❌ Error: No se encontró la columna 'Puesto de trabajo'")
            return []
        
        # Crear una columna combinada de SVPT + Puesto de trabajo
        df_actual['SVPT_Puesto'] = df_actual['Nivel SVPT'].astype(str) + ' - ' + df_actual['Puesto de trabajo'].astype(str)
        
        # Obtener combinaciones únicas
        combinaciones = df_actual['SVPT_Puesto'].unique()
        combinaciones = [str(c) for c in combinaciones if pd.notna(c) and str(c) != 'nan - nan']
        datos_combinaciones = []
        
        # Columnas para valores equiparados
        col_sin_comp = 'salario_base_equiparado'  # SB equiparado
        col_con_comp = 'sb_mas_comp_salariales_equiparado'  # SB + Complementos equiparado
        col_con_extra = 'sb_mas_comp_total_equiparado'  # SB + Complementos + Extrasalariales equiparado
        
        for combinacion in sorted(combinaciones):
            # Filtrar por la combinación específica
            df_comb = df_actual[df_actual['SVPT_Puesto'].astype(str) == str(combinacion)]
            
            # Conteos por género
            conteos = df_comb['SEXO'].value_counts()
            n_mujeres = conteos.get('M', 0)
            n_hombres = conteos.get('H', 0)
            
            # SIN complementos (solo SB equiparado) - solo personas con SB > 0
            datos_sin_comp = self.calcular_promedios_equiparados_sb(df_comb, col_sin_comp)
            sin_comp_m = datos_sin_comp['M']
            sin_comp_h = datos_sin_comp['H']
            brecha_sin_comp = ((sin_comp_h - sin_comp_m) / sin_comp_m * 100) if sin_comp_m > 0 else 0
            
            # CON complementos equiparados - todas las personas
            datos_con_comp = self.calcular_promedios_equiparados_sb_complementos(df_comb, col_con_comp)
            con_comp_m = datos_con_comp['M']
            con_comp_h = datos_con_comp['H']
            brecha_con_comp = ((con_comp_h - con_comp_m) / con_comp_m * 100) if con_comp_m > 0 else 0
            
            # CON extrasalariales equiparados - todas las personas
            datos_con_extra = self.calcular_promedios_equiparados_sb_complementos(df_comb, col_con_extra)
            con_extra_m = datos_con_extra['M']
            con_extra_h = datos_con_extra['H']
            brecha_con_extra = ((con_extra_h - con_extra_m) / con_extra_m * 100) if con_extra_m > 0 else 0
            
            datos_combinaciones.append({
                'svpt_puesto': combinacion,
                'n_mujeres': n_mujeres,
                'n_hombres': n_hombres,
                'sin_comp_m': sin_comp_m,
                'sin_comp_h': sin_comp_h,
                'brecha_sin_comp': brecha_sin_comp,
                'con_comp_m': con_comp_m,
                'con_comp_h': con_comp_h,
                'brecha_con_comp': brecha_con_comp,
                'con_extra_m': con_extra_m,
                'con_extra_h': con_extra_h,
                'brecha_con_extra': brecha_con_extra
            })
        
        # Calcular totales equiparados
        if datos_combinaciones:
            total_m = sum([d['n_mujeres'] for d in datos_combinaciones])
            total_h = sum([d['n_hombres'] for d in datos_combinaciones])
            
            # Totales para sin complementos equiparados
            datos_totales_sin = self.calcular_promedios_equiparados_sb(df_actual, col_sin_comp)
            brecha_total_sin = ((datos_totales_sin['H'] - datos_totales_sin['M']) / datos_totales_sin['M'] * 100) if datos_totales_sin['M'] > 0 else 0
            
            # Totales para con complementos equiparados
            datos_totales_con = self.calcular_promedios_equiparados_sb_complementos(df_actual, col_con_comp)
            brecha_total_con = ((datos_totales_con['H'] - datos_totales_con['M']) / datos_totales_con['M'] * 100) if datos_totales_con['M'] > 0 else 0
            
            # Totales para con extrasalariales equiparados
            datos_totales_extra = self.calcular_promedios_equiparados_sb_complementos(df_actual, col_con_extra)
            brecha_total_extra = ((datos_totales_extra['H'] - datos_totales_extra['M']) / datos_totales_extra['M'] * 100) if datos_totales_extra['M'] > 0 else 0
            
            # Insertar totales al principio
            datos_combinaciones.insert(0, {
                'svpt_puesto': 'Totales',
                'n_mujeres': total_m,
                'n_hombres': total_h,
                'sin_comp_m': datos_totales_sin['M'],
                'sin_comp_h': datos_totales_sin['H'],
                'brecha_sin_comp': brecha_total_sin,
                'con_comp_m': datos_totales_con['M'],
                'con_comp_h': datos_totales_con['H'],
                'brecha_con_comp': brecha_total_con,
                'con_extra_m': datos_totales_extra['M'],
                'con_extra_h': datos_totales_extra['H'],
                'brecha_con_extra': brecha_total_extra
            })
        
        return datos_combinaciones
    
    def crear_tabla_svpt_puesto(self, doc, datos_combinaciones):
        """Crea la tabla de retribución por SVPT + Puesto de trabajo combinados"""
        from docx.shared import Cm, Pt
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.shared import OxmlElement, qn
        
        if not datos_combinaciones:
            doc.add_paragraph("No hay datos disponibles para generar la tabla por SVPT + Puesto de trabajo.")
            return
        
        # Crear tabla con cabeceras
        num_filas = len(datos_combinaciones) + 1  # +1 para cabecera
        table = doc.add_table(rows=num_filas, cols=12)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        
        # Definir cabeceras
        headers = [
            'SVPT + Puesto de Trabajo', 'Nº M', 'Nº H',
            'Retribución Promedio SIN Complementos (Mujeres)', 'Retribución Promedio SIN Complementos (Hombres)', 'Brecha Salarial SIN Complementos',
            'Retribución Promedio CON Complementos (Mujeres)', 'Retribución Promedio CON Complementos (Hombres)', 'Brecha Salarial CON Complementos',
            'Retribución Promedio CON Complementos ES (Mujeres)', 'Retribución Promedio CON Complementos ES (Hombres)', 'Brecha Salarial CON Complementos ES'
        ]
        
        # Configurar cabeceras
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            if i < len(header_row.cells):
                cell = header_row.cells[i]
                cell.text = header
                # Aplicar formato de cabecera
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(8)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Llenar datos
        for row_idx, datos in enumerate(datos_combinaciones, 1):
            if row_idx < len(table.rows):
                data_row = table.rows[row_idx]
                
                # Preparar valores
                values = [
                    str(datos['svpt_puesto']),
                    str(datos['n_mujeres']),
                    str(datos['n_hombres']),
                    f"{self.formato_numero_es(datos['sin_comp_m'], 2)} €" if datos['sin_comp_m'] > 0 else "-",
                    f"{self.formato_numero_es(datos['sin_comp_h'], 2)} €" if datos['sin_comp_h'] > 0 else "-",
                    f"{self.formato_brecha_es(datos['brecha_sin_comp'], 2)}%" if datos['sin_comp_m'] > 0 and datos['sin_comp_h'] > 0 else "-",
                    f"{self.formato_numero_es(datos['con_comp_m'], 2)} €" if datos['con_comp_m'] > 0 else "-",
                    f"{self.formato_numero_es(datos['con_comp_h'], 2)} €" if datos['con_comp_h'] > 0 else "-",
                    f"{self.formato_brecha_es(datos['brecha_con_comp'], 2)}%" if datos['con_comp_m'] > 0 and datos['con_comp_h'] > 0 else "-",
                    f"{self.formato_numero_es(datos['con_extra_m'], 2)} €" if datos['con_extra_m'] > 0 else "-",
                    f"{self.formato_numero_es(datos['con_extra_h'], 2)} €" if datos['con_extra_h'] > 0 else "-",
                    f"{self.formato_brecha_es(datos['brecha_con_extra'], 2)}%" if datos['con_extra_m'] > 0 and datos['con_extra_h'] > 0 else "-"
                ]
                
                # Llenar celdas y aplicar colores
                for i, value in enumerate(values):
                    if i < len(data_row.cells):
                        cell = data_row.cells[i]
                        cell.text = value
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph.runs:
                                run.font.size = Pt(8)
                        
                        # Aplicar colores según las reglas de negocio
                        es_total = (datos['svpt_puesto'] == 'Totales')
                        self.aplicar_color_celda_grupo_profesional(cell, i, datos, es_total)
        
        # Agregar espacio después de la tabla
        doc.add_paragraph()
        
        return table
        """Crea la tabla de retribución por SVPT"""
        from docx.shared import Cm, Pt
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.shared import OxmlElement, qn
        
        if not datos_svpt:
            doc.add_paragraph("No hay datos disponibles para generar la tabla por SVPT.")
            return
        
        # Crear tabla con cabeceras
        num_filas = len(datos_svpt) + 1  # +1 para cabecera
        table = doc.add_table(rows=num_filas, cols=12)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        
        # Definir cabeceras
        headers = [
            'Nivel SVPT', 'Nº M', 'Nº H',
            'Retribución Promedio SIN Complementos (Mujeres)', 'Retribución Promedio SIN Complementos (Hombres)', 'Brecha Salarial SIN Complementos',
            'Retribución Promedio CON Complementos (Mujeres)', 'Retribución Promedio CON Complementos (Hombres)', 'Brecha Salarial CON Complementos',
            'Retribución Promedio CON Complementos ES (Mujeres)', 'Retribución Promedio CON Complementos ES (Hombres)', 'Brecha Salarial CON Complementos ES'
        ]
        
        # Configurar cabeceras
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            if i < len(header_row.cells):
                cell = header_row.cells[i]
                cell.text = header
                # Aplicar formato de cabecera
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(8)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Llenar datos
        for row_idx, datos in enumerate(datos_svpt, 1):
            if row_idx < len(table.rows):
                data_row = table.rows[row_idx]
                
                # Preparar valores
                values = [
                    str(datos['nivel']),
                    str(datos['n_mujeres']),
                    str(datos['n_hombres']),
                    f"{self.formato_numero_es(datos['sin_comp_m'], 2)} €" if datos['sin_comp_m'] > 0 else "-",
                    f"{self.formato_numero_es(datos['sin_comp_h'], 2)} €" if datos['sin_comp_h'] > 0 else "-",
                    f"{self.formato_brecha_es(datos['brecha_sin_comp'], 2)}%" if datos['sin_comp_m'] > 0 and datos['sin_comp_h'] > 0 else "-",
                    f"{self.formato_numero_es(datos['con_comp_m'], 2)} €" if datos['con_comp_m'] > 0 else "-",
                    f"{self.formato_numero_es(datos['con_comp_h'], 2)} €" if datos['con_comp_h'] > 0 else "-",
                    f"{self.formato_brecha_es(datos['brecha_con_comp'], 2)}%" if datos['con_comp_m'] > 0 and datos['con_comp_h'] > 0 else "-",
                    f"{self.formato_numero_es(datos['con_extra_m'], 2)} €" if datos['con_extra_m'] > 0 else "-",
                    f"{self.formato_numero_es(datos['con_extra_h'], 2)} €" if datos['con_extra_h'] > 0 else "-",
                    f"{self.formato_brecha_es(datos['brecha_con_extra'], 2)}%" if datos['con_extra_m'] > 0 and datos['con_extra_h'] > 0 else "-"
                ]
                
                # Llenar celdas y aplicar colores
                for i, value in enumerate(values):
                    if i < len(data_row.cells):
                        cell = data_row.cells[i]
                        cell.text = value
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph.runs:
                                run.font.size = Pt(8)
                        
                        # Aplicar colores según las reglas de negocio
                        es_total = (datos['nivel'] == 'Totales')
                        self.aplicar_color_celda_grupo_profesional(cell, i, datos, es_total)
        
        # Agregar espacio después de la tabla
        doc.add_paragraph()
        
        return table
    
    def crear_grafico_barras_svpt_puesto(self, datos_combinaciones, tipo="Efectivo"):
        """
        Crea un gráfico de barras horizontales por SVPT + Puesto de trabajo combinados
        Muestra salarios CON complementos + extrasalariales por género
        """
        # Filtrar solo los totales, pero incluir todas las combinaciones (incluso si solo tienen datos de un género)
        datos_validos = [d for d in datos_combinaciones if d['svpt_puesto'] != 'Totales']
        
        if not datos_validos:
            print("No hay datos válidos para el gráfico de barras por SVPT + Puesto de trabajo")
            return None
        
        # Preparar datos - usar 0 para valores sin datos en lugar de filtrarlos
        combinaciones = [d['svpt_puesto'] for d in datos_validos]
        salarios_m = [d['con_extra_m'] if d['con_extra_m'] > 0 else 0 for d in datos_validos]
        salarios_h = [d['con_extra_h'] if d['con_extra_h'] > 0 else 0 for d in datos_validos]
        
        # Crear figura con tamaño estandarizado
        figsize = self.calcular_tamaño_grafico(len(combinaciones), 'barra_horizontal')
        fig, ax = plt.subplots(figsize=figsize)
        self.configurar_grafico_base(fig, ax)
        
        # Configurar posiciones de las barras
        y_pos = np.arange(len(combinaciones))
        bar_height = 0.35
        
        # Crear barras horizontales con los mismos colores que los gráficos donut
        bars_m = ax.barh(y_pos - bar_height/2, salarios_m, bar_height, 
                        label='Mujeres', color=self.colores_genero['M'], alpha=0.8)
        bars_h = ax.barh(y_pos + bar_height/2, salarios_h, bar_height, 
                        label='Hombres', color=self.colores_genero['H'], alpha=0.8)
        
        # Agregar valores en las barras
        for i, (bar_m, bar_h, sal_m, sal_h) in enumerate(zip(bars_m, bars_h, salarios_m, salarios_h)):
            # Valor para mujeres (solo si tiene datos)
            if sal_m > 0:
                ax.text(bar_m.get_width() + max(salarios_m + salarios_h) * 0.01, 
                       bar_m.get_y() + bar_m.get_height()/2, 
                       f'{self.formato_numero_es(sal_m, 2)} €', 
                       va='center', ha='left', fontsize=7, fontweight='bold')
            
            # Valor para hombres (solo si tiene datos)
            if sal_h > 0:
                ax.text(bar_h.get_width() + max(salarios_m + salarios_h) * 0.01, 
                       bar_h.get_y() + bar_h.get_height()/2, 
                       f'{self.formato_numero_es(sal_h, 2)} €', 
                       va='center', ha='left', fontsize=7, fontweight='bold')
        
        # Configurar ejes
        ax.set_yticks(y_pos)
        ax.set_yticklabels(combinaciones, fontsize=7)
        ax.invert_yaxis()  # Para mostrar la primera combinación arriba
        
        # Configurar etiquetas y título
        ax.set_xlabel('Salario (€)', fontsize=12, fontweight='bold')
        ax.set_ylabel('SVPT + Puesto de Trabajo', fontsize=12, fontweight='bold')
        # Aplicar título estandarizado
        titulo = f'Salarios Medios CON Complementos + ES por SVPT y Puesto de Trabajo {tipo}'
        self.aplicar_estilo_titulo(ax, titulo)
        
        # Configurar leyenda
        ax.legend(loc='lower right', fontsize=11)
        
        # Configurar grid
        ax.grid(True, axis='x', alpha=0.3)
        ax.set_axisbelow(True)
        
        # Ajustar márgenes
        plt.tight_layout()
        
        return fig

    def generar_datos_por_escala_svpt(self, escala, tipo="Efectivo"):
        """
        Genera los datos para una escala SVPT específica (E1, E2, etc.)
        """
        # Filtrar datos actuales (sin "Ex" en primera columna)
        df_actual = self.data[self.data.iloc[:, 0] != 'Ex'].copy()
        
        # Verificar que existan las columnas necesarias
        if 'Nivel SVPT' not in df_actual.columns:
            print(f"❌ Error: No se encontró la columna 'Nivel SVPT' para escala {escala}")
            return []
        
        if 'Puesto de trabajo' not in df_actual.columns:
            print(f"❌ Error: No se encontró la columna 'Puesto de trabajo' para escala {escala}")
            return []
        
        # Filtrar por la escala específica
        df_escala = df_actual[df_actual['Nivel SVPT'].astype(str) == str(escala)]
        
        if len(df_escala) == 0:
            print(f"⚠️ No hay datos para la escala {escala}")
            return []
        
        # Obtener puestos únicos dentro de esta escala
        puestos = df_escala['Puesto de trabajo'].unique()
        puestos = [str(p) for p in puestos if pd.notna(p)]
        datos_puestos = []
        
        # Columnas según el tipo
        if tipo == "Efectivo":
            col_sin_comp = 'Salario base efectivo Total'
            col_con_comp = 'Salario base anual + complementos Total'
            col_con_extra = 'Salario base anual + complementos + Extrasalariales Total'
        else:  # Equiparado
            col_sin_comp = 'salario_base_equiparado'
            col_con_comp = 'sb_mas_comp_salariales_equiparado'
            col_con_extra = 'sb_mas_comp_total_equiparado'
        
        for puesto in sorted(puestos):
            # Filtrar datos del puesto dentro de la escala
            df_puesto = df_escala[df_escala['Puesto de trabajo'].astype(str) == puesto]
            
            # Contar por género
            conteos = df_puesto['SEXO'].value_counts()
            n_mujeres = conteos.get('M', 0)
            n_hombres = conteos.get('H', 0)
            
            if n_mujeres == 0 and n_hombres == 0:
                continue
            
            # Calcular retribución CON complementos + extrasalariales
            if tipo == "Efectivo":
                # Para efectivo, incluir todas las personas
                df_calc = df_puesto[df_puesto[col_con_extra].notna()]
            else:
                # Para equiparado, incluir todas las personas
                df_calc = df_puesto[df_puesto[col_con_extra].notna()]
            
            if len(df_calc) > 0:
                promedios = df_calc.groupby('SEXO')[col_con_extra].mean()
                con_extra_m = promedios.get('M', 0)
                con_extra_h = promedios.get('H', 0)
            else:
                con_extra_m = con_extra_h = 0
            
            # Calcular promedio de puntos para este puesto
            puntos_validos = df_puesto['Puntos'].dropna() if 'Puntos' in df_puesto.columns else []
            promedio_puntos = puntos_validos.mean() if len(puntos_validos) > 0 else 0
            
            datos_puestos.append({
                'puesto': puesto,
                'escala': escala,
                'n_mujeres': n_mujeres,
                'n_hombres': n_hombres,
                'con_extra_m': con_extra_m,
                'con_extra_h': con_extra_h,
                'puntos': promedio_puntos
            })
        
        return datos_puestos

    def obtener_escalas_svpt(self):
        """
        Obtiene las escalas SVPT disponibles en los datos
        """
        # Filtrar datos actuales (sin "Ex" en primera columna)
        df_actual = self.data[self.data.iloc[:, 0] != 'Ex'].copy()
        
        if 'Nivel SVPT' not in df_actual.columns:
            print("❌ Error: No se encontró la columna 'Nivel SVPT'")
            return []
        
        # Obtener escalas únicas
        escalas = df_actual['Nivel SVPT'].dropna().unique()
        escalas_validas = []
        
        for escala in escalas:
            escala_str = str(escala).strip()
            if escala_str and escala_str != 'nan' and len(escala_str) > 0:
                escalas_validas.append(escala_str)
        
        # Ordenar escalas de menor a mayor número (Escala 1, Escala 2, etc.)
        try:
            def ordenar_escala(escala_str):
                # Extraer el número de la escala
                import re
                match = re.search(r'(\d+)', escala_str)
                if match:
                    return int(match.group(1))
                return float('inf')
            
            # Ordenar de menor a mayor
            escalas_ordenadas = sorted(escalas_validas, key=ordenar_escala)
        except:
            escalas_ordenadas = sorted(escalas_validas)
        
        print(f"📋 Escalas SVPT encontradas: {escalas_ordenadas}")
        return escalas_ordenadas

    def crear_grafico_barras_por_escala(self, datos_puestos, escala, tipo="Efectivo"):
        """
        Crea un gráfico de barras verticales para una escala SVPT específica
        Similar al formato del ejemplo adjunto
        """
        if not datos_puestos:
            print(f"No hay datos válidos para el gráfico de escala {escala}")
            return None
        
        # Preparar datos
        puestos = [d['puesto'] for d in datos_puestos]
        salarios_m = [d['con_extra_m'] if d['con_extra_m'] > 0 else 0 for d in datos_puestos]
        salarios_h = [d['con_extra_h'] if d['con_extra_h'] > 0 else 0 for d in datos_puestos]
        puntos_promedio = [d['puntos'] for d in datos_puestos]
        
        # Crear figura con tamaño ajustado al documento
        # Calcular ancho óptimo: mínimo 10, máximo 16 para que no se salga del documento
        ancho_fig = max(10, min(16, len(puestos) * 1.8))
        alto_fig = 8  # Altura fija más compacta
        
        fig, ax1 = plt.subplots(figsize=(ancho_fig, alto_fig))
        fig.patch.set_facecolor('white')
        
        # Configurar posiciones de las barras
        x_pos = np.arange(len(puestos))
        bar_width = 0.35
        
        # Crear barras verticales
        bars_h = ax1.bar(x_pos - bar_width/2, salarios_h, bar_width, 
                        label='Hombres', color=self.colores_genero['H'], alpha=0.9)
        bars_m = ax1.bar(x_pos + bar_width/2, salarios_m, bar_width, 
                        label='Mujeres', color=self.colores_genero['M'], alpha=0.9)
        
        # Ajustar posición de valores usando configuración centralizada
        valores_salarios = salarios_h + salarios_m
        offset_salarios = self.ajustar_posicion_valores(ax1, valores_salarios)
        
        # Agregar valores en las barras con posicionamiento mejorado
        for i, (bar_h, bar_m, sal_h, sal_m) in enumerate(zip(bars_h, bars_m, salarios_h, salarios_m)):
            if sal_h > 0:
                ax1.text(bar_h.get_x() + bar_h.get_width()/2, 
                        bar_h.get_height() + offset_salarios, 
                        f'{self.formato_numero_es(sal_h, 2)} €', 
                        ha='center', va='bottom', 
                        fontsize=self.config_graficos['valores']['fontsize'], 
                        fontweight=self.config_graficos['valores']['fontweight'])
            
            if sal_m > 0:
                ax1.text(bar_m.get_x() + bar_m.get_width()/2, 
                        bar_m.get_height() + offset_salarios, 
                        f'{self.formato_numero_es(sal_m, 2)} €', 
                        ha='center', va='bottom', 
                        fontsize=self.config_graficos['valores']['fontsize'], 
                        fontweight=self.config_graficos['valores']['fontweight'])
        
        # Configurar primer eje Y (salarios)
        ax1.set_ylabel('Salario Medio CON Complementos (€)', fontsize=14, fontweight='bold')
        ax1.tick_params(axis='y', labelsize=12)
        
        # Configurar eje X
        ax1.set_xticks(x_pos)
        ax1.set_xticklabels(puestos, rotation=45, ha='right', fontsize=12)
        ax1.set_xlabel('Puesto de Trabajo', fontsize=14, fontweight='bold')
        
        # Crear segundo eje Y (puntos) - línea
        ax2 = ax1.twinx()
        
        # Solo agregar línea de puntos si hay datos válidos
        puntos_validos = [p for p in puntos_promedio if p > 0]
        if len(puntos_validos) > 0:
            # Crear línea de puntos
            line_puntos = ax2.plot(x_pos, puntos_promedio, 'o-', linewidth=3, markersize=8, 
                                  label='Puntos', color='black')
            
            # Ajustar posición de valores de puntos
            offset_puntos = self.ajustar_posicion_valores(ax2, puntos_promedio)
            
            # Agregar valores de puntos con posicionamiento mejorado
            for i, puntos in enumerate(puntos_promedio):
                if puntos > 0:
                    ax2.text(i, puntos + offset_puntos, 
                            f'{puntos:.1f}', 
                            ha='center', va='bottom', 
                            fontsize=self.config_graficos['valores']['fontsize'], 
                            fontweight=self.config_graficos['valores']['fontweight'],
                            color='black')
            
            # Configurar segundo eje Y
            ax2.set_ylabel('Puntos', fontsize=14, fontweight='bold')
            ax2.tick_params(axis='y', labelsize=12)
        
        # Configurar título siguiendo el formato del ejemplo
        if tipo == "Efectivo":
            titulo = f'Salarios Medios CON Complementos por Puesto Trabajo Efectivo ({escala})'
        else:
            titulo = f'Salarios Medios CON Complementos por Puesto Trabajo Equiparado ({escala})'
        
        # Aplicar título estandarizado
        self.aplicar_estilo_titulo(ax1, titulo)
        
        # Configurar leyendas
        lines1, labels1 = ax1.get_legend_handles_labels()
        if len(puntos_validos) > 0:
            lines2, labels2 = ax2.get_legend_handles_labels()
            ax1.legend(lines1 + lines2, labels1 + labels2, loc='upper left', fontsize=12)
        else:
            ax1.legend(loc='upper left', fontsize=12)
        
        # Configurar grid
        ax1.grid(True, axis='y', alpha=0.3, linestyle='-', linewidth=0.5)
        ax1.set_axisbelow(True)
        
        # Ajustar márgenes
        plt.tight_layout()
        
        return fig

    def crear_grafico_barras_vertical_svpt_puesto(self, datos_combinaciones, tipo="Efectivo"):
        """
        Crea un gráfico de barras verticales con doble eje Y por SVPT + Puesto de trabajo combinados
        Eje Y izquierdo: Salarios CON complementos + extrasalariales por género (barras)
        Eje Y derecho: Puntos promedio (línea)
        """
        # Filtrar solo los datos válidos (no totales)
        datos_validos = [d for d in datos_combinaciones if d['svpt_puesto'] != 'Totales']
        
        if not datos_validos:
            print("No hay datos válidos para el gráfico de barras verticales por SVPT + Puesto de trabajo")
            return None
        
        # Preparar datos para las barras
        combinaciones = [d['svpt_puesto'] for d in datos_validos]
        salarios_m = [d['con_extra_m'] if d['con_extra_m'] > 0 else 0 for d in datos_validos]
        salarios_h = [d['con_extra_h'] if d['con_extra_h'] > 0 else 0 for d in datos_validos]
        
        # Obtener datos de puntos promedio por cada combinación SVPT + Puesto
        puntos_promedio = []
        for datos in datos_validos:
            # Buscar datos de puntos en el dataframe original
            svpt_puesto = datos['svpt_puesto']
            
            # Filtrar el dataframe para esta combinación específica
            df_combinacion = self.data[
                (self.data['Nivel SVPT'].astype(str) + ' - ' + self.data['Puesto de trabajo'].astype(str)) == svpt_puesto
            ]
            
            # Calcular promedio de puntos para esta combinación
            if 'Puntos' in self.data.columns and len(df_combinacion) > 0:
                puntos_validos = df_combinacion['Puntos'].dropna()
                if len(puntos_validos) > 0:
                    promedio_puntos = puntos_validos.mean()
                    puntos_promedio.append(promedio_puntos)
                else:
                    puntos_promedio.append(0)
            else:
                puntos_promedio.append(0)
        
        # Crear figura con tamaño estandarizado
        figsize = self.calcular_tamaño_grafico(len(combinaciones), 'barra_vertical')
        fig, ax1 = plt.subplots(figsize=figsize)
        self.configurar_grafico_base(fig, ax1)
        
        # Configurar posiciones de las barras
        x_pos = np.arange(len(combinaciones))
        bar_width = 0.35
        
        # Crear barras verticales con los colores correctos de género
        bars_m = ax1.bar(x_pos - bar_width/2, salarios_m, bar_width, 
                        label='Mujeres', color=self.colores_genero['M'], alpha=0.8)
        bars_h = ax1.bar(x_pos + bar_width/2, salarios_h, bar_width, 
                        label='Hombres', color=self.colores_genero['H'], alpha=0.8)
        
        # Ajustar posición de valores usando configuración centralizada
        valores_salarios = salarios_m + salarios_h
        offset_salarios = self.ajustar_posicion_valores(ax1, valores_salarios)
        
        # Agregar valores en las barras con posicionamiento mejorado
        for i, (bar_m, bar_h, sal_m, sal_h) in enumerate(zip(bars_m, bars_h, salarios_m, salarios_h)):
            # Valor para mujeres (solo si tiene datos)
            if sal_m > 0:
                ax1.text(bar_m.get_x() + bar_m.get_width()/2, 
                        bar_m.get_height() + offset_salarios, 
                        f'{self.formato_numero_es(sal_m, 0)}', 
                        ha='center', va='bottom', 
                        fontsize=self.config_graficos['valores']['fontsize'], 
                        fontweight=self.config_graficos['valores']['fontweight'],
                        rotation=90)
            
            # Valor para hombres (solo si tiene datos)
            if sal_h > 0:
                ax1.text(bar_h.get_x() + bar_h.get_width()/2, 
                        bar_h.get_height() + offset_salarios, 
                        f'{self.formato_numero_es(sal_h, 0)}', 
                        ha='center', va='bottom', 
                        fontsize=self.config_graficos['valores']['fontsize'], 
                        fontweight=self.config_graficos['valores']['fontweight'],
                        rotation=90)
        
        # Configurar primer eje Y (salarios)
        ax1.set_xlabel('SVPT + Puesto de Trabajo', fontsize=14, fontweight='bold')
        ax1.set_ylabel('Salario Medio CON Complementos (€)', fontsize=14, fontweight='bold', color='black')
        ax1.tick_params(axis='y', labelcolor='black', labelsize=12)
        
        # Configurar eje X con mejor legibilidad
        ax1.set_xticks(x_pos)
        ax1.set_xticklabels(combinaciones, rotation=45, ha='right', fontsize=10)
        
        # Crear segundo eje Y (puntos) - línea
        ax2 = ax1.twinx()
        
        # Solo agregar línea de puntos si hay datos válidos
        puntos_validos = [p for p in puntos_promedio if p > 0]
        if len(puntos_validos) > 0:
            # Crear línea de puntos
            line_puntos = ax2.plot(x_pos, puntos_promedio, 'o-', linewidth=3, markersize=8, 
                                  label='Puntos', color='black')
            
            # Ajustar posición de valores de puntos
            offset_puntos = self.ajustar_posicion_valores(ax2, puntos_promedio)
            
            # Agregar valores de puntos con posicionamiento mejorado
            for i, puntos in enumerate(puntos_promedio):
                if puntos > 0:
                    ax2.text(i, puntos + offset_puntos, 
                            f'{puntos:.0f}', 
                            ha='center', va='bottom', 
                            fontsize=self.config_graficos['valores']['fontsize'], 
                            fontweight=self.config_graficos['valores']['fontweight'],
                            color='black')
            
            # Configurar segundo eje Y
            ax2.set_ylabel('Puntos', fontsize=14, fontweight='bold', color='black')
            ax2.tick_params(axis='y', labelcolor='black', labelsize=12)
        
        # Aplicar título estandarizado
        titulo = f'Salarios Medios CON Complementos por Puesto Trabajo {tipo}'
        self.aplicar_estilo_titulo(ax1, titulo)
        
        # Configurar leyendas
        lines1, labels1 = ax1.get_legend_handles_labels()
        if len(puntos_validos) > 0:
            lines2, labels2 = ax2.get_legend_handles_labels()
            ax1.legend(lines1 + lines2, labels1 + labels2, loc='upper left', fontsize=12)
        else:
            ax1.legend(loc='upper left', fontsize=12)
        
        # Configurar grid
        ax1.grid(True, axis='y', alpha=0.3)
        ax1.set_axisbelow(True)
        
        # Ajustar márgenes
        plt.tight_layout()
        
        return fig

# Archivo de configuración de ejemplo (report_config.yaml)
def create_sample_config():
    """Crea un archivo de configuración de ejemplo para registro retributivo"""
    config = {
        'excel_file': '',  # Se determina automáticamente
        'template_word': 'plantilla_informe.docx',  # Opcional
        'output_file': '05_INFORMES/registro_retributivo_final.docx',  # Se generará timestamp dinámico
        'charts': {
            'salario_base_efectivo': {
                'type': 'donut',
                'data_columns': ['Salario base efectivo Total'],
                'metodo': 'efectivos_sb',
                'title': 'Comparación Salario Base Efectivo Total por Género',
                'subtitulo': 'Análisis de igualdad retributiva - Salario base efectivamente percibido (solo SB > 0)',
                'marker': '{grafico_sb_efectivo}'
            },
            'sb_complementos_efectivo': {
                'type': 'donut',
                'data_columns': ['Salario base anual + complementos Total'],
                'metodo': 'efectivos_sb_complementos',
                'title': 'Salario Base + Complementos Salariales Efectivos por Género',
                'subtitulo': 'Incluye salario base y complementos salariales efectivamente percibidos (todas las personas)',
                'marker': '{grafico_sb_comp_efectivo}'
            },
            'sb_total_efectivo': {
                'type': 'donut',
                'data_columns': ['Salario base anual + complementos + Extrasalariales Total'],
                'metodo': 'efectivos_sb_complementos',
                'title': 'SB + Complementos + Extrasalariales Efectivos por Género',
                'subtitulo': 'Retribución total efectiva incluyendo todos los conceptos (todas las personas)',
                'marker': '{grafico_sb_total_efectivo}'
            },
            'salario_base_equiparado': {
                'type': 'donut',
                'data_columns': ['salario_base_equiparado'],
                'metodo': 'equiparados_sb',
                'title': 'Comparación Salario Base Equiparado por Género',
                'subtitulo': 'Salario base normalizado a jornada completa y 12 meses (solo SB > 0)',
                'marker': '{grafico_sb_equiparado}'
            },
            'sb_complementos_equiparado': {
                'type': 'donut',
                'data_columns': ['sb_mas_comp_salariales_equiparado'],
                'metodo': 'equiparados_sb_complementos',
                'title': 'Salario Base + Complementos Salariales Equiparados por Género',
                'subtitulo': 'SB + complementos salariales normalizados a jornada completa y 12 meses (todas las personas)',
                'marker': '{grafico_sb_comp_equiparado}'
            },
            'sb_total_equiparado': {
                'type': 'donut',
                'data_columns': ['sb_mas_comp_total_equiparado'],
                'metodo': 'equiparados_sb_complementos',
                'title': 'SB + Complementos + Extrasalariales Equiparados por Género',
                'subtitulo': 'Retribución total equiparada: SB + complementos salariales y extrasalariales (todas las personas)',
                'marker': '{grafico_sb_total_equiparado}'
            }
        }
    }
    
    with open('report_config.yaml', 'w', encoding='utf-8') as f:
        yaml.dump(config, f, default_flow_style=False, allow_unicode=True)
    
    print("Archivo de configuración creado: report_config.yaml")

# Script principal
if __name__ == "__main__":
    # Crear configuración de ejemplo si no existe
    if not os.path.exists('report_config.yaml'):
        create_sample_config()
    
    # Generar el reporte
    system = AutomatedReportSystem()
    system.generate_report()