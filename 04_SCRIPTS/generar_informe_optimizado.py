# -*- coding: utf-8 -*-
"""
Sistema Optimizado para Generación de Informes de Registro Retributivo
Versión refactorizada sin redundancia con funciones reutilizables
"""

import sys
import os
import io
from pathlib import Path
from datetime import datetime
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import warnings
import logging

# ==================== CONFIGURACIÓN GLOBAL ====================

# Configurar codificación UTF-8
if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding='utf-8')
        sys.stderr.reconfigure(encoding='utf-8')
    except:
        pass

# Configurar logging para capturar warnings en archivo
log_dir = Path(__file__).parent.parent / '03_LOGS'
log_dir.mkdir(exist_ok=True)
log_file = log_dir / f'informe_{datetime.now().strftime("%Y%m%d")}.log'

logging.basicConfig(
    level=logging.WARNING,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(log_file, encoding='utf-8'),
    ]
)

# Suprimir warnings de pandas y numpy en consola (pero se guardan en log)
warnings.filterwarnings('ignore', category=pd.errors.SettingWithCopyWarning)
warnings.filterwarnings('ignore', category=RuntimeWarning)
logging.captureWarnings(True)

# Estilos matplotlib
plt.rcParams['font.size'] = 14  # Aumentado de 10 a 14 para mejor legibilidad
plt.rcParams['font.family'] = 'sans-serif'
plt.rcParams['axes.labelsize'] = 16  # Tamaño de etiquetas de ejes
plt.rcParams['xtick.labelsize'] = 14  # Tamaño de etiquetas del eje X
plt.rcParams['ytick.labelsize'] = 14  # Tamaño de etiquetas del eje Y
plt.rcParams['legend.fontsize'] = 14  # Tamaño de la leyenda
sns.set_style("whitegrid")

# Colores corporativos

# Hex: Gráficos con matplotlib y seaborn
COLORES = {
    'mujer': '#1e4389',      # Azul
    'hombre': '#ea5d41',     # Rojo/Naranja
    'neutro': '#000000'      # Negro
}

# RGB: Word (requiere valores 0-255)
COLORES_RGB = {
    'mujer': RGBColor(30, 67, 137),
    'hombre': RGBColor(234, 93, 65),
    'neutro': RGBColor(0, 0, 0)
}

# Colores para tablas - estilo moderno y minimalista
COLOR_FONDO_ENCABEZADO = RGBColor(245, 245, 245)  # Gris muy clarito para encabezados
COLOR_FONDO_TOTAL = RGBColor(245, 245, 245)  # Mismo gris que encabezados para filas de totales

# Mapeo de columnas del Excel
COLS = {
    'reg': 'Reg.',
    'sexo': 'Sexo',
    'grupo_prof': 'Grupo profesional',
    'nivel_svpt': 'Nivel SVPT',
    'puesto': 'Puesto de trabajo',
    'nivel_convenio': 'Nivel Convenio Colectivo',

    # Efectivos (siempre usar columnas Total)
    'sb_efectivo': 'Salario base efectivo Total',
    'sbc_efectivo': 'Salario base anual + complementos Total',
    'sbce_efectivo': 'Salario base anual + complementos + Extrasalariales Total',

    # Equiparados
    'sb_equiparado': 'salario_base_equiparado',
    'sbc_equiparado': 'sb_mas_comp_salariales_equiparado',
    'sbce_equiparado': 'sb_mas_comp_total_equiparado',

    # Complementos Salariales
    'comp_efectivo': 'Compltos Salariales efectivo Total',
    'comp_equiparado': 'complementos_salariales_equiparados',

    # Complementos Extrasalariales
    'extra_efectivo': 'Compltos Extrasalariales efectivo Total',
    'extra_equiparado': 'complementos_extrasalariales_equiparados',

    # Puntos
    'puntos': 'Puntos'
}

# Configuración de gráficos
GRAFICO_CFG = {
    'dpi': 300,
    'ancho_doc': 6.5,  # inches (ancho completo de página)
    'ancho_donut': 3.0,  # inches (50% de la página para 2 donuts lado a lado)
    'figsize_donut': (8, 8),  # Tamaño del gráfico
    'figsize_barras': (14, 8),
    'titulo_size': 18,  # Aumentado de 14 a 18
    'titulo_color': '#cc0000',
    'titulo_weight': 'bold',
    'etiqueta_size': 14,  # Aumentado de 10 a 14
    'valor_size': 13  # Aumentado de 9 a 13
}


# ==================== FUNCIONES AUXILIARES ====================

def log(mensaje, tipo='INFO'):
    """Log estandarizado"""
    prefijos = {'INFO': '[INFO]', 'OK': '[✓]', 'ERROR': '[✗]', 'WARN': '[!]'}
    print(f"{prefijos.get(tipo, '[INFO]')} {mensaje}", flush=True)


def formato_numero_es(numero, decimales=2):
    """Formatea número con estilo español"""
    if pd.isna(numero) or numero is None:
        return "0,00" if decimales == 2 else "0"
    formato = f"{{:,.{decimales}f}}"
    num_str = formato.format(numero)
    return num_str.replace(',', 'X').replace('.', ',').replace('X', '.')


def calcular_brecha(valor_h, valor_m):
    """Calcula brecha salarial: ((H - M) / H) * 100"""
    if valor_h is None or valor_h <= 0 or valor_m is None or valor_m < 0:
        return None
    return ((valor_h - valor_m) / valor_h) * 100


def aplicar_filtro_privacidad_datos(datos_lista):
    """
    Filtra elementos de una lista de datos para proteger la privacidad.
    Oculta datos cuando hay exactamente 1 empleado de cualquier sexo (con al menos 1 del otro).
    Esto previene la identificación de empleados individuales (protección LOPD/RGPD).

    Args:
        datos_lista: Lista de diccionarios con claves 'n_M' y 'n_H' (o 'total')

    Returns:
        Lista filtrada sin elementos donde n_M == 1 o n_H == 1
    """
    if not datos_lista:
        return datos_lista

    datos_filtrados = []
    items_eliminados = 0

    for item in datos_lista:
        # Obtener número de empleados por sexo
        n_m = item.get('n_M', 0)
        n_h = item.get('n_H', 0)

        # Ocultar si hay exactamente 1 empleado de cualquier sexo
        # (permite mostrar cuando hay 0 o 2+ de cada sexo)
        if n_m == 1 or n_h == 1:
            items_eliminados += 1
            categoria = item.get('categoria', item.get('grupo', item.get('complemento', 'desconocido')))
            log(f"Privacidad: Eliminado '{categoria}' (n_M={n_m}, n_H={n_h}) - empleado identificable", 'WARN')
        else:
            datos_filtrados.append(item)

    if items_eliminados > 0:
        log(f"Privacidad: {items_eliminados} elemento(s) eliminado(s) por tener empleados identificables", 'WARN')

    return datos_filtrados


def verificar_privacidad_tabla(datos):
    """
    Verifica si los datos de una tabla cumplen con los requisitos de privacidad.
    Retorna None si los datos no deben mostrarse (hay exactamente 1 empleado de cualquier sexo).

    Args:
        datos: dict con estructura {'efectivo': {...}, 'equiparado': {...}}
               donde cada valor tiene 'n_M' y 'n_H'

    Returns:
        datos originales si es válido, None si debe ocultarse
    """
    if not datos or 'efectivo' not in datos:
        return datos

    # Verificar número de empleados por sexo
    n_m = datos['efectivo'].get('n_M', 0)
    n_h = datos['efectivo'].get('n_H', 0)

    # Ocultar si hay exactamente 1 empleado de cualquier sexo
    if n_m == 1 or n_h == 1:
        log(f"Privacidad: Tabla eliminada (n_M={n_m}, n_H={n_h}) - empleado identificable", 'WARN')
        return None

    return datos


def filtrar_datos(df, columna_salario, tipo_calculo):
    """
    Filtra datos según el tipo de cálculo.

    tipo_calculo:
    - 'efectivos_sb': solo SB > 0 (excluir 'Ex' en Reg.)
    - 'efectivos_sb_complementos': todos (excluir 'Ex' en Reg.)
    - 'equiparados_sb': solo SB > 0 (excluir 'Ex' en Reg.)
    - 'equiparados_sb_complementos': todos (excluir 'Ex' en Reg.)
    """
    # Excluir registros con Reg. = 'Ex'
    df_filtrado = df[df[COLS['reg']] != 'Ex'].copy()

    # Para cálculos de salario base, filtrar solo salarios > 0
    if 'sb' in tipo_calculo and 'complementos' not in tipo_calculo:
        # Solo personas con salario base > 0
        df_filtrado = df_filtrado[df_filtrado[columna_salario] > 0]

    return df_filtrado


def calcular_estadistico(df, columna_salario, tipo_calculo, metodo='media', grupo_col=None, grupo_valor=None):
    """
    Calcula estadístico (media o mediana) por género.

    Returns:
        dict: {'M': valor_mujeres, 'H': valor_hombres, 'n_M': count_m, 'n_H': count_h}
    """
    # Filtrar datos
    df_filtrado = filtrar_datos(df, columna_salario, tipo_calculo)

    # Aplicar filtro adicional si hay grupo
    if grupo_col and grupo_valor is not None:
        # Convertir ambos a string para comparación consistente
        # Esto maneja casos donde la columna tiene int (4, 5) y str ("4 turnos")
        df_filtrado = df_filtrado[df_filtrado[grupo_col].astype(str) == str(grupo_valor)]

    if df_filtrado.empty:
        return {'M': 0, 'H': 0, 'n_M': 0, 'n_H': 0}

    # Calcular por género
    resultado = {}
    for sexo in ['M', 'H']:
        df_sexo = df_filtrado[df_filtrado[COLS['sexo']] == sexo]
        resultado[f'n_{sexo}'] = len(df_sexo)

        if len(df_sexo) > 0:
            if metodo == 'media':
                resultado[sexo] = df_sexo[columna_salario].mean()
            else:  # mediana
                resultado[sexo] = df_sexo[columna_salario].median()
        else:
            resultado[sexo] = 0

    return resultado


def reformatear_etiqueta_escala(etiqueta):
    """
    Reformatea etiquetas de 'Escala X + Nombre' a 'Nombre - EX'
    Ejemplo: 'Escala 2 + Offside Leader' -> 'Offside Leader - E2'
    """
    import re

    if not etiqueta or not isinstance(etiqueta, str):
        return etiqueta

    # Patrón para detectar 'Escala X + Nombre'
    patron = r'Escala\s+(\d+)\s*\+\s*(.+)'
    match = re.match(patron, etiqueta.strip(), re.IGNORECASE)

    if match:
        numero_escala = match.group(1)
        nombre = match.group(2).strip()
        return f'{nombre} - E{numero_escala}'

    return etiqueta


def obtener_escalas_svpt(df):
    """Extrae las escalas de la columna Nivel SVPT (ej: 'Escala 1', 'Escala 2', etc.)"""
    escalas = set()
    for nivel in df[COLS['nivel_svpt']].dropna().unique():
        nivel_str = str(nivel).strip()
        # Buscar patrones como "Escala 1", "Escala 2", etc.
        if 'escala' in nivel_str.lower():
            # Extraer el número de escala
            partes = nivel_str.lower().split('escala')
            if len(partes) > 1:
                num_escala = partes[1].strip().split()[0] if partes[1].strip() else ''
                if num_escala.isdigit():
                    escalas.add(f'Escala {num_escala}')

    # Ordenar las escalas numéricamente
    def ordenar_escala(escala_str):
        try:
            return int(escala_str.split()[-1])
        except:
            return 999

    return sorted(list(escalas), key=ordenar_escala)


# ==================== GENERACIÓN DE GRÁFICOS ====================

def ocultar_bordes_tabla(tabla):
    """Oculta completamente los bordes de una tabla de Word"""
    # Configurar bordes de la tabla completa
    tbl = tabla._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Crear bordes de tabla como "none"
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'FFFFFF')
        tblBorders.append(border)

    tblPr.append(tblBorders)

    # También configurar bordes de celda
    for row in tabla.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()

            # Crear elemento de bordes de celda
            tcBorders = OxmlElement('w:tcBorders')

            # Configurar cada borde como ninguno
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                border.set(qn('w:sz'), '0')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'FFFFFF')
                tcBorders.append(border)

            tcPr.append(tcBorders)

def crear_grafico_donut(datos, titulo, nombre_archivo="temp_donut.png"):
    """
    Crea un gráfico de donut comparando hombres y mujeres.

    datos: dict con claves 'M', 'H', 'brecha', opcionalmente 'n_M' y 'n_H'
    """
    valor_m = datos['M']
    valor_h = datos['H']
    brecha = datos.get('brecha', 0)

    # Verificar privacidad si hay información de conteo
    n_m = datos.get('n_M', 0)
    n_h = datos.get('n_H', 0)
    if n_m > 0 or n_h > 0:  # Si hay información de conteo
        # Ocultar si hay exactamente 1 empleado de cualquier sexo
        if n_m == 1 or n_h == 1:
            log(f"Gráfico '{titulo}' omitido por privacidad (n_M={n_m}, n_H={n_h}) - empleado identificable", 'WARN')
            return None

    if valor_m == 0 and valor_h == 0:
        log(f"Sin datos para gráfico: {titulo}", 'WARN')
        return None

    fig, ax = plt.subplots(figsize=GRAFICO_CFG['figsize_donut'])
    fig.patch.set_facecolor('white')

    # Datos
    valores = [valor_m, valor_h]
    colores = [COLORES['mujer'], COLORES['hombre']]

    # Crear donut con mayor grosor
    wedges, _ = ax.pie(
        valores,
        labels=['', ''],  # Sin etiquetas automáticas
        colors=colores,
        startangle=90,
        wedgeprops=dict(width=0.5, edgecolor='white', linewidth=2)  # Grosor del donut aumentado de 0.4 a 0.5
    )

    # Agregar valores monetarios en cada segmento
    for wedge, valor in zip(wedges, valores):
        ang = (wedge.theta2 - wedge.theta1) / 2 + wedge.theta1

        # Posición en el centro del anillo del donut
        radio_centro = 0.80
        x_centro = np.cos(np.deg2rad(ang)) * radio_centro
        y_centro = np.sin(np.deg2rad(ang)) * radio_centro

        # Valor monetario en blanco
        ax.text(x_centro, y_centro, formato_numero_es(valor, 2) + ' €',
                ha='center', va='center',
                fontsize=14, fontweight='bold', color='white')  # Aumentado de 14 a 16

    # Crear hueco central (más pequeño por el mayor grosor)
    centro = plt.Circle((0, 0), 0.55, fc='white')
    ax.add_artist(centro)

    # Determinar color de la brecha según si favorece a hombres o mujeres
    # Brecha positiva = favorece a hombres (rojo), brecha negativa = favorece a mujeres (azul)
    if brecha > 0:
        color_brecha = '#E74C3C'  # Rojo cuando favorece a hombres
    elif brecha < 0:
        color_brecha = '#3498DB'  # Azul cuando favorece a mujeres
    else:
        color_brecha = GRAFICO_CFG['titulo_color']  # Color por defecto si brecha es 0

    # Añadir brecha en el centro con color según favorece
    texto_brecha = f"Brecha Salarial\n{formato_numero_es(abs(brecha), 2)}%"
    ax.text(0, 0, texto_brecha, ha='center', va='center',
            fontsize=22, fontweight='bold', color=color_brecha)  # Aumentado de 18 a 22

    # Título
    ax.set_title(titulo, fontsize=GRAFICO_CFG['titulo_size'],
                 fontweight=GRAFICO_CFG['titulo_weight'],
                 color=GRAFICO_CFG['titulo_color'], pad=20)

    # Agregar leyenda pequeña más hacia el centro
    from matplotlib.patches import Patch
    legend_elements = [
        Patch(facecolor=COLORES['mujer'], label='Mujeres'),
        Patch(facecolor=COLORES['hombre'], label='Hombres')
    ]
    ax.legend(handles=legend_elements, loc='lower right', frameon=False,
              fontsize=14, bbox_to_anchor=(1.0, -0.05))  # Aumentado de 10 a 14

    plt.tight_layout()
    plt.savefig(nombre_archivo, dpi=GRAFICO_CFG['dpi'], bbox_inches='tight', facecolor='white')
    plt.close()

    return nombre_archivo


def crear_grafico_barras(datos_lista, titulo, tipo_valor, nombre_archivo="temp_barras.png", orientacion='horizontal', mostrar_titulo=True, mostrar_leyenda=True, aplicar_filtro_privacidad=True, etiqueta_categoria=None):
    """
    Crea un gráfico de barras agrupadas.

    datos_lista: lista de dicts con claves 'categoria', 'M', 'H', 'n_M', 'n_H'
    tipo_valor: string para el eje (ej: "Salario Medio (€)")
    orientacion: 'horizontal' o 'vertical'
    aplicar_filtro_privacidad: si True (por defecto), aplica filtro de privacidad
    etiqueta_categoria: string opcional para etiqueta del eje de categorías (solo para gráficos verticales)
    """
    if not datos_lista:
        log(f"Sin datos para gráfico: {titulo}", 'WARN')
        return None

    # Aplicar filtro de privacidad solo si se solicita
    if aplicar_filtro_privacidad:
        datos_lista = aplicar_filtro_privacidad_datos(datos_lista)

    if not datos_lista:
        log(f"Gráfico '{titulo}' omitido - ningún elemento cumple requisitos de privacidad", 'WARN')
        return None

    # Filtrar categorías sin datos
    datos_lista = [d for d in datos_lista if d['M'] > 0 or d['H'] > 0]

    if not datos_lista:
        return None

    categorias = [reformatear_etiqueta_escala(d['categoria']) for d in datos_lista]
    valores_m = [d['M'] for d in datos_lista]
    valores_h = [d['H'] for d in datos_lista]

    # Ajustar tamaño según número de categorías
    if orientacion == 'horizontal':
        alto = max(6, len(categorias) * 0.8)
        figsize = (12, alto)
    else:
        ancho = max(10, len(categorias) * 0.8)
        figsize = (ancho, 8)

    fig, ax = plt.subplots(figsize=figsize)
    fig.patch.set_facecolor('white')

    if orientacion == 'horizontal':
        y_pos = np.arange(len(categorias))
        bar_height = 0.35

        # IMPORTANTE: Dibujar TODAS las barras (incluso las de valor 0) para mantener agrupación visual
        # Esto asegura que Mujeres y Hombres aparezcan siempre juntas para cada categoría
        bars_m = ax.barh(y_pos - bar_height/2, valores_m, bar_height,
                         label='Mujeres', color=COLORES['mujer'])
        bars_h = ax.barh(y_pos + bar_height/2, valores_h, bar_height,
                         label='Hombres', color=COLORES['hombre'])

        ax.set_yticks(y_pos)
        ax.set_yticklabels(categorias, fontsize=13)  # Añadido tamaño de fuente
        ax.set_xlabel(tipo_valor, fontsize=16, fontweight='bold')  # Aumentado de 12 a 16
        ax.invert_yaxis()

        # Añadir valores solo para barras con valor > 0
        for i, (bar, valor) in enumerate(zip(bars_m, valores_m)):
            if valor > 0:
                ax.text(valor, bar.get_y() + bar.get_height()/2,
                       f' {formato_numero_es(valor, 2)}€',
                       va='center', ha='left', fontsize=13, fontweight='bold')  # Aumentado de 9 a 13

        for i, (bar, valor) in enumerate(zip(bars_h, valores_h)):
            if valor > 0:
                ax.text(valor, bar.get_y() + bar.get_height()/2,
                       f' {formato_numero_es(valor, 2)}€',
                       va='center', ha='left', fontsize=13, fontweight='bold')  # Aumentado de 9 a 13

    else:  # vertical
        x_pos = np.arange(len(categorias))
        bar_width = 0.35

        bars_m = ax.bar(x_pos - bar_width/2, valores_m, bar_width,
                        label='Mujeres', color=COLORES['mujer'])
        bars_h = ax.bar(x_pos + bar_width/2, valores_h, bar_width,
                        label='Hombres', color=COLORES['hombre'])

        ax.set_xticks(x_pos)
        ax.set_xticklabels(categorias, rotation=45, ha='right', fontsize=13)  # Añadido tamaño de fuente
        ax.set_ylabel(tipo_valor, fontsize=16, fontweight='bold')  # Aumentado de 12 a 16
        if etiqueta_categoria:
            ax.set_xlabel(etiqueta_categoria, fontsize=16, fontweight='bold')  # Etiqueta para eje X

        # Añadir valores dentro de las barras, en vertical y centrados
        for bar in bars_m + bars_h:
            height = bar.get_height()
            if height > 0:
                ax.text(bar.get_x() + bar.get_width()/2, height/2,
                       f'{formato_numero_es(height, 2)}€',
                       ha='center', va='center', fontsize=13, fontweight='bold',
                       rotation=90, color='white')  # Texto vertical, centrado y en blanco

    if mostrar_leyenda:
        ax.legend(loc='best', fontsize=15)  # Aumentado de 11 a 15

    ax.grid(True, axis='x' if orientacion == 'horizontal' else 'y', alpha=0.3)
    ax.set_axisbelow(True)

    if mostrar_titulo:
        ax.set_title(titulo, fontsize=GRAFICO_CFG['titulo_size'],
                     fontweight=GRAFICO_CFG['titulo_weight'],
                     color=GRAFICO_CFG['titulo_color'], pad=20)

    plt.tight_layout()
    plt.savefig(nombre_archivo, dpi=GRAFICO_CFG['dpi'], bbox_inches='tight', facecolor='white')
    plt.close()

    return nombre_archivo


def crear_grafico_barras_acumulativo(datos_lista, titulo, nombre_archivo="temp_barras_acum.png", mostrar_titulo=True, mostrar_leyenda=True, aplicar_filtro_privacidad=True):
    """
    Crea un gráfico de barras horizontales acumulativas.

    Para cada categoría muestra:
    - Primer tramo: Salario Base (sb_M, sb_H)
    - Segundo tramo apilado: Complementos Salariales (diferencia entre sbc y sb)
    - Tercer tramo apilado: Complementos Extrasalariales (diferencia entre sbce y sbc)
    - Etiqueta al final: Total (sbce_M, sbce_H)

    datos_lista: lista de dicts con claves 'categoria', 'sb_M', 'sb_H', 'sbc_M', 'sbc_H', 'sbce_M', 'sbce_H', 'n_M', 'n_H'
    """
    if not datos_lista:
        log(f"Sin datos para gráfico: {titulo}", 'WARN')
        return None

    # Aplicar filtro de privacidad solo si se solicita
    if aplicar_filtro_privacidad:
        datos_lista = aplicar_filtro_privacidad_datos(datos_lista)

    if not datos_lista:
        log(f"Gráfico '{titulo}' omitido - ningún elemento cumple requisitos de privacidad", 'WARN')
        return None

    # Filtrar categorías sin datos
    datos_lista = [d for d in datos_lista if d['sbce_M'] > 0 or d['sbce_H'] > 0]

    if not datos_lista:
        return None

    categorias = [reformatear_etiqueta_escala(d['categoria']) for d in datos_lista]

    # Primer tramo: Salario Base
    sb_m = [d['sb_M'] for d in datos_lista]
    sb_h = [d['sb_H'] for d in datos_lista]

    # Segundo tramo: Complementos Salariales (diferencia entre sbc y sb)
    comp_sal_m = [d['sbc_M'] - d['sb_M'] for d in datos_lista]
    comp_sal_h = [d['sbc_H'] - d['sb_H'] for d in datos_lista]

    # Tercer tramo: Complementos Extrasalariales (diferencia entre sbce y sbc)
    comp_ext_m = [d['sbce_M'] - d['sbc_M'] for d in datos_lista]
    comp_ext_h = [d['sbce_H'] - d['sbc_H'] for d in datos_lista]

    # Total para etiquetas
    total_m = [d['sbce_M'] for d in datos_lista]
    total_h = [d['sbce_H'] for d in datos_lista]

    # Ajustar tamaño según número de categorías
    alto = max(6, len(categorias) * 0.8)
    figsize = (14, alto)

    fig, ax = plt.subplots(figsize=figsize)
    fig.patch.set_facecolor('white')

    y_pos = np.arange(len(categorias))
    bar_height = 0.35

    # Barras de Mujeres (apiladas en 3 tramos)
    bars_m_base = ax.barh(y_pos - bar_height/2, sb_m, bar_height,
                          label='Mujeres', color=COLORES['mujer'])
    bars_m_comp_sal = ax.barh(y_pos - bar_height/2, comp_sal_m, bar_height,
                              left=sb_m, color=COLORES['mujer'], alpha=0.7)
    bars_m_comp_ext = ax.barh(y_pos - bar_height/2, comp_ext_m, bar_height,
                              left=[sb_m[i] + comp_sal_m[i] for i in range(len(sb_m))],
                              color=COLORES['mujer'], alpha=0.5)

    # Barras de Hombres (apiladas en 3 tramos)
    bars_h_base = ax.barh(y_pos + bar_height/2, sb_h, bar_height,
                          label='Hombres', color=COLORES['hombre'])
    bars_h_comp_sal = ax.barh(y_pos + bar_height/2, comp_sal_h, bar_height,
                              left=sb_h, color=COLORES['hombre'], alpha=0.7)
    bars_h_comp_ext = ax.barh(y_pos + bar_height/2, comp_ext_h, bar_height,
                              left=[sb_h[i] + comp_sal_h[i] for i in range(len(sb_h))],
                              color=COLORES['hombre'], alpha=0.5)

    ax.set_yticks(y_pos)
    ax.set_yticklabels(categorias, fontsize=13)  # Añadido tamaño de fuente
    ax.set_xlabel('Retribución (€)', fontsize=16, fontweight='bold')  # Aumentado de 12 a 16
    ax.invert_yaxis()

    # Umbral: mostrar etiqueta interna solo si el tramo representa al menos 8% del total
    umbral_porcentaje = 0.08

    # Añadir etiquetas de valores
    # Etiqueta en el primer tramo (Salario Base) - solo si hay espacio suficiente
    for i, (bar, valor, total) in enumerate(zip(bars_m_base, sb_m, total_m)):
        if valor > 0 and total > 0:
            # Mostrar solo si el tramo es suficientemente ancho (al menos 8% del total)
            if (valor / total) >= umbral_porcentaje:
                ax.text(valor/2, bar.get_y() + bar.get_height()/2,
                       f'{formato_numero_es(valor, 2)}€',
                       va='center', ha='center', fontsize=12, fontweight='bold', color='white')

    for i, (bar, valor, total) in enumerate(zip(bars_h_base, sb_h, total_h)):
        if valor > 0 and total > 0:
            if (valor / total) >= umbral_porcentaje:
                ax.text(valor/2, bar.get_y() + bar.get_height()/2,
                       f'{formato_numero_es(valor, 2)}€',
                       va='center', ha='center', fontsize=12, fontweight='bold', color='white')

    # Etiqueta en el segundo tramo (Complementos Salariales) - solo si hay espacio suficiente
    for i, (valor_base, valor_comp_sal, total) in enumerate(zip(sb_m, comp_sal_m, total_m)):
        if valor_comp_sal > 0 and total > 0:
            if (valor_comp_sal / total) >= umbral_porcentaje:
                ax.text(valor_base + valor_comp_sal/2, y_pos[i] - bar_height/2,
                       f'{formato_numero_es(valor_comp_sal, 2)}€',
                       va='center', ha='center', fontsize=12, fontweight='bold', color='white')

    for i, (valor_base, valor_comp_sal, total) in enumerate(zip(sb_h, comp_sal_h, total_h)):
        if valor_comp_sal > 0 and total > 0:
            if (valor_comp_sal / total) >= umbral_porcentaje:
                ax.text(valor_base + valor_comp_sal/2, y_pos[i] + bar_height/2,
                       f'{formato_numero_es(valor_comp_sal, 2)}€',
                       va='center', ha='center', fontsize=12, fontweight='bold', color='white')

    # Etiqueta en el tercer tramo (Complementos Extrasalariales) - solo si hay espacio suficiente
    for i in range(len(sb_m)):
        valor_comp_ext = comp_ext_m[i]
        total = total_m[i]
        if valor_comp_ext > 0 and total > 0:
            if (valor_comp_ext / total) >= umbral_porcentaje:
                ax.text(sb_m[i] + comp_sal_m[i] + valor_comp_ext/2, y_pos[i] - bar_height/2,
                       f'{formato_numero_es(valor_comp_ext, 2)}€',
                       va='center', ha='center', fontsize=12, fontweight='bold', color='white')

    for i in range(len(sb_h)):
        valor_comp_ext = comp_ext_h[i]
        total = total_h[i]
        if valor_comp_ext > 0 and total > 0:
            if (valor_comp_ext / total) >= umbral_porcentaje:
                ax.text(sb_h[i] + comp_sal_h[i] + valor_comp_ext/2, y_pos[i] + bar_height/2,
                       f'{formato_numero_es(valor_comp_ext, 2)}€',
                       va='center', ha='center', fontsize=12, fontweight='bold', color='white')

    # Etiqueta al final con el total
    for i, (bar, valor_total) in enumerate(zip(bars_m_comp_ext, total_m)):
        if valor_total > 0:
            ax.text(valor_total, bar.get_y() + bar.get_height()/2,
                   f' {formato_numero_es(valor_total, 2)}€',
                   va='center', ha='left', fontsize=13, fontweight='bold')

    for i, (bar, valor_total) in enumerate(zip(bars_h_comp_ext, total_h)):
        if valor_total > 0:
            ax.text(valor_total, bar.get_y() + bar.get_height()/2,
                   f' {formato_numero_es(valor_total, 2)}€',
                   va='center', ha='left', fontsize=13, fontweight='bold')

    if mostrar_leyenda:
        ax.legend(loc='best', fontsize=15)  # Aumentado de 11 a 15

    ax.grid(True, axis='x', alpha=0.3)
    ax.set_axisbelow(True)

    if mostrar_titulo:
        ax.set_title(titulo, fontsize=GRAFICO_CFG['titulo_size'],
                     fontweight=GRAFICO_CFG['titulo_weight'],
                     color=GRAFICO_CFG['titulo_color'], pad=20)

    plt.tight_layout()
    plt.savefig(nombre_archivo, dpi=GRAFICO_CFG['dpi'], bbox_inches='tight', facecolor='white')
    plt.close()

    return nombre_archivo


def crear_grafico_barras_doble_eje_denso(datos_lista, titulo, nombre_archivo="temp_barras_dual.png"):
    """
    Crea un gráfico de barras vertical con doble eje para ALTA DENSIDAD de datos (muchas categorías):
    - Eje izquierdo (Y1): Barras de salarios para Mujeres y Hombres
    - Eje derecho (Y2): Línea de Puntos
    - Fuentes más pequeñas para acomodar muchas categorías

    Usado en: Apartado 3 - RETRIBUCIÓN POR AGRUPACIÓN (NIVEL SVPT) Y PUESTO DE TRABAJO

    datos_lista: lista de dicts con claves 'categoria', 'M', 'H', 'puntos', 'n_M', 'n_H'
    """
    if not datos_lista:
        log(f"Sin datos para gráfico: {titulo}", 'WARN')
        return None

    # Aplicar filtro de privacidad primero
    datos_lista = aplicar_filtro_privacidad_datos(datos_lista)

    if not datos_lista:
        log(f"Gráfico '{titulo}' omitido - ningún elemento cumple requisitos de privacidad", 'WARN')
        return None

    # Filtrar categorías sin datos
    datos_lista = [d for d in datos_lista if d.get('M', 0) > 0 or d.get('H', 0) > 0]

    if not datos_lista:
        return None

    categorias = [reformatear_etiqueta_escala(d['categoria']) for d in datos_lista]
    valores_m = [d.get('M', 0) for d in datos_lista]
    valores_h = [d.get('H', 0) for d in datos_lista]
    puntos = [d.get('puntos', 0) for d in datos_lista]

    # Ajustar tamaño según número de categorías - MUY GRANDE para máxima legibilidad
    ancho = max(35, len(categorias) * 3.0)  # Más ancho por categoría
    altura = max(24, ancho * 0.65)  # Mayor altura proporcional
    fig, ax1 = plt.subplots(figsize=(ancho, altura))
    fig.patch.set_facecolor('white')

    # Eje 1 (izquierdo): Barras de salarios
    x_pos = np.arange(len(categorias))
    bar_width = 0.45

    bars_m = ax1.bar(x_pos - bar_width/2, valores_m, bar_width,
                     label='Mujeres', color=COLORES['mujer'], alpha=0.9, edgecolor='white', linewidth=3)
    bars_h = ax1.bar(x_pos + bar_width/2, valores_h, bar_width,
                     label='Hombres', color=COLORES['hombre'], alpha=0.9, edgecolor='white', linewidth=3)

    ax1.set_xlabel('Puesto de Trabajo', fontsize=16, fontweight='bold', labelpad=15)  # Etiqueta específica para puestos
    ax1.set_ylabel('Salario (€)', fontsize=16, fontweight='bold', color='black', labelpad=15)  # Igual que gráficos normales
    ax1.set_xticks(x_pos)
    ax1.set_xticklabels(categorias, rotation=45, ha='right', fontsize=13, fontweight='normal')  # Igual que gráficos normales
    ax1.tick_params(axis='y', labelcolor='black', labelsize=14, width=3, length=10)  # Igual que gráficos normales
    ax1.tick_params(axis='x', labelsize=13, width=3, length=10)  # Igual que gráficos normales
    ax1.grid(True, axis='y', alpha=0.6, linestyle='--', linewidth=2)
    ax1.set_axisbelow(True)

    # Añadir valores en las barras - DENTRO DE LAS BARRAS, EN VERTICAL Y BLANCO
    for i, bar in enumerate(bars_m):
        height = bar.get_height()
        if height > 0:
            # Valores de mujeres: centrado en la barra, vertical, blanco
            ax1.text(bar.get_x() + bar.get_width()/2, height / 2,
                    f'{formato_numero_es(height, 0)}',
                    ha='center', va='center', fontsize=24, fontweight='bold',  # AUMENTADO para mejor visibilidad
                    color='white', rotation=90)

    for i, bar in enumerate(bars_h):
        height = bar.get_height()
        if height > 0:
            # Valores de hombres: centrado en la barra, vertical, blanco
            ax1.text(bar.get_x() + bar.get_width()/2, height / 2,
                    f'{formato_numero_es(height, 0)}',
                    ha='center', va='center', fontsize=24, fontweight='bold',  # AUMENTADO para mejor visibilidad
                    color='white', rotation=90)

    # Eje 2 (derecho): Línea de puntos - MUY VISIBLE
    ax2 = ax1.twinx()
    line_puntos = ax2.plot(x_pos, puntos, color='#2ca02c', marker='o',
                           linewidth=7, markersize=26, label='Puntos', linestyle='-',
                           markeredgecolor='white', markeredgewidth=5)
    ax2.set_ylabel('Puntos', fontsize=16, fontweight='bold', color='#2ca02c', labelpad=20)  # Igual que gráficos normales
    ax2.tick_params(axis='y', labelcolor='#2ca02c', labelsize=14, width=3.5, length=12)  # Igual que gráficos normales

    # Añadir valores en la línea de puntos - POSICIONADOS DENTRO DEL GRÁFICO
    max_puntos = max(puntos) if puntos else 1
    for i, (x, y) in enumerate(zip(x_pos, puntos)):
        if y > 0:
            # Posicionar el texto DEBAJO del punto si está muy arriba, o arriba si hay espacio
            if y > max_puntos * 0.85:  # Si el punto está en el 85% superior
                ax2.text(x, y - (max_puntos * 0.06), f'{int(y)}', ha='center', va='top',
                        fontsize=26, fontweight='bold', color='white',  # AUMENTADO para mejor visibilidad
                        bbox=dict(boxstyle='round,pad=0.8', facecolor='#2ca02c', edgecolor='white',
                                 alpha=0.95, linewidth=3.5))
            else:
                ax2.text(x, y + (max_puntos * 0.04), f'{int(y)}', ha='center', va='bottom',
                        fontsize=26, fontweight='bold', color='white',  # AUMENTADO para mejor visibilidad
                        bbox=dict(boxstyle='round,pad=0.8', facecolor='#2ca02c', edgecolor='white',
                                 alpha=0.95, linewidth=3.5))

    # Expandir límites del eje Y de puntos para dar espacio arriba
    y_max_puntos = max(puntos) if puntos else 1
    ax2.set_ylim(0, y_max_puntos * 1.15)  # 15% más espacio arriba

    # Título con más padding
    ax1.set_title(titulo, fontsize=GRAFICO_CFG['titulo_size'] + 8,
                 fontweight=GRAFICO_CFG['titulo_weight'],
                 color=GRAFICO_CFG['titulo_color'], pad=50)

    # Leyendas combinadas
    lines1, labels1 = ax1.get_legend_handles_labels()
    lines2, labels2 = ax2.get_legend_handles_labels()
    legend = ax1.legend(lines1 + lines2, labels1 + labels2,
                       loc='upper right', fontsize=15, frameon=True, shadow=True,  # Igual que gráficos normales
                       fancybox=True, framealpha=0.98, edgecolor='black')
    legend.get_frame().set_linewidth(3.5)

    plt.tight_layout()
    plt.savefig(nombre_archivo, dpi=GRAFICO_CFG['dpi'], bbox_inches='tight', facecolor='white')
    plt.close()

    return nombre_archivo


def crear_grafico_barras_doble_eje_simple(datos_lista, titulo, nombre_archivo="temp_barras_dual.png"):
    """
    Crea un gráfico de barras vertical con doble eje para BAJA DENSIDAD de datos (pocas categorías):
    - Eje izquierdo (Y1): Barras de salarios para Mujeres y Hombres
    - Eje derecho (Y2): Línea de Puntos
    - Fuentes MUY GRANDES para máxima legibilidad con pocas categorías

    Usado en: Apartado 4 - ANÁLISIS POR ESCALA SVPT (E1-E5)

    datos_lista: lista de dicts con claves 'categoria', 'M', 'H', 'puntos', 'n_M', 'n_H'
    """
    if not datos_lista:
        log(f"Sin datos para gráfico: {titulo}", 'WARN')
        return None

    # Aplicar filtro de privacidad primero
    datos_lista = aplicar_filtro_privacidad_datos(datos_lista)

    if not datos_lista:
        log(f"Gráfico '{titulo}' omitido - ningún elemento cumple requisitos de privacidad", 'WARN')
        return None

    # Filtrar categorías sin datos
    datos_lista = [d for d in datos_lista if d.get('M', 0) > 0 or d.get('H', 0) > 0]

    if not datos_lista:
        return None

    categorias = [reformatear_etiqueta_escala(d['categoria']) for d in datos_lista]
    valores_m = [d.get('M', 0) for d in datos_lista]
    valores_h = [d.get('H', 0) for d in datos_lista]
    puntos = [d.get('puntos', 0) for d in datos_lista]

    # Ajustar tamaño según número de categorías - MUY GRANDE para máxima legibilidad
    ancho = max(35, len(categorias) * 3.0)  # Más ancho por categoría
    altura = max(24, ancho * 0.65)  # Mayor altura proporcional
    fig, ax1 = plt.subplots(figsize=(ancho, altura))
    fig.patch.set_facecolor('white')

    # Eje 1 (izquierdo): Barras de salarios
    x_pos = np.arange(len(categorias))
    bar_width = 0.45

    bars_m = ax1.bar(x_pos - bar_width/2, valores_m, bar_width,
                     label='Mujeres', color=COLORES['mujer'], alpha=0.9, edgecolor='white', linewidth=3)
    bars_h = ax1.bar(x_pos + bar_width/2, valores_h, bar_width,
                     label='Hombres', color=COLORES['hombre'], alpha=0.9, edgecolor='white', linewidth=3)

    ax1.set_xlabel('Puesto de Trabajo', fontsize=48, fontweight='bold', labelpad=30)  # Etiqueta específica para puestos dentro de escalas
    ax1.set_ylabel('Salario (€)', fontsize=52, fontweight='bold', color='black', labelpad=30)  # MUY GRANDE
    ax1.set_xticks(x_pos)
    ax1.set_xticklabels(categorias, rotation=45, ha='right', fontsize=36, fontweight='normal')  # MUY GRANDE
    ax1.tick_params(axis='y', labelcolor='black', labelsize=40, width=3, length=10)  # MUY GRANDE
    ax1.tick_params(axis='x', labelsize=36, width=3, length=10)  # MUY GRANDE
    ax1.grid(True, axis='y', alpha=0.6, linestyle='--', linewidth=2)
    ax1.set_axisbelow(True)

    # Añadir valores en las barras - DENTRO DE LAS BARRAS, EN VERTICAL Y BLANCO
    for i, bar in enumerate(bars_m):
        height = bar.get_height()
        if height > 0:
            # Valores de mujeres: centrado en la barra, vertical, blanco
            ax1.text(bar.get_x() + bar.get_width()/2, height / 2,
                    f'{formato_numero_es(height, 0)}',
                    ha='center', va='center', fontsize=38, fontweight='bold',  # MUY GRANDE
                    color='white', rotation=90)

    for i, bar in enumerate(bars_h):
        height = bar.get_height()
        if height > 0:
            # Valores de hombres: centrado en la barra, vertical, blanco
            ax1.text(bar.get_x() + bar.get_width()/2, height / 2,
                    f'{formato_numero_es(height, 0)}',
                    ha='center', va='center', fontsize=38, fontweight='bold',  # MUY GRANDE
                    color='white', rotation=90)

    # Eje 2 (derecho): Línea de puntos - MUY VISIBLE
    ax2 = ax1.twinx()
    line_puntos = ax2.plot(x_pos, puntos, color='#2ca02c', marker='o',
                           linewidth=9, markersize=32, label='Puntos', linestyle='-',
                           markeredgecolor='white', markeredgewidth=6)
    ax2.set_ylabel('Puntos', fontsize=54, fontweight='bold', color='#2ca02c', labelpad=35)  # MUY GRANDE
    ax2.tick_params(axis='y', labelcolor='#2ca02c', labelsize=42, width=4, length=14)  # MUY GRANDE

    # Añadir valores en la línea de puntos - POSICIONADOS DENTRO DEL GRÁFICO
    max_puntos = max(puntos) if puntos else 1
    for i, (x, y) in enumerate(zip(x_pos, puntos)):
        if y > 0:
            # Posicionar el texto DEBAJO del punto si está muy arriba, o arriba si hay espacio
            if y > max_puntos * 0.85:  # Si el punto está en el 85% superior
                ax2.text(x, y - (max_puntos * 0.06), f'{int(y)}', ha='center', va='top',
                        fontsize=38, fontweight='bold', color='white',  # MUY GRANDE
                        bbox=dict(boxstyle='round,pad=1.0', facecolor='#2ca02c', edgecolor='white',
                                 alpha=0.95, linewidth=4))
            else:
                ax2.text(x, y + (max_puntos * 0.04), f'{int(y)}', ha='center', va='bottom',
                        fontsize=38, fontweight='bold', color='white',  # MUY GRANDE
                        bbox=dict(boxstyle='round,pad=1.0', facecolor='#2ca02c', edgecolor='white',
                                 alpha=0.95, linewidth=4))

    # Expandir límites del eje Y de puntos para dar espacio arriba
    y_max_puntos = max(puntos) if puntos else 1
    ax2.set_ylim(0, y_max_puntos * 1.15)  # 15% más espacio arriba

    # Título con más padding
    ax1.set_title(titulo, fontsize=GRAFICO_CFG['titulo_size'] + 14,  # Título aún más grande
                 fontweight=GRAFICO_CFG['titulo_weight'],
                 color=GRAFICO_CFG['titulo_color'], pad=60)

    # Leyendas combinadas - MUY GRANDE
    lines1, labels1 = ax1.get_legend_handles_labels()
    lines2, labels2 = ax2.get_legend_handles_labels()
    legend = ax1.legend(lines1 + lines2, labels1 + labels2,
                       loc='upper right', fontsize=38, frameon=True, shadow=True,  # MUY GRANDE
                       fancybox=True, framealpha=0.98, edgecolor='black')
    legend.get_frame().set_linewidth(4)

    plt.tight_layout()
    plt.savefig(nombre_archivo, dpi=GRAFICO_CFG['dpi'], bbox_inches='tight', facecolor='white')
    plt.close()

    return nombre_archivo


def tiene_datos_puntos(datos_lista):
    """
    Verifica si los datos contienen información válida de Puntos.
    Retorna True si hay al menos un valor de 'puntos' > 0.
    """
    if not datos_lista:
        return False

    for dato in datos_lista:
        puntos = dato.get('puntos', 0)
        if puntos and puntos > 0:
            return True

    return False


def crear_grafico_salarios_con_puntos_condicional(datos_lista, titulo, nombre_archivo="temp_barras.png", tipo_grafico='denso'):
    """
    Crea un gráfico de salarios que puede incluir o no la línea de Puntos:
    - Si hay datos de Puntos válidos (> 0): crea gráfico de doble eje
    - Si NO hay datos de Puntos: crea gráfico de barras vertical normal

    datos_lista: lista de dicts con claves 'categoria', 'M', 'H', 'puntos' (opcional), 'n_M', 'n_H'
    tipo_grafico: 'denso' (muchas categorías, fuentes más pequeñas) o 'simple' (pocas categorías, fuentes grandes)
    """
    if not datos_lista:
        log(f"Sin datos para gráfico: {titulo}", 'WARN')
        return None

    # Aplicar filtro de privacidad primero
    datos_lista = aplicar_filtro_privacidad_datos(datos_lista)

    if not datos_lista:
        log(f"Gráfico '{titulo}' omitido - ningún elemento cumple requisitos de privacidad", 'WARN')
        return None

    # Verificar si hay datos válidos de Puntos
    if tiene_datos_puntos(datos_lista):
        # Usar gráfico de doble eje (salarios + puntos) según el tipo
        log(f"Creando gráfico con doble eje (incluye Puntos) tipo '{tipo_grafico}': {titulo}", 'DEBUG')
        if tipo_grafico == 'simple':
            return crear_grafico_barras_doble_eje_simple(datos_lista, titulo, nombre_archivo)
        else:
            return crear_grafico_barras_doble_eje_denso(datos_lista, titulo, nombre_archivo)
    else:
        # Usar gráfico de barras normal (solo salarios)
        log(f"Creando gráfico de barras normal (sin Puntos): {titulo}", 'DEBUG')
        return crear_grafico_barras(
            datos_lista,
            titulo,
            tipo_valor="Salario (€)",
            nombre_archivo=nombre_archivo,
            orientacion='vertical',
            etiqueta_categoria='Puesto de Trabajo',
            mostrar_titulo=True,
            mostrar_leyenda=True
        )


# ==================== GENERACIÓN DE TABLAS WORD ====================

def colorear_texto_celda(cell, color_rgb):
    """Aplica color al texto de una celda (no al fondo)"""
    if cell.paragraphs:
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = color_rgb


def aplicar_formato_celda(cell, texto, alineacion=WD_ALIGN_PARAGRAPH.CENTER, bold=False, font_size=7, color_rgb=None):
    """Aplica formato unificado a una celda de tabla con estilo moderno y minimalista"""
    cell.text = texto
    paragraph = cell.paragraphs[0]
    paragraph.alignment = alineacion

    if paragraph.runs:
        run = paragraph.runs[0]
        run.font.name = 'Calibri'
        run.font.size = Pt(font_size)
        run.font.bold = bold
        if color_rgb:
            run.font.color.rgb = color_rgb


def colorear_celda_fondo(cell, color_rgb):
    """Aplica color de fondo a una celda (usado solo para encabezados)"""
    # Si es RGBColor, obtener los valores
    if isinstance(color_rgb, RGBColor):
        color_hex = '%02x%02x%02x' % (color_rgb[0], color_rgb[1], color_rgb[2])
    elif isinstance(color_rgb, tuple):
        color_hex = '%02x%02x%02x' % color_rgb
    else:
        color_hex = 'ffffff'  # blanco por defecto

    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._element.get_or_add_tcPr().append(shading_elm)


def aplicar_bordes_tabla(table):
    """
    Aplica bordes finos y grises a todas las celdas de una tabla para un estilo moderno y minimalista.
    Bordes color gris claro (C0C0C0) con grosor fino (2 = 0.25pt).
    """
    from docx.oxml import parse_xml

    # Definir bordes finos y grises para toda la tabla (sz="2" = muy fino, color gris claro)
    borders_xml = '''
    <w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
        <w:top w:val="single" w:sz="2" w:space="0" w:color="C0C0C0"/>
        <w:left w:val="single" w:sz="2" w:space="0" w:color="C0C0C0"/>
        <w:bottom w:val="single" w:sz="2" w:space="0" w:color="C0C0C0"/>
        <w:right w:val="single" w:sz="2" w:space="0" w:color="C0C0C0"/>
        <w:insideH w:val="single" w:sz="2" w:space="0" w:color="C0C0C0"/>
        <w:insideV w:val="single" w:sz="2" w:space="0" w:color="C0C0C0"/>
    </w:tblBorders>
    '''

    # Obtener el elemento de propiedades de la tabla
    tbl_pr = table._element.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        table._element.insert(0, tbl_pr)

    # Eliminar bordes existentes si los hay
    for child in list(tbl_pr):
        if child.tag.endswith('tblBorders'):
            tbl_pr.remove(child)

    # Agregar los nuevos bordes
    tbl_pr.append(parse_xml(borders_xml))


def crear_tabla_resumen(doc, titulo, datos, metodo='PROMEDIO'):
    """
    Crea tabla de resumen con valores efectivos y equiparados.
    Formato limpio con colores solo en texto según lógica de brecha.

    datos estructura:
    {
        'efectivo': {'M': val, 'H': val, 'brecha': val, 'n_M': count, 'n_H': count},
        'equiparado': {'M': val, 'H': val, 'brecha': val, 'n_M': count, 'n_H': count}
    }
    """
    # Verificar privacidad - no mostrar si hay empleados identificables
    if verificar_privacidad_tabla(datos) is None:
        log(f"Tabla '{titulo}' omitida por privacidad (empleado identificable)", 'WARN')
        return

    doc.add_heading(titulo, level=3)

    # Extraer el código del concepto del título (ej: "(SB)", "(SB+C)", "(SB+C+ES)")
    # Buscar el texto entre paréntesis al final del título
    import re
    match = re.search(r'\(([^)]+)\)$', titulo)
    concepto = match.group(1) if match else 'SB'

    # Crear tabla con 8 columnas (según imagen)
    table = doc.add_table(rows=2, cols=8)
    aplicar_bordes_tabla(table)

    # Fila de encabezados
    row_headers = table.rows[0]
    headers = [
        'Nº M',
        'Nº H',
        f'Promedio ({concepto}) Efectivo [Mujeres]',
        f'Promedio ({concepto}) Efectivo [Hombres]',
        f'Brecha ({concepto}) Efectivo',
        f'Promedio ({concepto}) Equiparado [Mujeres]',
        f'Promedio ({concepto}) Equiparado [Hombres]',
        f'Brecha ({concepto}) Equiparado'
    ]

    for i, header in enumerate(headers):
        cell = row_headers.cells[i]
        aplicar_formato_celda(cell, header, bold=True, font_size=7)
        colorear_celda_fondo(cell, COLOR_FONDO_ENCABEZADO)

    # Fila de datos
    efe = datos['efectivo']
    equi = datos['equiparado']
    row_data = table.rows[1]

    # Valores con símbolos
    valores = [
        str(efe.get('n_M', 0)),  # N Mujeres
        str(efe.get('n_H', 0)),  # N Hombres
        formato_numero_es(efe['M'], 2) + ' €',  # Promedio Efectivo Mujeres
        formato_numero_es(efe['H'], 2) + ' €',  # Promedio Efectivo Hombres
        formato_numero_es(abs(efe['brecha']), 2) + ' %',  # Brecha Efectivo
        formato_numero_es(equi['M'], 2) + ' €',  # Promedio Equiparado Mujeres
        formato_numero_es(equi['H'], 2) + ' €',  # Promedio Equiparado Hombres
        formato_numero_es(abs(equi['brecha']), 2) + ' %'  # Brecha Equiparado
    ]

    # Rellenar celdas con formato unificado
    for i, valor in enumerate(valores):
        aplicar_formato_celda(row_data.cells[i], valor, font_size=7)

    # Aplicar colores según lógica de brecha:
    # Si H > M (brecha positiva, favorable a hombres):
    #   - Valor de M en naranja (desfavorable)
    #   - Valor de H en azul (favorable)
    #   - Brecha en naranja (indica desventaja para mujeres)
    # Si M > H (brecha negativa, favorable a mujeres):
    #   - Valor de M en azul (favorable)
    #   - Valor de H en naranja (desfavorable)
    #   - Brecha en azul (indica ventaja para mujeres)

    # Para efectivo
    if efe['H'] > efe['M']:
        # Favorable a hombres
        colorear_texto_celda(row_data.cells[2], COLORES_RGB['hombre'])  # M en naranja
        colorear_texto_celda(row_data.cells[3], COLORES_RGB['mujer'])   # H en azul
        colorear_texto_celda(row_data.cells[4], COLORES_RGB['hombre'])  # Brecha en naranja
    elif efe['M'] > efe['H']:
        # Favorable a mujeres
        colorear_texto_celda(row_data.cells[2], COLORES_RGB['mujer'])   # M en azul
        colorear_texto_celda(row_data.cells[3], COLORES_RGB['hombre'])  # H en naranja
        colorear_texto_celda(row_data.cells[4], COLORES_RGB['mujer'])   # Brecha en azul
    # Si son iguales, no se aplica color (queda en negro)

    # Para equiparado
    if equi['H'] > equi['M']:
        # Favorable a hombres
        colorear_texto_celda(row_data.cells[5], COLORES_RGB['hombre'])  # M en naranja
        colorear_texto_celda(row_data.cells[6], COLORES_RGB['mujer'])   # H en azul
        colorear_texto_celda(row_data.cells[7], COLORES_RGB['hombre'])  # Brecha en naranja
    elif equi['M'] > equi['H']:
        # Favorable a mujeres
        colorear_texto_celda(row_data.cells[5], COLORES_RGB['mujer'])   # M en azul
        colorear_texto_celda(row_data.cells[6], COLORES_RGB['hombre'])  # H en naranja
        colorear_texto_celda(row_data.cells[7], COLORES_RGB['mujer'])   # Brecha en azul
    # Si son iguales, no se aplica color (queda en negro)

    doc.add_paragraph()


def crear_tabla_complementos(doc, titulo, datos_lista, aplicar_privacidad=False):
    """
    Crea tabla de complementos con distribución por género.

    datos_lista: lista de dicts con estructura:
    {
        'complemento': nombre (ej: 'PS1'),
        'n_M': count,
        'n_H': count,
        'total': count,
        'pct_M': porcentaje,
        'pct_H': porcentaje
    }
    aplicar_privacidad: si True, aplica filtro de privacidad (por defecto False para distribuciones generales)
    """
    # Aplicar filtro de privacidad solo si se solicita
    if aplicar_privacidad:
        datos_lista = aplicar_filtro_privacidad_datos(datos_lista)

    if not datos_lista:
        log(f"Tabla '{titulo}' omitida - sin datos válidos", 'WARN')
        return

    doc.add_heading(titulo, level=3)

    # Tabla: | Complemento | Nº M | Nº H | Total | % M | % H |
    table = doc.add_table(rows=len(datos_lista) + 1, cols=6)
    aplicar_bordes_tabla(table)

    # Encabezados
    headers = ['Complemento', 'Nº M', 'Nº H', 'Total', '% M', '% H']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        aplicar_formato_celda(cell, header, bold=True, font_size=7)
        colorear_celda_fondo(cell, COLOR_FONDO_ENCABEZADO)

    # Filas de datos
    for idx, datos in enumerate(datos_lista, start=1):
        row = table.rows[idx]

        pct_m = datos.get('pct_M', 0)
        pct_h = datos.get('pct_H', 0)

        valores = [
            str(datos['complemento']),
            str(datos.get('n_M', 0)),
            str(datos.get('n_H', 0)),
            str(datos.get('total', 0)),
            formato_numero_es(pct_m, 1) + ' %' if pct_m > 0 else '0,0 %',
            formato_numero_es(pct_h, 1) + ' %' if pct_h > 0 else '0,0 %'
        ]

        # Rellenar celdas con formato unificado
        for i, text in enumerate(valores):
            aplicar_formato_celda(row.cells[i], text, font_size=7)

        # Aplicar colores en porcentajes: el mayor en azul, el menor en naranja
        if pct_h > pct_m:
            # Hombres tienen mayor porcentaje
            colorear_texto_celda(row.cells[4], COLORES_RGB['hombre'])  # % M en naranja
            colorear_texto_celda(row.cells[5], COLORES_RGB['mujer'])   # % H en azul
        elif pct_m > pct_h:
            # Mujeres tienen mayor porcentaje
            colorear_texto_celda(row.cells[4], COLORES_RGB['mujer'])   # % M en azul
            colorear_texto_celda(row.cells[5], COLORES_RGB['hombre'])  # % H en naranja

    doc.add_paragraph()


def crear_tabla_complementos_por_grupo(doc, titulo, datos_lista):
    """
    Crea tabla de complementos por grupo profesional con promedio y mediana.

    datos_lista: lista de dicts con estructura:
    {
        'grupo': nombre del grupo,
        'n_M': count,
        'n_H': count,
        'promedio_M': valor,
        'promedio_H': valor,
        'promedio_brecha': valor,
        'mediana_M': valor,
        'mediana_H': valor,
        'mediana_brecha': valor
    }
    """
    # Aplicar filtro de privacidad
    datos_lista = aplicar_filtro_privacidad_datos(datos_lista)

    if not datos_lista:
        log(f"Tabla '{titulo}' omitida - ningún elemento cumple requisitos de privacidad", 'WARN')
        return

    doc.add_heading(titulo, level=3)

    # Tabla: | Grupo | Nº M | Nº H | Promedio(M) | Promedio(H) | Promedio(Brecha) | Mediana(M) | Mediana(H) | Mediana(Brecha) |
    table = doc.add_table(rows=len(datos_lista) + 2, cols=9)  # +2 para header y total
    aplicar_bordes_tabla(table)

    # Encabezados
    headers = ['Grupos profesionales', 'Nº M', 'Nº H',
               'Promedio (Mujeres)', 'Promedio (Hombres)', 'Promedio (Brecha)',
               'Mediana (Mujeres)', 'Mediana (Hombres)', 'Mediana (Brecha)']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        aplicar_formato_celda(cell, header, bold=True, font_size=7)
        colorear_celda_fondo(cell, COLOR_FONDO_ENCABEZADO)

    # Calcular totales
    total_n_m = sum(d.get('n_M', 0) for d in datos_lista)
    total_n_h = sum(d.get('n_H', 0) for d in datos_lista)

    # Promedios ponderados para el total
    sum_promedio_m = sum(d.get('promedio_M', 0) * d.get('n_M', 0) for d in datos_lista)
    sum_promedio_h = sum(d.get('promedio_H', 0) * d.get('n_H', 0) for d in datos_lista)
    total_promedio_m = (sum_promedio_m / total_n_m) if total_n_m > 0 else 0
    total_promedio_h = (sum_promedio_h / total_n_h) if total_n_h > 0 else 0
    total_promedio_brecha = calcular_brecha(total_promedio_h, total_promedio_m) or 0

    sum_mediana_m = sum(d.get('mediana_M', 0) * d.get('n_M', 0) for d in datos_lista)
    sum_mediana_h = sum(d.get('mediana_H', 0) * d.get('n_H', 0) for d in datos_lista)
    total_mediana_m = (sum_mediana_m / total_n_m) if total_n_m > 0 else 0
    total_mediana_h = (sum_mediana_h / total_n_h) if total_n_h > 0 else 0
    total_mediana_brecha = calcular_brecha(total_mediana_h, total_mediana_m) or 0

    # Filas de datos
    for idx, datos in enumerate(datos_lista, start=1):
        row = table.rows[idx]

        valores = [
            reformatear_etiqueta_escala(str(datos['grupo'])),
            str(datos.get('n_M', 0)),
            str(datos.get('n_H', 0)),
            formato_numero_es(datos.get('promedio_M', 0), 2) + ' €' if datos.get('promedio_M', 0) > 0 else '-',
            formato_numero_es(datos.get('promedio_H', 0), 2) + ' €' if datos.get('promedio_H', 0) > 0 else '-',
            formato_numero_es(abs(datos.get('promedio_brecha', 0)), 2) + ' %' if datos.get('promedio_M', 0) > 0 and datos.get('promedio_H', 0) > 0 else '-',
            formato_numero_es(datos.get('mediana_M', 0), 2) + ' €' if datos.get('mediana_M', 0) > 0 else '-',
            formato_numero_es(datos.get('mediana_H', 0), 2) + ' €' if datos.get('mediana_H', 0) > 0 else '-',
            formato_numero_es(abs(datos.get('mediana_brecha', 0)), 2) + ' %' if datos.get('mediana_M', 0) > 0 and datos.get('mediana_H', 0) > 0 else '-'
        ]

        # Rellenar celdas
        for i, text in enumerate(valores):
            aplicar_formato_celda(row.cells[i], text, font_size=7)

        # Aplicar colores - Promedio
        promedio_h = datos.get('promedio_H', 0)
        promedio_m = datos.get('promedio_M', 0)
        if promedio_m > 0 and promedio_h > 0:
            if promedio_h > promedio_m:
                colorear_texto_celda(row.cells[3], COLORES_RGB['hombre'])  # M en naranja
                colorear_texto_celda(row.cells[4], COLORES_RGB['mujer'])   # H en azul
                colorear_texto_celda(row.cells[5], COLORES_RGB['hombre'])  # Brecha en naranja
            elif promedio_m > promedio_h:
                colorear_texto_celda(row.cells[3], COLORES_RGB['mujer'])   # M en azul
                colorear_texto_celda(row.cells[4], COLORES_RGB['hombre'])  # H en naranja
                colorear_texto_celda(row.cells[5], COLORES_RGB['mujer'])   # Brecha en azul
            # Si son iguales, no se aplica color (queda en negro)

        # Aplicar colores - Mediana
        mediana_h = datos.get('mediana_H', 0)
        mediana_m = datos.get('mediana_M', 0)
        if mediana_m > 0 and mediana_h > 0:
            if mediana_h > mediana_m:
                colorear_texto_celda(row.cells[6], COLORES_RGB['hombre'])
                colorear_texto_celda(row.cells[7], COLORES_RGB['mujer'])
                colorear_texto_celda(row.cells[8], COLORES_RGB['hombre'])
            elif mediana_m > mediana_h:
                colorear_texto_celda(row.cells[6], COLORES_RGB['mujer'])
                colorear_texto_celda(row.cells[7], COLORES_RGB['hombre'])
                colorear_texto_celda(row.cells[8], COLORES_RGB['mujer'])
            # Si son iguales, no se aplica color (queda en negro)

    # Fila de totales
    row_total = table.rows[-1]
    valores_totales = [
        'Total',
        str(total_n_m),
        str(total_n_h),
        formato_numero_es(total_promedio_m, 2) + ' €',
        formato_numero_es(total_promedio_h, 2) + ' €',
        formato_numero_es(abs(total_promedio_brecha), 2) + ' %',
        formato_numero_es(total_mediana_m, 2) + ' €',
        formato_numero_es(total_mediana_h, 2) + ' €',
        formato_numero_es(abs(total_mediana_brecha), 2) + ' %'
    ]

    for i, text in enumerate(valores_totales):
        aplicar_formato_celda(row_total.cells[i], text, font_size=7, bold=True)
        colorear_celda_fondo(row_total.cells[i], COLOR_FONDO_TOTAL)

    # Colores para totales - Promedio
    if total_promedio_h > total_promedio_m:
        colorear_texto_celda(row_total.cells[3], COLORES_RGB['hombre'])
        colorear_texto_celda(row_total.cells[4], COLORES_RGB['mujer'])
        colorear_texto_celda(row_total.cells[5], COLORES_RGB['hombre'])
    else:
        colorear_texto_celda(row_total.cells[3], COLORES_RGB['mujer'])
        colorear_texto_celda(row_total.cells[4], COLORES_RGB['hombre'])
        colorear_texto_celda(row_total.cells[5], COLORES_RGB['mujer'])

    # Colores para totales - Mediana
    if total_mediana_h > total_mediana_m:
        colorear_texto_celda(row_total.cells[6], COLORES_RGB['hombre'])
        colorear_texto_celda(row_total.cells[7], COLORES_RGB['mujer'])
        colorear_texto_celda(row_total.cells[8], COLORES_RGB['hombre'])
    else:
        colorear_texto_celda(row_total.cells[6], COLORES_RGB['mujer'])
        colorear_texto_celda(row_total.cells[7], COLORES_RGB['hombre'])
        colorear_texto_celda(row_total.cells[8], COLORES_RGB['mujer'])

    doc.add_paragraph()


def crear_tabla_complementos_por_puesto(doc, titulo, datos_lista):
    """
    Crea tabla de complementos por puesto de trabajo con distribución por género.

    datos_lista: lista de dicts con estructura:
    {
        'puesto': nombre del puesto (ej: 'Factory Manager - E1'),
        'complemento': nombre (ej: 'PS1'),
        'n_M': count,
        'n_H': count,
        'total': count,
        'pct_M': porcentaje,
        'pct_H': porcentaje
    }
    """
    # Aplicar filtro de privacidad
    datos_lista = aplicar_filtro_privacidad_datos(datos_lista)

    if not datos_lista:
        log(f"Tabla '{titulo}' omitida - ningún elemento cumple requisitos de privacidad", 'WARN')
        return

    doc.add_heading(titulo, level=3)

    # Tabla: | Puesto de trabajo | Complemento | Nº M | Nº H | Total | % M | % H |
    table = doc.add_table(rows=len(datos_lista) + 1, cols=7)
    aplicar_bordes_tabla(table)

    # Encabezados
    headers = ['Puesto de trabajo', 'Complemento', 'Nº M', 'Nº H', 'Total', '% M', '% H']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        aplicar_formato_celda(cell, header, bold=True, font_size=7)
        colorear_celda_fondo(cell, COLOR_FONDO_ENCABEZADO)

    # Filas de datos
    for idx, datos in enumerate(datos_lista, start=1):
        row = table.rows[idx]

        pct_m = datos.get('pct_M', 0)
        pct_h = datos.get('pct_H', 0)

        valores = [
            reformatear_etiqueta_escala(str(datos['puesto'])),
            str(datos['complemento']),
            str(datos.get('n_M', 0)),
            str(datos.get('n_H', 0)),
            str(datos.get('total', 0)),
            formato_numero_es(pct_m, 1) + ' %' if pct_m > 0 else '0,0 %',
            formato_numero_es(pct_h, 1) + ' %' if pct_h > 0 else '0,0 %'
        ]

        # Rellenar celdas con formato unificado
        for i, text in enumerate(valores):
            aplicar_formato_celda(row.cells[i], text, font_size=7)

        # Aplicar colores en porcentajes: el mayor en azul, el menor en naranja
        if pct_h > pct_m:
            # Hombres tienen mayor porcentaje
            colorear_texto_celda(row.cells[5], COLORES_RGB['hombre'])  # % M en naranja
            colorear_texto_celda(row.cells[6], COLORES_RGB['mujer'])   # % H en azul
        elif pct_m > pct_h:
            # Mujeres tienen mayor porcentaje
            colorear_texto_celda(row.cells[5], COLORES_RGB['mujer'])   # % M en azul
            colorear_texto_celda(row.cells[6], COLORES_RGB['hombre'])  # % H en naranja

    doc.add_paragraph()


def crear_tabla_por_grupo(doc, titulo, datos_lista, metodo='PROMEDIO'):
    """
    Crea tabla con desglose por grupos con formato limpio y colores en texto.

    datos_lista: lista de dicts con estructura:
    {
        'grupo': nombre,
        'n_M': count, 'n_H': count,
        'sb_M': val, 'sb_H': val, 'sb_brecha': val,
        'sbc_M': val, 'sbc_H': val, 'sbc_brecha': val,
        'sbce_M': val, 'sbce_H': val, 'sbce_brecha': val
    }
    """
    # Aplicar filtro de privacidad
    datos_lista = aplicar_filtro_privacidad_datos(datos_lista)

    if not datos_lista:
        log(f"Tabla '{titulo}' omitida - ningún elemento cumple requisitos de privacidad", 'WARN')
        return

    doc.add_heading(titulo, level=3)

    # Tabla: | Grupo | N M | N H | SB M | SB H | Brecha | SBC M | SBC H | Brecha | SBCE M | SBCE H | Brecha |
    table = doc.add_table(rows=len(datos_lista) + 1, cols=12)
    aplicar_bordes_tabla(table)

    # Encabezados
    headers = [
        'Categoría',
        'N º M',
        'N º H',
        f'Retribución {metodo} SIN Complementos (Mujeres)',
        f'Retribución {metodo} SIN Complementos (Hombres)',
        'Brecha Salarial SIN Complementos',
        f'Retribución {metodo} CON Complementos (Mujeres)',
        f'Retribución {metodo} CON Complementos (Hombres)',
        'Brecha Salarial CON Complementos',
        f'Retribución {metodo} CON Complementos ES (Mujeres)',
        f'Retribución {metodo} CON Complementos ES (Hombres)',
        'Brecha Salarial CON Complementos ES'
    ]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        aplicar_formato_celda(cell, header, bold=True, font_size=7)
        colorear_celda_fondo(cell, COLOR_FONDO_ENCABEZADO)

    # Calcular totales
    totales = {
        'n_M': sum(d.get('n_M', 0) for d in datos_lista),
        'n_H': sum(d.get('n_H', 0) for d in datos_lista),
        'sb_M': sum(d.get('sb_M', 0) for d in datos_lista) / len(datos_lista) if datos_lista else 0,
        'sb_H': sum(d.get('sb_H', 0) for d in datos_lista) / len(datos_lista) if datos_lista else 0,
        'sbc_M': sum(d.get('sbc_M', 0) for d in datos_lista) / len(datos_lista) if datos_lista else 0,
        'sbc_H': sum(d.get('sbc_H', 0) for d in datos_lista) / len(datos_lista) if datos_lista else 0,
        'sbce_M': sum(d.get('sbce_M', 0) for d in datos_lista) / len(datos_lista) if datos_lista else 0,
        'sbce_H': sum(d.get('sbce_H', 0) for d in datos_lista) / len(datos_lista) if datos_lista else 0,
    }

    # Filas de datos
    for idx, datos in enumerate(datos_lista, start=1):
        row = table.rows[idx]

        # Preparar valores con símbolos € y %, mostrar "-" si valor es 0
        sb_m = datos.get('sb_M', 0)
        sb_h = datos.get('sb_H', 0)
        sb_brecha = datos.get('sb_brecha', 0)
        sbc_m = datos.get('sbc_M', 0)
        sbc_h = datos.get('sbc_H', 0)
        sbc_brecha = datos.get('sbc_brecha', 0)
        sbce_m = datos.get('sbce_M', 0)
        sbce_h = datos.get('sbce_H', 0)
        sbce_brecha = datos.get('sbce_brecha', 0)

        valores = [
            reformatear_etiqueta_escala(str(datos['grupo'])),
            str(datos.get('n_M', 0)) if datos.get('n_M', 0) > 0 else '-',
            str(datos.get('n_H', 0)) if datos.get('n_H', 0) > 0 else '-',
            formato_numero_es(sb_m, 2) + ' €' if sb_m > 0 else '-',
            formato_numero_es(sb_h, 2) + ' €' if sb_h > 0 else '-',
            formato_numero_es(abs(sb_brecha), 2) + ' %' if (sb_m > 0 and sb_h > 0) else '-',
            formato_numero_es(sbc_m, 2) + ' €' if sbc_m != 0 else '-',
            formato_numero_es(sbc_h, 2) + ' €' if sbc_h != 0 else '-',
            formato_numero_es(abs(sbc_brecha), 2) + ' %' if (sbc_m != 0 and sbc_h != 0) else '-',
            formato_numero_es(sbce_m, 2) + ' €' if sbce_m != 0 else '-',
            formato_numero_es(sbce_h, 2) + ' €' if sbce_h != 0 else '-',
            formato_numero_es(abs(sbce_brecha), 2) + ' %' if (sbce_m != 0 and sbce_h != 0) else '-'
        ]

        # Rellenar celdas con formato unificado
        for i, text in enumerate(valores):
            aplicar_formato_celda(row.cells[i], text, font_size=7)

        # Aplicar colores en texto según lógica de brecha (solo si ambos valores existen)
        # SB (columnas 3,4,5)
        if sb_m > 0 and sb_h > 0:
            if sb_h > sb_m:
                colorear_texto_celda(row.cells[3], COLORES_RGB['hombre'])  # M naranja
                colorear_texto_celda(row.cells[4], COLORES_RGB['mujer'])   # H azul
                colorear_texto_celda(row.cells[5], COLORES_RGB['hombre'])  # Brecha naranja
            elif sb_m > sb_h:
                colorear_texto_celda(row.cells[3], COLORES_RGB['mujer'])   # M azul
                colorear_texto_celda(row.cells[4], COLORES_RGB['hombre'])  # H naranja
                colorear_texto_celda(row.cells[5], COLORES_RGB['mujer'])   # Brecha azul
            # Si son iguales, no se aplica color (queda en negro)

        # SBC (columnas 6,7,8)
        if sbc_m != 0 and sbc_h != 0:
            if sbc_h > sbc_m:
                colorear_texto_celda(row.cells[6], COLORES_RGB['hombre'])
                colorear_texto_celda(row.cells[7], COLORES_RGB['mujer'])
                colorear_texto_celda(row.cells[8], COLORES_RGB['hombre'])
            elif sbc_m > sbc_h:
                colorear_texto_celda(row.cells[6], COLORES_RGB['mujer'])
                colorear_texto_celda(row.cells[7], COLORES_RGB['hombre'])
                colorear_texto_celda(row.cells[8], COLORES_RGB['mujer'])
            # Si son iguales, no se aplica color (queda en negro)

        # SBCE (columnas 9,10,11)
        if sbce_m != 0 and sbce_h != 0:
            if sbce_h > sbce_m:
                colorear_texto_celda(row.cells[9], COLORES_RGB['hombre'])
                colorear_texto_celda(row.cells[10], COLORES_RGB['mujer'])
                colorear_texto_celda(row.cells[11], COLORES_RGB['hombre'])
            elif sbce_m > sbce_h:
                colorear_texto_celda(row.cells[9], COLORES_RGB['mujer'])
                colorear_texto_celda(row.cells[10], COLORES_RGB['hombre'])
                colorear_texto_celda(row.cells[11], COLORES_RGB['mujer'])
            # Si son iguales, no se aplica color (queda en negro)

    # Añadir fila de totales si hay más de una fila de datos
    if len(datos_lista) > 1:
        row_total = table.add_row()

        # Calcular brechas de totales
        brecha_sb = ((totales['sb_H'] - totales['sb_M']) / totales['sb_H'] * 100) if totales['sb_H'] != 0 else 0
        brecha_sbc = ((totales['sbc_H'] - totales['sbc_M']) / totales['sbc_H'] * 100) if totales['sbc_H'] != 0 else 0
        brecha_sbce = ((totales['sbce_H'] - totales['sbce_M']) / totales['sbce_H'] * 100) if totales['sbce_H'] != 0 else 0

        valores_totales = [
            'TOTAL',
            str(totales['n_M']),
            str(totales['n_H']),
            formato_numero_es(totales['sb_M'], 2) + ' €',
            formato_numero_es(totales['sb_H'], 2) + ' €',
            formato_numero_es(abs(brecha_sb), 2) + ' %',
            formato_numero_es(totales['sbc_M'], 2) + ' €',
            formato_numero_es(totales['sbc_H'], 2) + ' €',
            formato_numero_es(abs(brecha_sbc), 2) + ' %',
            formato_numero_es(totales['sbce_M'], 2) + ' €',
            formato_numero_es(totales['sbce_H'], 2) + ' €',
            formato_numero_es(abs(brecha_sbce), 2) + ' %'
        ]

        for i, text in enumerate(valores_totales):
            aplicar_formato_celda(row_total.cells[i], text, font_size=7, bold=True)
            colorear_celda_fondo(row_total.cells[i], COLOR_FONDO_TOTAL)

        # Aplicar colores a la fila de totales
        if totales['sb_H'] > totales['sb_M']:
            colorear_texto_celda(row_total.cells[3], COLORES_RGB['hombre'])
            colorear_texto_celda(row_total.cells[4], COLORES_RGB['mujer'])
            colorear_texto_celda(row_total.cells[5], COLORES_RGB['hombre'])
        elif totales['sb_M'] > totales['sb_H']:
            colorear_texto_celda(row_total.cells[3], COLORES_RGB['mujer'])
            colorear_texto_celda(row_total.cells[4], COLORES_RGB['hombre'])
            colorear_texto_celda(row_total.cells[5], COLORES_RGB['mujer'])
        # Si son iguales, no se aplica color (queda en negro)

        if totales['sbc_H'] > totales['sbc_M']:
            colorear_texto_celda(row_total.cells[6], COLORES_RGB['hombre'])
            colorear_texto_celda(row_total.cells[7], COLORES_RGB['mujer'])
            colorear_texto_celda(row_total.cells[8], COLORES_RGB['hombre'])
        elif totales['sbc_M'] > totales['sbc_H']:
            colorear_texto_celda(row_total.cells[6], COLORES_RGB['mujer'])
            colorear_texto_celda(row_total.cells[7], COLORES_RGB['hombre'])
            colorear_texto_celda(row_total.cells[8], COLORES_RGB['mujer'])
        # Si son iguales, no se aplica color (queda en negro)

        if totales['sbce_H'] > totales['sbce_M']:
            colorear_texto_celda(row_total.cells[9], COLORES_RGB['hombre'])
            colorear_texto_celda(row_total.cells[10], COLORES_RGB['mujer'])
            colorear_texto_celda(row_total.cells[11], COLORES_RGB['hombre'])
        elif totales['sbce_M'] > totales['sbce_H']:
            colorear_texto_celda(row_total.cells[9], COLORES_RGB['mujer'])
            colorear_texto_celda(row_total.cells[10], COLORES_RGB['hombre'])
            colorear_texto_celda(row_total.cells[11], COLORES_RGB['mujer'])
        # Si son iguales, no se aplica color (queda en negro)

    doc.add_paragraph()


# ==================== CLASE PRINCIPAL ====================

class GeneradorInformeOptimizado:
    """Generador de informes de registro retributivo optimizado"""

    def __init__(self):
        self.df = None
        self.ruta_salida = None
        self.archivos_temp = []

    def cargar_datos(self):
        """Carga el archivo Excel más reciente de la carpeta 02_RESULTADOS"""
        log("Buscando archivo de datos...")

        ruta_base = Path(__file__).parent.parent
        carpeta_resultados = ruta_base / '02_RESULTADOS'

        if not carpeta_resultados.exists():
            log("Carpeta 02_RESULTADOS no encontrada", 'ERROR')
            return False

        archivos = list(carpeta_resultados.glob('REPORTE_*.xlsx'))
        archivos = [f for f in archivos if not f.name.startswith('~$')]

        if not archivos:
            log("No se encontraron archivos REPORTE_*.xlsx", 'ERROR')
            return False

        # Tomar el más reciente
        archivo = max(archivos, key=lambda x: x.stat().st_mtime)
        log(f"Cargando datos desde: {archivo.name}")

        try:
            self.df = pd.read_excel(archivo)

            # Mapear valores de la columna Sexo a formato corto
            if 'Sexo' in self.df.columns:
                self.df['Sexo'] = self.df['Sexo'].map({
                    'Hombres': 'H',
                    'Mujeres': 'M'
                }).fillna(self.df['Sexo'])  # Mantener valores no mapeados

            log(f"Datos cargados: {len(self.df)} registros", 'OK')
            return True
        except Exception as e:
            log(f"Error al cargar datos: {e}", 'ERROR')
            return False

    def cargar_datos_desde_bytes(self, excel_bytes):
        """
        Carga datos desde BytesIO para uso en Streamlit
        
        Args:
            excel_bytes: BytesIO o bytes del Excel procesado
        
        Returns:
            bool: True si se cargó correctamente
        """
        try:
            if isinstance(excel_bytes, bytes):
                excel_bytes = io.BytesIO(excel_bytes)
            
            excel_bytes.seek(0)
            self.df = pd.read_excel(excel_bytes, sheet_name='DATOS_PROCESADOS')
            
            # Limpiar espacios en nombres de columnas
            self.df.columns = self.df.columns.str.strip()
            
            # Mapear valores de la columna Sexo a formato corto
            if 'Sexo' in self.df.columns:
                self.df['Sexo'] = self.df['Sexo'].map({
                    'Hombres': 'H',
                    'Mujeres': 'M'
                }).fillna(self.df['Sexo'])
            
            log(f"Datos cargados desde BytesIO: {len(self.df)} registros", 'OK')
            return True
        except Exception as e:
            log(f"Error al cargar datos desde BytesIO: {e}", 'ERROR')
            return False

    def generar_informe_bytes(self, tipo_informe='CONSOLIDADO'):
        """
        Genera el informe Word y lo retorna como BytesIO para Streamlit
        
        Args:
            tipo_informe: 'CONSOLIDADO', 'PROMEDIO', 'MEDIANA', o 'COMPLEMENTOS'
        
        Returns:
            BytesIO con el documento Word generado o None si hay error
        """
        if self.df is None:
            log("No hay datos cargados", 'ERROR')
            return None
        
        log(f"Generando informe tipo: {tipo_informe}")
        
        # Generar el informe normalmente
        if not self.generar_informe(tipo_informe):
            return None
        
        # Leer el archivo generado y convertirlo a BytesIO
        try:
            if self.ruta_salida and self.ruta_salida.exists():
                with open(self.ruta_salida, 'rb') as f:
                    doc_bytes = io.BytesIO(f.read())
                
                # Limpiar el archivo temporal
                try:
                    self.ruta_salida.unlink()
                except:
                    pass
                
                return doc_bytes
            else:
                log("No se encontró el archivo generado", 'ERROR')
                return None
        except Exception as e:
            log(f"Error al convertir informe a BytesIO: {e}", 'ERROR')
            return None

    def generar_informe(self, tipo_informe='CONSOLIDADO'):
        """
        Genera el informe Word completo.

        tipo_informe: 'CONSOLIDADO', 'PROMEDIO', 'MEDIANA', o 'COMPLEMENTOS'
        """
        if self.df is None:
            log("No hay datos cargados", 'ERROR')
            return False

        log(f"Generando informe tipo: {tipo_informe}")

        # Determinar qué secciones generar
        configuraciones = {
            'CONSOLIDADO': {'metodos': ['media', 'mediana'], 'complementos': True},
            'PROMEDIO': {'metodos': ['media'], 'complementos': False},
            'MEDIANA': {'metodos': ['mediana'], 'complementos': False},
            'COMPLEMENTOS': {'metodos': [], 'complementos': True}
        }

        config = configuraciones.get(tipo_informe, {'metodos': ['media'], 'complementos': True})
        metodos_a_usar = config['metodos']
        incluir_complementos = config['complementos']

        # Cargar plantilla existente - Buscar en múltiples ubicaciones
        # 1. En templates/ (para Streamlit Cloud y despliegue)
        # 2. En 00_DOCUMENTACION/ (para uso local)
        plantilla_paths = [
            Path(__file__).parent.parent / 'templates' / 'plantilla_informe.docx',
            Path(__file__).parent.parent / '00_DOCUMENTACION' / 'Registro retributivo' / 'Reg Retributivo  NUEVA PLANTILLA.docx'
        ]
        
        plantilla_path = None
        for path in plantilla_paths:
            if path.exists():
                plantilla_path = path
                break
        
        if plantilla_path is None:
            log(f"ADVERTENCIA: No se encontró la plantilla en ninguna ubicación. Creando documento desde cero.", 'WARNING')
            doc = Document()

            # Configurar márgenes
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)

            # Título principal
            titulo = doc.add_heading('INFORME DE REGISTRO RETRIBUTIVO', 0)
            titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

            fecha = datetime.now().strftime('%d/%m/%Y')
            p = doc.add_paragraph(f'Fecha: {fecha}')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_page_break()

        else:
            log(f"Usando plantilla: {plantilla_path}")
            doc = Document(str(plantilla_path))


        # Generar secciones de PROMEDIOS y MEDIANAS (si aplica)
        if metodos_a_usar:
            for metodo in metodos_a_usar:
                metodo_label = 'PROMEDIO' if metodo == 'media' else 'MEDIANA'

                # Título del bloque principal
                doc.add_heading(f'ANÁLISIS CON {metodo_label}', 1)
                doc.add_page_break()

                # 1. Análisis General
                self._generar_analisis_general(doc, metodo, metodo_label)

                # 2. Análisis por Grupo Profesional
                self._generar_analisis_grupo_profesional(doc, metodo, metodo_label)

                # 3. Análisis por Puesto de Trabajo
                self._generar_analisis_puesto(doc, metodo, metodo_label)

                # 4. Análisis por Escala SVPT
                self._generar_analisis_escalas(doc, metodo, metodo_label)

                # 5. Análisis por Nivel
                self._generar_analisis_nivel(doc, metodo, metodo_label)

                # Separador entre bloques principales (Promedios/Medianas y Complementos)
                if metodo == metodos_a_usar[-1] and incluir_complementos:
                    doc.add_page_break()

        # BLOQUE 3: ANÁLISIS DE COMPLEMENTOS (si aplica)
        if incluir_complementos:
            self._generar_analisis_complementos(doc)

        # Guardar documento
        self._guardar_documento(doc, tipo_informe)

        # Limpiar archivos temporales
        self._limpiar_temporales()

        log(f"Informe generado exitosamente: {self.ruta_salida}", 'OK')
        return True

    def _generar_analisis_general(self, doc, metodo, metodo_label):
        """Genera la sección de análisis general de salarios"""
        log(f"Generando análisis general ({metodo_label})...")

        doc.add_heading(f'1. ANÁLISIS GENERAL DE SALARIOS ({metodo_label})', 2)

        # Calcular para SB, SB+C, SB+C+ES (efectivo y equiparado)
        configs = [
            ('Salario Base (SB)', COLS['sb_efectivo'], COLS['sb_equiparado'], 'efectivos_sb', 'equiparados_sb'),
            ('SB + Complementos (SB+C)', COLS['sbc_efectivo'], COLS['sbc_equiparado'], 'efectivos_sb_complementos', 'equiparados_sb_complementos'),
            ('SB + Complementos + Extrasalariales (SB+C+ES)', COLS['sbce_efectivo'], COLS['sbce_equiparado'], 'efectivos_sb_complementos', 'equiparados_sb_complementos')
        ]

        for titulo, col_efe, col_equi, tipo_efe, tipo_equi in configs:
            # Calcular efectivo
            stats_efe = calcular_estadistico(self.df, col_efe, tipo_efe, metodo)
            brecha_efe = calcular_brecha(stats_efe['H'], stats_efe['M']) or 0

            # Calcular equiparado
            stats_equi = calcular_estadistico(self.df, col_equi, tipo_equi, metodo)
            brecha_equi = calcular_brecha(stats_equi['H'], stats_equi['M']) or 0

            # Usar los conteos reales de cada cálculo (pueden diferir por filtros aplicados)
            n_m_efe = stats_efe.get('n_M', 0)
            n_h_efe = stats_efe.get('n_H', 0)
            n_m_equi = stats_equi.get('n_M', 0)
            n_h_equi = stats_equi.get('n_H', 0)

            datos = {
                'efectivo': {**stats_efe, 'brecha': brecha_efe},
                'equiparado': {**stats_equi, 'brecha': brecha_equi}
            }

            # Tabla
            crear_tabla_resumen(doc, titulo, datos, metodo_label)

            # Gráficos donut - crear archivos
            archivo_efe = f"temp_donut_{titulo.replace(' ', '_').replace('(', '').replace(')', '')}_efectivo.png"
            archivo_equi = f"temp_donut_{titulo.replace(' ', '_').replace('(', '').replace(')', '')}_equiparado.png"

            # Extraer código del concepto (ej: "SB", "SB+C", "SB+C+ES")
            import re
            match = re.search(r'\(([^)]+)\)', titulo)
            codigo_concepto = match.group(1) if match else titulo

            img_efe_creada = crear_grafico_donut(
                {'M': stats_efe['M'], 'H': stats_efe['H'], 'brecha': brecha_efe, 'n_M': n_m_efe, 'n_H': n_h_efe},
                f'{metodo_label} {codigo_concepto} EFECTIVO',
                archivo_efe
            )

            img_equi_creada = crear_grafico_donut(
                {'M': stats_equi['M'], 'H': stats_equi['H'], 'brecha': brecha_equi, 'n_M': n_m_equi, 'n_H': n_h_equi},
                f'{metodo_label} {codigo_concepto} EQUIPARADO',
                archivo_equi
            )

            # Insertar gráficos lado a lado (2 por fila) usando tabla
            if img_efe_creada and img_equi_creada:
                self.archivos_temp.extend([archivo_efe, archivo_equi])

                # Crear tabla de 1 fila x 2 columnas para gráficos lado a lado
                tabla_graficos = doc.add_table(rows=1, cols=2)

                # Ocultar bordes de la tabla
                ocultar_bordes_tabla(tabla_graficos)

                # Añadir gráfico efectivo (izquierda - 50%)
                celda_izq = tabla_graficos.rows[0].cells[0]
                paragraph_izq = celda_izq.paragraphs[0]
                run_izq = paragraph_izq.add_run()
                run_izq.add_picture(archivo_efe, width=Inches(GRAFICO_CFG['ancho_donut']))
                paragraph_izq.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Añadir gráfico equiparado (derecha - 50%)
                celda_der = tabla_graficos.rows[0].cells[1]
                paragraph_der = celda_der.paragraphs[0]
                run_der = paragraph_der.add_run()
                run_der.add_picture(archivo_equi, width=Inches(GRAFICO_CFG['ancho_donut']))
                paragraph_der.alignment = WD_ALIGN_PARAGRAPH.CENTER

                doc.add_paragraph()  # Espaciado después de los gráficos

        doc.add_page_break()

    def _generar_analisis_grupo_profesional(self, doc, metodo, metodo_label):
        """Genera análisis por grupo profesional"""
        log(f"Generando análisis por grupo profesional ({metodo_label})...")

        doc.add_heading(f'2. ANÁLISIS POR GRUPO PROFESIONAL ({metodo_label})', 2)

        # Obtener grupos únicos
        df_actual = self.df[self.df[COLS['reg']] != 'Ex']
        grupos = sorted([str(g) for g in df_actual[COLS['grupo_prof']].dropna().unique()])

        # Calcular para efectivo y equiparado
        for tipo_datos, sufijo in [('efectivo', 'EFECTIVO'), ('equiparado', 'EQUIPARADO')]:
            datos_grupos = []

            for grupo in grupos:
                # SB
                tipo_calc_sb = f'efectivos_sb' if tipo_datos == 'efectivo' else 'equiparados_sb'
                col_sb = COLS[f'sb_{tipo_datos}']
                stats_sb = calcular_estadistico(self.df, col_sb, tipo_calc_sb, metodo, COLS['grupo_prof'], grupo)
                brecha_sb = calcular_brecha(stats_sb['H'], stats_sb['M']) or 0

                # SBC
                tipo_calc_sbc = f'efectivos_sb_complementos' if tipo_datos == 'efectivo' else 'equiparados_sb_complementos'
                col_sbc = COLS[f'sbc_{tipo_datos}']
                stats_sbc = calcular_estadistico(self.df, col_sbc, tipo_calc_sbc, metodo, COLS['grupo_prof'], grupo)
                brecha_sbc = calcular_brecha(stats_sbc['H'], stats_sbc['M']) or 0

                # SBCE
                col_sbce = COLS[f'sbce_{tipo_datos}']
                stats_sbce = calcular_estadistico(self.df, col_sbce, tipo_calc_sbc, metodo, COLS['grupo_prof'], grupo)
                brecha_sbce = calcular_brecha(stats_sbce['H'], stats_sbce['M']) or 0

                datos_grupos.append({
                    'grupo': grupo,
                    'n_M': stats_sb['n_M'],
                    'n_H': stats_sb['n_H'],
                    'sb_M': stats_sb['M'],
                    'sb_H': stats_sb['H'],
                    'sb_brecha': brecha_sb,
                    'sbc_M': stats_sbc['M'],
                    'sbc_H': stats_sbc['H'],
                    'sbc_brecha': brecha_sbc,
                    'sbce_M': stats_sbce['M'],
                    'sbce_H': stats_sbce['H'],
                    'sbce_brecha': brecha_sbce
                })

            # Crear tabla
            crear_tabla_por_grupo(doc, f'Retribución {metodo_label} por Grupo Profesional {sufijo}', datos_grupos, metodo_label)

            # Gráfico de barras acumulativo (SB + Complementos Salariales + Complementos Extrasalariales)
            datos_grafico = [{
                'categoria': d['grupo'],
                'sb_M': d['sb_M'],
                'sb_H': d['sb_H'],
                'sbc_M': d['sbc_M'],
                'sbc_H': d['sbc_H'],
                'sbce_M': d['sbce_M'],
                'sbce_H': d['sbce_H'],
                'n_M': d['n_M'],
                'n_H': d['n_H']
            } for d in datos_grupos]
            archivo = f'temp_barras_grupo_prof_{tipo_datos}.png'
            resultado = crear_grafico_barras_acumulativo(
                datos_grafico,
                f'Salarios {metodo_label} CON Complementos + ES por Grupo Profesional {sufijo}',
                archivo
            )
            if resultado is not None and os.path.exists(archivo):
                self.archivos_temp.append(archivo)
                doc.add_picture(archivo, width=Inches(GRAFICO_CFG['ancho_doc']))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()

        doc.add_page_break()

    def _generar_analisis_puesto(self, doc, metodo, metodo_label):
        """Genera análisis por Agrupación (Nivel SVPT) y Puesto de trabajo"""
        log(f"Generando análisis por puesto ({metodo_label})...")

        doc.add_heading(f'3. RETRIBUCIÓN POR AGRUPACIÓN (NIVEL SVPT) Y PUESTO DE TRABAJO ({metodo_label})', 2)

        # Obtener combinaciones únicas de Nivel SVPT + Puesto
        df_actual = self.df[self.df[COLS['reg']] != 'Ex'].copy()
        df_actual['combinacion'] = df_actual[COLS['nivel_svpt']].astype(str) + ' + ' + df_actual[COLS['puesto']].astype(str)
        combinaciones = sorted(df_actual['combinacion'].dropna().unique())

        # Para cada tipo (efectivo/equiparado)
        for tipo_datos, sufijo in [('efectivo', 'EFECTIVO'), ('equiparado', 'EQUIPARADO')]:
            datos_puestos = []

            for comb in combinaciones:
                df_comb = df_actual[df_actual['combinacion'] == comb]

                # SB
                col_sb = COLS[f'sb_{tipo_datos}']
                stats_sb = {'M': 0, 'H': 0, 'n_M': 0, 'n_H': 0}
                for sexo in ['M', 'H']:
                    df_sexo = df_comb[(df_comb[COLS['sexo']] == sexo) & (df_comb[col_sb] > 0)]
                    stats_sb[f'n_{sexo}'] = len(df_sexo)
                    if len(df_sexo) > 0:
                        stats_sb[sexo] = df_sexo[col_sb].mean() if metodo == 'media' else df_sexo[col_sb].median()
                brecha_sb = calcular_brecha(stats_sb['H'], stats_sb['M']) or 0

                # SBC
                col_sbc = COLS[f'sbc_{tipo_datos}']
                stats_sbc = {'M': 0, 'H': 0}
                for sexo in ['M', 'H']:
                    df_sexo = df_comb[df_comb[COLS['sexo']] == sexo]
                    if len(df_sexo) > 0:
                        stats_sbc[sexo] = df_sexo[col_sbc].mean() if metodo == 'media' else df_sexo[col_sbc].median()
                brecha_sbc = calcular_brecha(stats_sbc['H'], stats_sbc['M']) or 0

                # SBCE
                col_sbce = COLS[f'sbce_{tipo_datos}']
                stats_sbce = {'M': 0, 'H': 0}
                for sexo in ['M', 'H']:
                    df_sexo = df_comb[df_comb[COLS['sexo']] == sexo]
                    if len(df_sexo) > 0:
                        stats_sbce[sexo] = df_sexo[col_sbce].mean() if metodo == 'media' else df_sexo[col_sbce].median()
                brecha_sbce = calcular_brecha(stats_sbce['H'], stats_sbce['M']) or 0

                # Obtener puntos promedio para esta combinación
                puntos_promedio = 0
                if COLS['puntos'] in df_comb.columns and len(df_comb) > 0:
                    puntos_promedio = df_comb[COLS['puntos']].mean()

                if stats_sb['n_M'] > 0 or stats_sb['n_H'] > 0:
                    datos_puestos.append({
                        'grupo': comb,
                        'n_M': stats_sb['n_M'],
                        'n_H': stats_sb['n_H'],
                        'sb_M': stats_sb['M'],
                        'sb_H': stats_sb['H'],
                        'sb_brecha': brecha_sb,
                        'sbc_M': stats_sbc['M'],
                        'sbc_H': stats_sbc['H'],
                        'sbc_brecha': brecha_sbc,
                        'sbce_M': stats_sbce['M'],
                        'sbce_H': stats_sbce['H'],
                        'sbce_brecha': brecha_sbce,
                        'puntos': puntos_promedio
                    })

            # Usar la función estándar de tablas
            crear_tabla_por_grupo(doc, f'Retribución {metodo_label} por Puesto de Trabajo {sufijo}', datos_puestos, metodo_label)

            # Gráfico de barras acumulativo horizontal (SB + Complementos Salariales + Complementos Extrasalariales)
            datos_grafico = [{
                'categoria': d['grupo'],
                'sb_M': d['sb_M'],
                'sb_H': d['sb_H'],
                'sbc_M': d['sbc_M'],
                'sbc_H': d['sbc_H'],
                'sbce_M': d['sbce_M'],
                'sbce_H': d['sbce_H'],
                'n_M': d['n_M'],
                'n_H': d['n_H']
            } for d in datos_puestos]
            archivo = f'temp_barras_puesto_{tipo_datos}.png'
            resultado = crear_grafico_barras_acumulativo(
                datos_grafico,
                f'Salarios {metodo_label} CON Complementos + ES por Puesto de Trabajo {sufijo}',
                archivo
            )
            if resultado is not None and os.path.exists(archivo):
                self.archivos_temp.append(archivo)
                doc.add_picture(archivo, width=Inches(GRAFICO_CFG['ancho_doc']))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()

            # Gráfico de barras vertical con doble eje (salarios + puntos) o barras simple si no hay puntos
            # APARTADO 3: usar tipo 'denso' porque hay muchas categorías (puestos de trabajo)
            datos_grafico_dual = [{'categoria': d['grupo'], 'M': d['sbc_M'], 'H': d['sbc_H'], 'puntos': d.get('puntos', 0), 'n_M': d.get('n_M', 0), 'n_H': d.get('n_H', 0)} for d in datos_puestos]
            archivo_dual = f'temp_barras_dual_puesto_{tipo_datos}.png'
            resultado = crear_grafico_salarios_con_puntos_condicional(
                datos_grafico_dual,
                f'Salarios {metodo_label} CON Complementos por Puesto de Trabajo {sufijo}',
                archivo_dual,
                tipo_grafico='denso'  # Muchas categorías: usar fuentes más pequeñas
            )
            if resultado is not None and os.path.exists(archivo_dual):
                self.archivos_temp.append(archivo_dual)
                doc.add_picture(archivo_dual, width=Inches(GRAFICO_CFG['ancho_doc']))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()

        doc.add_page_break()

    def _generar_analisis_nivel(self, doc, metodo, metodo_label):
        """Genera análisis por nivel de convenio"""
        log(f"Generando análisis por nivel ({metodo_label})...")

        doc.add_heading(f'5. ANÁLISIS POR NIVEL ({metodo_label})', 2)

        # Obtener niveles únicos
        df_actual = self.df[self.df[COLS['reg']] != 'Ex']
        niveles = sorted([str(n) for n in df_actual[COLS['nivel_convenio']].dropna().unique()])

        # Para cada tipo
        for tipo_datos, sufijo in [('efectivo', 'EFECTIVO'), ('equiparado', 'EQUIPARADO')]:
            datos_niveles = []

            for nivel in niveles:
                # Filtrar salarios > 0 para SB
                tipo_calc_sb = f'efectivos_sb' if tipo_datos == 'efectivo' else 'equiparados_sb'
                tipo_calc_sbc = f'efectivos_sb_complementos' if tipo_datos == 'efectivo' else 'equiparados_sb_complementos'

                # SB
                col_sb = COLS[f'sb_{tipo_datos}']
                stats_sb = calcular_estadistico(self.df, col_sb, tipo_calc_sb, metodo, COLS['nivel_convenio'], nivel)
                brecha_sb = calcular_brecha(stats_sb['H'], stats_sb['M']) or 0

                # SBC
                col_sbc = COLS[f'sbc_{tipo_datos}']
                stats_sbc = calcular_estadistico(self.df, col_sbc, tipo_calc_sbc, metodo, COLS['nivel_convenio'], nivel)
                brecha_sbc = calcular_brecha(stats_sbc['H'], stats_sbc['M']) or 0

                # SBCE
                col_sbce = COLS[f'sbce_{tipo_datos}']
                stats_sbce = calcular_estadistico(self.df, col_sbce, tipo_calc_sbc, metodo, COLS['nivel_convenio'], nivel)
                brecha_sbce = calcular_brecha(stats_sbce['H'], stats_sbce['M']) or 0

                datos_niveles.append({
                    'grupo': nivel,
                    'n_M': stats_sb['n_M'],
                    'n_H': stats_sb['n_H'],
                    'sb_M': stats_sb['M'],
                    'sb_H': stats_sb['H'],
                    'sb_brecha': brecha_sb,
                    'sbc_M': stats_sbc['M'],
                    'sbc_H': stats_sbc['H'],
                    'sbc_brecha': brecha_sbc,
                    'sbce_M': stats_sbce['M'],
                    'sbce_H': stats_sbce['H'],
                    'sbce_brecha': brecha_sbce
                })

            # Usar la función estándar de tablas
            crear_tabla_por_grupo(doc, f'Retribución {metodo_label} por Nivel {sufijo}', datos_niveles, metodo_label)

            doc.add_paragraph()

            # Gráfico de barras (SB+C+ES)
            datos_grafico = [{'categoria': d['grupo'], 'M': d['sbce_M'], 'H': d['sbce_H'], 'n_M': d['n_M'], 'n_H': d['n_H']} for d in datos_niveles]
            archivo = f'temp_barras_nivel_{tipo_datos}.png'
            resultado = crear_grafico_barras(
                datos_grafico,
                f'Salarios {metodo_label} CON Complementos + ES por Nivel {sufijo}',
                'Salario (€)',
                archivo,
                'vertical',
                etiqueta_categoria='Nivel convenio colectivo'
            )
            if resultado is not None and os.path.exists(archivo):
                self.archivos_temp.append(archivo)
                doc.add_picture(archivo, width=Inches(GRAFICO_CFG['ancho_doc']))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()

        doc.add_page_break()

    def _generar_analisis_escalas(self, doc, metodo, metodo_label):
        """Genera análisis agrupado por escalas E1-E5"""
        log(f"Generando análisis por escalas SVPT ({metodo_label})...")

        doc.add_heading(f'4. ANÁLISIS POR ESCALA SVPT (E1-E5) ({metodo_label})', 2)

        escalas = obtener_escalas_svpt(self.df)

        if not escalas:
            log("No se encontraron escalas SVPT", 'WARN')
            return

        # Para cada escala
        for escala in escalas:
            # Evitar duplicar "Escala" si ya está en el nombre
            nombre_escala = escala if escala.lower().startswith('escala') else f'Escala {escala}'
            doc.add_heading(nombre_escala, level=3)

            # Filtrar puestos de esta escala
            df_actual = self.df[self.df[COLS['reg']] != 'Ex'].copy()
            df_escala = df_actual[df_actual[COLS['nivel_svpt']].astype(str).str.contains(escala, case=False, na=False)]

            if df_escala.empty:
                doc.add_paragraph(f'Sin datos para escala {escala}')
                continue

            puestos = sorted(df_escala[COLS['puesto']].dropna().unique())

            # Para cada tipo
            for tipo_datos, sufijo in [('efectivo', 'EFECTIVO'), ('equiparado', 'EQUIPARADO')]:
                datos_puestos = []

                for puesto in puestos:
                    df_puesto = df_escala[df_escala[COLS['puesto']] == puesto]

                    tipo_calc = f'efectivos_sb_complementos' if tipo_datos == 'efectivo' else 'equiparados_sb_complementos'
                    col_sbc = COLS[f'sbc_{tipo_datos}']

                    stats = {'M': 0, 'H': 0, 'n_M': 0, 'n_H': 0}
                    for sexo in ['M', 'H']:
                        df_sexo = df_puesto[df_puesto[COLS['sexo']] == sexo]
                        stats[f'n_{sexo}'] = len(df_sexo)
                        if len(df_sexo) > 0:
                            if metodo == 'media':
                                stats[sexo] = df_sexo[col_sbc].mean()
                            else:
                                stats[sexo] = df_sexo[col_sbc].median()

                    # Obtener puntos promedio para este puesto
                    puntos_promedio = 0
                    if COLS['puntos'] in df_puesto.columns and len(df_puesto) > 0:
                        puntos_promedio = df_puesto[COLS['puntos']].mean()

                    if stats['n_M'] > 0 or stats['n_H'] > 0:
                        datos_puestos.append({
                            'categoria': puesto,
                            'M': stats['M'],
                            'H': stats['H'],
                            'n_M': stats['n_M'],
                            'n_H': stats['n_H'],
                            'puntos': puntos_promedio
                        })

                # Gráfico VERTICAL con doble eje (salarios + puntos) o barras simple si no hay puntos
                # APARTADO 4: usar tipo 'simple' porque hay pocas categorías (escalas)
                if datos_puestos:
                    archivo = f'temp_barras_dual_escala_{escala}_{tipo_datos}.png'
                    resultado = crear_grafico_salarios_con_puntos_condicional(
                        datos_puestos,
                        f'Salarios {metodo_label} CON Complementos - Escala {escala} {sufijo}',
                        archivo,
                        tipo_grafico='simple'  # Pocas categorías: usar fuentes MUY GRANDES
                    )
                    if resultado is not None and os.path.exists(archivo):
                        self.archivos_temp.append(archivo)
                        doc.add_picture(archivo, width=Inches(GRAFICO_CFG['ancho_doc']))
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        doc.add_paragraph()

        doc.add_page_break()

    def _generar_analisis_complementos(self, doc):
        """Genera análisis de complementos salariales y extrasalariales (bloque independiente)"""
        log("Generando análisis de complementos...")

        doc.add_heading('ANÁLISIS DE COMPLEMENTOS', 1)

        # ==================== COMPLEMENTOS SALARIALES ====================
        doc.add_heading('Complementos Salariales', level=2)

        # Gráficos donut - Efectivo y Equiparado
        col_comp_efe = COLS['comp_efectivo']
        col_comp_equi = COLS['comp_equiparado']

        stats_efe = calcular_estadistico(self.df, col_comp_efe, 'efectivos_sb_complementos', 'media')
        brecha_efe = calcular_brecha(stats_efe['H'], stats_efe['M']) or 0

        stats_equi = calcular_estadistico(self.df, col_comp_equi, 'equiparados_sb_complementos', 'media')
        brecha_equi = calcular_brecha(stats_equi['H'], stats_equi['M']) or 0

        archivo_efe = 'temp_donut_comp_salariales_efectivo.png'
        archivo_equi = 'temp_donut_comp_salariales_equiparado.png'

        img_efe_creada = crear_grafico_donut(
            {'M': stats_efe['M'], 'H': stats_efe['H'], 'brecha': brecha_efe},
            'PROMEDIO Complementos Salariales - EFECTIVO',
            archivo_efe
        )

        img_equi_creada = crear_grafico_donut(
            {'M': stats_equi['M'], 'H': stats_equi['H'], 'brecha': brecha_equi},
            'PROMEDIO Complementos Salariales - EQUIPARADO',
            archivo_equi
        )

        # Insertar gráficos lado a lado
        if img_efe_creada and img_equi_creada:
            self.archivos_temp.extend([archivo_efe, archivo_equi])

            tabla_graficos = doc.add_table(rows=1, cols=2)

            # Ocultar bordes
            ocultar_bordes_tabla(tabla_graficos)

            # Añadir gráfico efectivo
            celda_izq = tabla_graficos.rows[0].cells[0]
            paragraph_izq = celda_izq.paragraphs[0]
            run_izq = paragraph_izq.add_run()
            run_izq.add_picture(archivo_efe, width=Inches(GRAFICO_CFG['ancho_donut']))
            paragraph_izq.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Añadir gráfico equiparado
            celda_der = tabla_graficos.rows[0].cells[1]
            paragraph_der = celda_der.paragraphs[0]
            run_der = paragraph_der.add_run()
            run_der.add_picture(archivo_equi, width=Inches(GRAFICO_CFG['ancho_donut']))
            paragraph_der.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()

        # Tabla de distribución de complementos salariales
        df_actual = self.df[self.df[COLS['reg']] != 'Ex']

        # Buscar todas las columnas PS (formato: "PS 1 Antigüedad" o "PS1")
        import re
        cols_ps = [col for col in self.df.columns
                   if col.startswith('PS') and bool(re.match(r'^PS\s*\d+', col))]

        # Filtrar solo las columnas base (sin _equiparado)
        cols_ps_base = [col for col in cols_ps if '_equiparado' not in col]

        # Ordenar numéricamente por el número después de PS
        def extraer_numero_ps(col):
            try:
                # Extraer el número que sigue a PS (puede haber espacio o no)
                match = re.search(r'PS\s*(\d+)', col)
                return int(match.group(1)) if match else 999
            except:
                return 999
        cols_ps_base = sorted(cols_ps_base, key=extraer_numero_ps)

        datos_complementos = []
        for col_comp in cols_ps_base:
            df_con_comp = df_actual[df_actual[col_comp] > 0]
            n_m = len(df_con_comp[df_con_comp[COLS['sexo']] == 'M'])
            n_h = len(df_con_comp[df_con_comp[COLS['sexo']] == 'H'])
            total = n_m + n_h

            if total > 0:
                # Calcular porcentaje respecto al total de personas con ese complemento
                pct_m = (n_m / total * 100) if total > 0 else 0
                pct_h = (n_h / total * 100) if total > 0 else 0

                datos_complementos.append({
                    'complemento': col_comp,
                    'n_M': n_m,
                    'n_H': n_h,
                    'total': total,
                    'pct_M': pct_m,
                    'pct_H': pct_h
                })

        if datos_complementos:
            crear_tabla_complementos(doc, 'Distribución de Complementos Salariales', datos_complementos)

        # Tabla de complementos salariales por puesto de trabajo
        puestos = sorted(df_actual[COLS['puesto']].dropna().unique())

        datos_comp_puesto = []
        for puesto in puestos:
            df_puesto = df_actual[df_actual[COLS['puesto']] == puesto]

            for col_comp in cols_ps_base:
                df_puesto_comp = df_puesto[df_puesto[col_comp] > 0]
                n_m = len(df_puesto_comp[df_puesto_comp[COLS['sexo']] == 'M'])
                n_h = len(df_puesto_comp[df_puesto_comp[COLS['sexo']] == 'H'])
                total = n_m + n_h

                if total > 0:
                    pct_m = (n_m / total * 100) if total > 0 else 0
                    pct_h = (n_h / total * 100) if total > 0 else 0

                    datos_comp_puesto.append({
                        'puesto': puesto,
                        'complemento': col_comp,
                        'n_M': n_m,
                        'n_H': n_h,
                        'total': total,
                        'pct_M': pct_m,
                        'pct_H': pct_h
                    })

        if datos_comp_puesto:
            crear_tabla_complementos_por_puesto(doc, 'Distribución de Complementos Salariales por Puesto de Trabajo', datos_comp_puesto)

        # Gráficos de barras horizontales de promedio de complementos salariales
        # Efectivos
        datos_grafico_efectivos = []
        for col_comp in cols_ps_base:
            df_con_comp = df_actual[df_actual[col_comp] > 0]
            if len(df_con_comp) > 0:
                promedio_m = df_con_comp[df_con_comp[COLS['sexo']] == 'M'][col_comp].mean() if len(df_con_comp[df_con_comp[COLS['sexo']] == 'M']) > 0 else 0
                promedio_h = df_con_comp[df_con_comp[COLS['sexo']] == 'H'][col_comp].mean() if len(df_con_comp[df_con_comp[COLS['sexo']] == 'H']) > 0 else 0

                if promedio_m > 0 or promedio_h > 0:
                    datos_grafico_efectivos.append({
                        'categoria': col_comp,
                        'M': promedio_m,
                        'H': promedio_h,
                        'n_M': len(df_con_comp[df_con_comp[COLS['sexo']] == 'M']),
                        'n_H': len(df_con_comp[df_con_comp[COLS['sexo']] == 'H'])
                    })

        # Aplicar filtro de privacidad antes de dividir en bloques
        datos_grafico_efectivos = aplicar_filtro_privacidad_datos(datos_grafico_efectivos)

        # Dividir en bloques de 20 complementos (ya filtrados por privacidad)
        if datos_grafico_efectivos:
            tamanio_bloque = 20
            for idx in range(0, len(datos_grafico_efectivos), tamanio_bloque):
                bloque = datos_grafico_efectivos[idx:idx + tamanio_bloque]
                es_primer_bloque = (idx == 0)

                archivo = f'temp_barras_comp_salariales_efectivos_{idx}.png'
                # Los datos ya están filtrados, no aplicar filtro de nuevo en crear_grafico_barras
                resultado = crear_grafico_barras(
                    bloque,
                    'Complementos Promedio Efectivos por Complemento Salarial * Sexo',
                    'Promedio (€)',
                    archivo,
                    'horizontal',
                    mostrar_titulo=es_primer_bloque,
                    mostrar_leyenda=es_primer_bloque,
                    aplicar_filtro_privacidad=False  # Ya aplicado antes
                )
                if resultado is not None and os.path.exists(archivo):
                    self.archivos_temp.append(archivo)
                    doc.add_picture(archivo, width=Inches(GRAFICO_CFG['ancho_doc']))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()

        # Equiparados
        datos_grafico_equiparados = []
        for col_comp in cols_ps_base:
            # Buscar la columna equiparada correspondiente
            col_comp_equi = f"{col_comp}_equiparado"
            if col_comp_equi in self.df.columns:
                df_con_comp = df_actual[df_actual[col_comp_equi] > 0]
                if len(df_con_comp) > 0:
                    promedio_m = df_con_comp[df_con_comp[COLS['sexo']] == 'M'][col_comp_equi].mean() if len(df_con_comp[df_con_comp[COLS['sexo']] == 'M']) > 0 else 0
                    promedio_h = df_con_comp[df_con_comp[COLS['sexo']] == 'H'][col_comp_equi].mean() if len(df_con_comp[df_con_comp[COLS['sexo']] == 'H']) > 0 else 0

                    if promedio_m > 0 or promedio_h > 0:
                        datos_grafico_equiparados.append({
                            'categoria': col_comp,
                            'M': promedio_m,
                            'H': promedio_h,
                            'n_M': len(df_con_comp[df_con_comp[COLS['sexo']] == 'M']),
                            'n_H': len(df_con_comp[df_con_comp[COLS['sexo']] == 'H'])
                        })

        # Aplicar filtro de privacidad antes de dividir en bloques
        datos_grafico_equiparados = aplicar_filtro_privacidad_datos(datos_grafico_equiparados)

        # Dividir en bloques de 20 complementos (ya filtrados por privacidad)
        if datos_grafico_equiparados:
            tamanio_bloque = 20
            for idx in range(0, len(datos_grafico_equiparados), tamanio_bloque):
                bloque = datos_grafico_equiparados[idx:idx + tamanio_bloque]
                es_primer_bloque = (idx == 0)

                archivo = f'temp_barras_comp_salariales_equiparados_{idx}.png'
                # Los datos ya están filtrados, no aplicar filtro de nuevo en crear_grafico_barras
                resultado = crear_grafico_barras(
                    bloque,
                    'Complementos Promedio Equiparados por Complemento Salarial * Sexo',
                    'Promedio (€)',
                    archivo,
                    'horizontal',
                    mostrar_titulo=es_primer_bloque,
                    mostrar_leyenda=es_primer_bloque,
                    aplicar_filtro_privacidad=False  # Ya aplicado antes
                )
                if resultado is not None and os.path.exists(archivo):
                    self.archivos_temp.append(archivo)
                    doc.add_picture(archivo, width=Inches(GRAFICO_CFG['ancho_doc']))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()

        # Tablas de complementos salariales por grupo profesional
        grupos_prof = sorted([str(g) for g in df_actual[COLS['grupo_prof']].dropna().unique()])

        # Efectivos
        datos_grupos_efectivos = []
        for grupo in grupos_prof:
            df_grupo = df_actual[df_actual[COLS['grupo_prof']].astype(str) == grupo]

            # Calcular promedio y mediana del total de complementos salariales para este grupo
            df_grupo_m = df_grupo[df_grupo[COLS['sexo']] == 'M']
            df_grupo_h = df_grupo[df_grupo[COLS['sexo']] == 'H']

            col_comp_efe = COLS['comp_efectivo']

            # Promedio
            promedio_m = df_grupo_m[col_comp_efe].mean() if len(df_grupo_m) > 0 else 0
            promedio_h = df_grupo_h[col_comp_efe].mean() if len(df_grupo_h) > 0 else 0
            promedio_brecha = calcular_brecha(promedio_h, promedio_m) or 0

            # Mediana
            mediana_m = df_grupo_m[col_comp_efe].median() if len(df_grupo_m) > 0 else 0
            mediana_h = df_grupo_h[col_comp_efe].median() if len(df_grupo_h) > 0 else 0
            mediana_brecha = calcular_brecha(mediana_h, mediana_m) or 0

            datos_grupos_efectivos.append({
                'grupo': grupo,
                'n_M': len(df_grupo_m),
                'n_H': len(df_grupo_h),
                'promedio_M': promedio_m,
                'promedio_H': promedio_h,
                'promedio_brecha': promedio_brecha,
                'mediana_M': mediana_m,
                'mediana_H': mediana_h,
                'mediana_brecha': mediana_brecha
            })

        if datos_grupos_efectivos:
            crear_tabla_complementos_por_grupo(doc, 'Complementos Salariales por Grupos Profesionales - EFECTIVOS', datos_grupos_efectivos)

        # Equiparados
        datos_grupos_equiparados = []
        for grupo in grupos_prof:
            df_grupo = df_actual[df_actual[COLS['grupo_prof']].astype(str) == grupo]

            df_grupo_m = df_grupo[df_grupo[COLS['sexo']] == 'M']
            df_grupo_h = df_grupo[df_grupo[COLS['sexo']] == 'H']

            col_comp_equi = COLS['comp_equiparado']

            # Promedio
            promedio_m = df_grupo_m[col_comp_equi].mean() if len(df_grupo_m) > 0 else 0
            promedio_h = df_grupo_h[col_comp_equi].mean() if len(df_grupo_h) > 0 else 0
            promedio_brecha = calcular_brecha(promedio_h, promedio_m) or 0

            # Mediana
            mediana_m = df_grupo_m[col_comp_equi].median() if len(df_grupo_m) > 0 else 0
            mediana_h = df_grupo_h[col_comp_equi].median() if len(df_grupo_h) > 0 else 0
            mediana_brecha = calcular_brecha(mediana_h, mediana_m) or 0

            datos_grupos_equiparados.append({
                'grupo': grupo,
                'n_M': len(df_grupo_m),
                'n_H': len(df_grupo_h),
                'promedio_M': promedio_m,
                'promedio_H': promedio_h,
                'promedio_brecha': promedio_brecha,
                'mediana_M': mediana_m,
                'mediana_H': mediana_h,
                'mediana_brecha': mediana_brecha
            })

        if datos_grupos_equiparados:
            crear_tabla_complementos_por_grupo(doc, 'Complementos Salariales por Grupos Profesionales - EQUIPARADOS', datos_grupos_equiparados)

        # Gráficos de barras horizontales por grupo profesional - Efectivos
        if datos_grupos_efectivos:
            datos_grafico_grupos_efe = []
            for d in datos_grupos_efectivos:
                if d['promedio_M'] > 0 or d['promedio_H'] > 0:
                    datos_grafico_grupos_efe.append({
                        'categoria': d['grupo'],
                        'M': d['promedio_M'],
                        'H': d['promedio_H'],
                        'n_M': d['n_M'],
                        'n_H': d['n_H']
                    })

            if datos_grafico_grupos_efe:
                archivo = 'temp_barras_comp_grupos_efectivos.png'
                resultado = crear_grafico_barras(
                    datos_grafico_grupos_efe,
                    'Promedio de Complementos Salariales por Grupo Profesional - EFECTIVO',
                    'Promedio (€)',
                    archivo,
                    'horizontal'
                )
                if resultado is not None and os.path.exists(archivo):
                    self.archivos_temp.append(archivo)
                    doc.add_picture(archivo, width=Inches(GRAFICO_CFG['ancho_doc']))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()

        # Gráficos de barras horizontales por grupo profesional - Equiparados
        if datos_grupos_equiparados:
            datos_grafico_grupos_equi = []
            for d in datos_grupos_equiparados:
                if d['promedio_M'] > 0 or d['promedio_H'] > 0:
                    datos_grafico_grupos_equi.append({
                        'categoria': d['grupo'],
                        'M': d['promedio_M'],
                        'H': d['promedio_H'],
                        'n_M': d['n_M'],
                        'n_H': d['n_H']
                    })

            if datos_grafico_grupos_equi:
                archivo = 'temp_barras_comp_grupos_equiparados.png'
                resultado = crear_grafico_barras(
                    datos_grafico_grupos_equi,
                    'Promedio de Complementos Salariales por Grupo Profesional - EQUIPARADO',
                    'Promedio (€)',
                    archivo,
                    'horizontal'
                )
                if resultado is not None and os.path.exists(archivo):
                    self.archivos_temp.append(archivo)
                    doc.add_picture(archivo, width=Inches(GRAFICO_CFG['ancho_doc']))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()

        # Tablas de complementos salariales por agrupaciones (Nivel SVPT) y puestos de trabajo
        df_actual['combinacion'] = df_actual[COLS['nivel_svpt']].astype(str) + ' + ' + df_actual[COLS['puesto']].astype(str)
        combinaciones = sorted(df_actual['combinacion'].dropna().unique())

        # Efectivos
        datos_puestos_efectivos = []
        for comb in combinaciones:
            df_comb = df_actual[df_actual['combinacion'] == comb]

            df_comb_m = df_comb[df_comb[COLS['sexo']] == 'M']
            df_comb_h = df_comb[df_comb[COLS['sexo']] == 'H']

            col_comp_efe = COLS['comp_efectivo']

            # Promedio
            promedio_m = df_comb_m[col_comp_efe].mean() if len(df_comb_m) > 0 else 0
            promedio_h = df_comb_h[col_comp_efe].mean() if len(df_comb_h) > 0 else 0
            promedio_brecha = calcular_brecha(promedio_h, promedio_m) or 0

            # Mediana
            mediana_m = df_comb_m[col_comp_efe].median() if len(df_comb_m) > 0 else 0
            mediana_h = df_comb_h[col_comp_efe].median() if len(df_comb_h) > 0 else 0
            mediana_brecha = calcular_brecha(mediana_h, mediana_m) or 0

            # Puntos promedio
            puntos_promedio = 0
            if COLS['puntos'] in df_comb.columns and len(df_comb) > 0:
                puntos_promedio = df_comb[COLS['puntos']].mean()

            if len(df_comb_m) > 0 or len(df_comb_h) > 0:
                datos_puestos_efectivos.append({
                    'grupo': comb,
                    'n_M': len(df_comb_m),
                    'n_H': len(df_comb_h),
                    'promedio_M': promedio_m,
                    'promedio_H': promedio_h,
                    'promedio_brecha': promedio_brecha,
                    'mediana_M': mediana_m,
                    'mediana_H': mediana_h,
                    'mediana_brecha': mediana_brecha,
                    'puntos': puntos_promedio
                })

        if datos_puestos_efectivos:
            crear_tabla_complementos_por_grupo(doc, 'Complementos Salariales por Agrupaciones (Nivel SVPT) y Puestos de Trabajo - EFECTIVOS', datos_puestos_efectivos)

        # Equiparados
        datos_puestos_equiparados = []
        for comb in combinaciones:
            df_comb = df_actual[df_actual['combinacion'] == comb]

            df_comb_m = df_comb[df_comb[COLS['sexo']] == 'M']
            df_comb_h = df_comb[df_comb[COLS['sexo']] == 'H']

            col_comp_equi = COLS['comp_equiparado']

            # Promedio
            promedio_m = df_comb_m[col_comp_equi].mean() if len(df_comb_m) > 0 else 0
            promedio_h = df_comb_h[col_comp_equi].mean() if len(df_comb_h) > 0 else 0
            promedio_brecha = calcular_brecha(promedio_h, promedio_m) or 0

            # Mediana
            mediana_m = df_comb_m[col_comp_equi].median() if len(df_comb_m) > 0 else 0
            mediana_h = df_comb_h[col_comp_equi].median() if len(df_comb_h) > 0 else 0
            mediana_brecha = calcular_brecha(mediana_h, mediana_m) or 0

            # Puntos promedio
            puntos_promedio = 0
            if COLS['puntos'] in df_comb.columns and len(df_comb) > 0:
                puntos_promedio = df_comb[COLS['puntos']].mean()

            if len(df_comb_m) > 0 or len(df_comb_h) > 0:
                datos_puestos_equiparados.append({
                    'grupo': comb,
                    'n_M': len(df_comb_m),
                    'n_H': len(df_comb_h),
                    'promedio_M': promedio_m,
                    'promedio_H': promedio_h,
                    'promedio_brecha': promedio_brecha,
                    'mediana_M': mediana_m,
                    'mediana_H': mediana_h,
                    'mediana_brecha': mediana_brecha,
                    'puntos': puntos_promedio
                })

        if datos_puestos_equiparados:
            crear_tabla_complementos_por_grupo(doc, 'Complementos Salariales por Agrupaciones (Nivel SVPT) y Puestos de Trabajo - EQUIPARADOS', datos_puestos_equiparados)

        # Gráficos de barras horizontales por puesto - Efectivos
        if datos_puestos_efectivos:
            datos_grafico_puestos_efe = []
            for d in datos_puestos_efectivos:
                if d['promedio_M'] > 0 or d['promedio_H'] > 0:
                    datos_grafico_puestos_efe.append({
                        'categoria': d['grupo'],
                        'M': d['promedio_M'],
                        'H': d['promedio_H'],
                        'n_M': d['n_M'],
                        'n_H': d['n_H']
                    })

            if datos_grafico_puestos_efe:
                archivo = 'temp_barras_comp_puestos_efectivos.png'
                resultado = crear_grafico_barras(
                    datos_grafico_puestos_efe,
                    'Promedio de Complementos Salariales por Puesto de Trabajo - EFECTIVO',
                    'Promedio (€)',
                    archivo,
                    'horizontal'
                )
                if resultado is not None and os.path.exists(archivo):
                    self.archivos_temp.append(archivo)
                    doc.add_picture(archivo, width=Inches(GRAFICO_CFG['ancho_doc']))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()

        # Gráfico dual-eje por puesto - Efectivos (UN SOLO gráfico con todos los puestos)
        if datos_puestos_efectivos:
            datos_grafico_dual = []
            for d in datos_puestos_efectivos:
                if d['promedio_M'] > 0 or d['promedio_H'] > 0:
                    datos_grafico_dual.append({
                        'categoria': d['grupo'],
                        'M': d['promedio_M'],
                        'H': d['promedio_H'],
                        'puntos': d.get('puntos', 0),
                        'n_M': d.get('n_M', 0),
                        'n_H': d.get('n_H', 0)
                    })

            if datos_grafico_dual:
                archivo = 'temp_barras_dual_comp_puestos_efectivos.png'
                resultado = crear_grafico_salarios_con_puntos_condicional(
                    datos_grafico_dual,
                    'Complementos Salariales Medios por Puesto de Trabajo - EFECTIVO',
                    archivo,
                    tipo_grafico='denso'  # Muchas categorías: usar fuentes más pequeñas
                )
                if resultado is not None and os.path.exists(archivo):
                    self.archivos_temp.append(archivo)
                    doc.add_picture(archivo, width=Inches(GRAFICO_CFG['ancho_doc']))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()

        # Gráficos de barras horizontales por puesto - Equiparados
        if datos_puestos_equiparados:
            datos_grafico_puestos_equi = []
            for d in datos_puestos_equiparados:
                if d['promedio_M'] > 0 or d['promedio_H'] > 0:
                    datos_grafico_puestos_equi.append({
                        'categoria': d['grupo'],
                        'M': d['promedio_M'],
                        'H': d['promedio_H'],
                        'n_M': d['n_M'],
                        'n_H': d['n_H']
                    })

            if datos_grafico_puestos_equi:
                archivo = 'temp_barras_comp_puestos_equiparados.png'
                resultado = crear_grafico_barras(
                    datos_grafico_puestos_equi,
                    'Promedio de Complementos Salariales por Puesto de Trabajo - EQUIPARADO',
                    'Promedio (€)',
                    archivo,
                    'horizontal'
                )
                if resultado is not None and os.path.exists(archivo):
                    self.archivos_temp.append(archivo)
                    doc.add_picture(archivo, width=Inches(GRAFICO_CFG['ancho_doc']))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()

        # Gráfico dual-eje por puesto - Equiparados (UN SOLO gráfico con todos los puestos)
        if datos_puestos_equiparados:
            datos_grafico_dual = []
            for d in datos_puestos_equiparados:
                if d['promedio_M'] > 0 or d['promedio_H'] > 0:
                    datos_grafico_dual.append({
                        'categoria': d['grupo'],
                        'M': d['promedio_M'],
                        'H': d['promedio_H'],
                        'puntos': d.get('puntos', 0),
                        'n_M': d.get('n_M', 0),
                        'n_H': d.get('n_H', 0)
                    })

            if datos_grafico_dual:
                archivo = 'temp_barras_dual_comp_puestos_equiparados.png'
                resultado = crear_grafico_salarios_con_puntos_condicional(
                    datos_grafico_dual,
                    'Complementos Salariales Medios por Puesto de Trabajo - EQUIPARADO',
                    archivo,
                    tipo_grafico='denso'  # Muchas categorías: usar fuentes más pequeñas
                )
                if resultado is not None and os.path.exists(archivo):
                    self.archivos_temp.append(archivo)
                    doc.add_picture(archivo, width=Inches(GRAFICO_CFG['ancho_doc']))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()

        # Tablas de complementos salariales por nivel convenio colectivo
        niveles = sorted([str(n) for n in df_actual[COLS['nivel_convenio']].dropna().unique()])

        # Efectivos
        datos_niveles_efectivos = []
        for nivel in niveles:
            df_nivel = df_actual[df_actual[COLS['nivel_convenio']].astype(str) == nivel]

            df_nivel_m = df_nivel[df_nivel[COLS['sexo']] == 'M']
            df_nivel_h = df_nivel[df_nivel[COLS['sexo']] == 'H']

            col_comp_efe = COLS['comp_efectivo']

            # Promedio
            promedio_m = df_nivel_m[col_comp_efe].mean() if len(df_nivel_m) > 0 else 0
            promedio_h = df_nivel_h[col_comp_efe].mean() if len(df_nivel_h) > 0 else 0
            promedio_brecha = calcular_brecha(promedio_h, promedio_m) or 0

            # Mediana
            mediana_m = df_nivel_m[col_comp_efe].median() if len(df_nivel_m) > 0 else 0
            mediana_h = df_nivel_h[col_comp_efe].median() if len(df_nivel_h) > 0 else 0
            mediana_brecha = calcular_brecha(mediana_h, mediana_m) or 0

            datos_niveles_efectivos.append({
                'grupo': nivel,
                'n_M': len(df_nivel_m),
                'n_H': len(df_nivel_h),
                'promedio_M': promedio_m,
                'promedio_H': promedio_h,
                'promedio_brecha': promedio_brecha,
                'mediana_M': mediana_m,
                'mediana_H': mediana_h,
                'mediana_brecha': mediana_brecha
            })

        if datos_niveles_efectivos:
            crear_tabla_complementos_por_grupo(doc, 'Complementos Salariales por Nivel Convenio Colectivo - EFECTIVOS', datos_niveles_efectivos)

        # Equiparados
        datos_niveles_equiparados = []
        for nivel in niveles:
            df_nivel = df_actual[df_actual[COLS['nivel_convenio']].astype(str) == nivel]

            df_nivel_m = df_nivel[df_nivel[COLS['sexo']] == 'M']
            df_nivel_h = df_nivel[df_nivel[COLS['sexo']] == 'H']

            col_comp_equi = COLS['comp_equiparado']

            # Promedio
            promedio_m = df_nivel_m[col_comp_equi].mean() if len(df_nivel_m) > 0 else 0
            promedio_h = df_nivel_h[col_comp_equi].mean() if len(df_nivel_h) > 0 else 0
            promedio_brecha = calcular_brecha(promedio_h, promedio_m) or 0

            # Mediana
            mediana_m = df_nivel_m[col_comp_equi].median() if len(df_nivel_m) > 0 else 0
            mediana_h = df_nivel_h[col_comp_equi].median() if len(df_nivel_h) > 0 else 0
            mediana_brecha = calcular_brecha(mediana_h, mediana_m) or 0

            datos_niveles_equiparados.append({
                'grupo': nivel,
                'n_M': len(df_nivel_m),
                'n_H': len(df_nivel_h),
                'promedio_M': promedio_m,
                'promedio_H': promedio_h,
                'promedio_brecha': promedio_brecha,
                'mediana_M': mediana_m,
                'mediana_H': mediana_h,
                'mediana_brecha': mediana_brecha
            })

        if datos_niveles_equiparados:
            crear_tabla_complementos_por_grupo(doc, 'Complementos Salariales por Nivel Convenio Colectivo - EQUIPARADOS', datos_niveles_equiparados)

        # Gráficos de barras horizontales por nivel - Efectivos
        if datos_niveles_efectivos:
            datos_grafico_niveles_efe = []
            for d in datos_niveles_efectivos:
                if d['promedio_M'] > 0 or d['promedio_H'] > 0:
                    datos_grafico_niveles_efe.append({
                        'categoria': d['grupo'],
                        'M': d['promedio_M'],
                        'H': d['promedio_H'],
                        'n_M': d['n_M'],
                        'n_H': d['n_H']
                    })

            if datos_grafico_niveles_efe:
                archivo = 'temp_barras_comp_niveles_efectivos.png'
                resultado = crear_grafico_barras(
                    datos_grafico_niveles_efe,
                    'Promedio de Complementos Salariales por Nivel Convenio Colectivo - EFECTIVO',
                    'Promedio (€)',
                    archivo,
                    'horizontal'
                )
                if resultado is not None and os.path.exists(archivo):
                    self.archivos_temp.append(archivo)
                    doc.add_picture(archivo, width=Inches(GRAFICO_CFG['ancho_doc']))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()

        # Gráficos de barras horizontales por nivel - Equiparados
        if datos_niveles_equiparados:
            datos_grafico_niveles_equi = []
            for d in datos_niveles_equiparados:
                if d['promedio_M'] > 0 or d['promedio_H'] > 0:
                    datos_grafico_niveles_equi.append({
                        'categoria': d['grupo'],
                        'M': d['promedio_M'],
                        'H': d['promedio_H'],
                        'n_M': d['n_M'],
                        'n_H': d['n_H']
                    })

            if datos_grafico_niveles_equi:
                archivo = 'temp_barras_comp_niveles_equiparados.png'
                resultado = crear_grafico_barras(
                    datos_grafico_niveles_equi,
                    'Promedio de Complementos Salariales por Nivel Convenio Colectivo - EQUIPARADO',
                    'Promedio (€)',
                    archivo,
                    'horizontal'
                )
                if resultado is not None and os.path.exists(archivo):
                    self.archivos_temp.append(archivo)
                    doc.add_picture(archivo, width=Inches(GRAFICO_CFG['ancho_doc']))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()

        # ==================== COMPLEMENTOS EXTRASALARIALES ====================
        doc.add_heading('Complementos Extrasalariales', level=2)

        # Gráficos donut - Efectivo y Equiparado
        col_extra_efe = COLS['extra_efectivo']
        col_extra_equi = COLS['extra_equiparado']

        stats_extra_efe = calcular_estadistico(self.df, col_extra_efe, 'efectivos_sb_complementos', 'media')
        brecha_extra_efe = calcular_brecha(stats_extra_efe['H'], stats_extra_efe['M']) or 0

        stats_extra_equi = calcular_estadistico(self.df, col_extra_equi, 'equiparados_sb_complementos', 'media')
        brecha_extra_equi = calcular_brecha(stats_extra_equi['H'], stats_extra_equi['M']) or 0

        archivo_extra_efe = 'temp_donut_comp_extrasalariales_efectivo.png'
        archivo_extra_equi = 'temp_donut_comp_extrasalariales_equiparado.png'

        img_extra_efe_creada = crear_grafico_donut(
            {'M': stats_extra_efe['M'], 'H': stats_extra_efe['H'], 'brecha': brecha_extra_efe},
            'PROMEDIO Complementos Extrasalariales - EFECTIVO',
            archivo_extra_efe
        )

        img_extra_equi_creada = crear_grafico_donut(
            {'M': stats_extra_equi['M'], 'H': stats_extra_equi['H'], 'brecha': brecha_extra_equi},
            'PROMEDIO Complementos Extrasalariales - EQUIPARADO',
            archivo_extra_equi
        )

        # Insertar gráficos lado a lado
        if img_extra_efe_creada and img_extra_equi_creada:
            self.archivos_temp.extend([archivo_extra_efe, archivo_extra_equi])

            tabla_graficos = doc.add_table(rows=1, cols=2)

            # Ocultar bordes
            ocultar_bordes_tabla(tabla_graficos)

            # Añadir gráfico efectivo
            celda_izq = tabla_graficos.rows[0].cells[0]
            paragraph_izq = celda_izq.paragraphs[0]
            run_izq = paragraph_izq.add_run()
            run_izq.add_picture(archivo_extra_efe, width=Inches(GRAFICO_CFG['ancho_donut']))
            paragraph_izq.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Añadir gráfico equiparado
            celda_der = tabla_graficos.rows[0].cells[1]
            paragraph_der = celda_der.paragraphs[0]
            run_der = paragraph_der.add_run()
            run_der.add_picture(archivo_extra_equi, width=Inches(GRAFICO_CFG['ancho_donut']))
            paragraph_der.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()

        # Tabla de distribución de complementos extrasalariales
        # Buscar todas las columnas PE (formato: "PE 1 Nombre" o "PE1")
        cols_pe = [col for col in self.df.columns
                   if col.startswith('PE') and bool(re.match(r'^PE\s*\d+', col))]

        # Filtrar solo las columnas base (sin _equiparado)
        cols_pe_base = [col for col in cols_pe if '_equiparado' not in col]

        # Ordenar numéricamente por el número después de PE
        def extraer_numero_pe(col):
            try:
                # Extraer el número que sigue a PE (puede haber espacio o no)
                match = re.search(r'PE\s*(\d+)', col)
                return int(match.group(1)) if match else 999
            except:
                return 999
        cols_pe_base = sorted(cols_pe_base, key=extraer_numero_pe)

        datos_extra = []
        for col_comp in cols_pe_base:
            df_con_comp = df_actual[df_actual[col_comp] > 0]
            n_m = len(df_con_comp[df_con_comp[COLS['sexo']] == 'M'])
            n_h = len(df_con_comp[df_con_comp[COLS['sexo']] == 'H'])
            total = n_m + n_h

            if total > 0:
                # Calcular porcentaje respecto al total de personas con ese complemento
                pct_m = (n_m / total * 100) if total > 0 else 0
                pct_h = (n_h / total * 100) if total > 0 else 0

                datos_extra.append({
                    'complemento': col_comp,
                    'n_M': n_m,
                    'n_H': n_h,
                    'total': total,
                    'pct_M': pct_m,
                    'pct_H': pct_h
                })

        if datos_extra:
            crear_tabla_complementos(doc, 'Distribución de Complementos Extrasalariales', datos_extra)

        # Tabla de complementos extrasalariales por puesto de trabajo
        datos_extra_puesto = []
        for puesto in puestos:
            df_puesto = df_actual[df_actual[COLS['puesto']] == puesto]

            for col_comp in cols_pe_base:
                df_puesto_comp = df_puesto[df_puesto[col_comp] > 0]
                n_m = len(df_puesto_comp[df_puesto_comp[COLS['sexo']] == 'M'])
                n_h = len(df_puesto_comp[df_puesto_comp[COLS['sexo']] == 'H'])
                total = n_m + n_h

                if total > 0:
                    pct_m = (n_m / total * 100) if total > 0 else 0
                    pct_h = (n_h / total * 100) if total > 0 else 0

                    datos_extra_puesto.append({
                        'puesto': puesto,
                        'complemento': col_comp,
                        'n_M': n_m,
                        'n_H': n_h,
                        'total': total,
                        'pct_M': pct_m,
                        'pct_H': pct_h
                    })

        if datos_extra_puesto:
            crear_tabla_complementos_por_puesto(doc, 'Distribución de Complementos Extrasalariales por Puesto de Trabajo', datos_extra_puesto)

        # Gráficos de barras horizontales de promedio de complementos extrasalariales
        # Efectivos
        datos_grafico_extra_efectivos = []
        for col_comp in cols_pe_base:
            df_con_comp = df_actual[df_actual[col_comp] > 0]
            if len(df_con_comp) > 0:
                promedio_m = df_con_comp[df_con_comp[COLS['sexo']] == 'M'][col_comp].mean() if len(df_con_comp[df_con_comp[COLS['sexo']] == 'M']) > 0 else 0
                promedio_h = df_con_comp[df_con_comp[COLS['sexo']] == 'H'][col_comp].mean() if len(df_con_comp[df_con_comp[COLS['sexo']] == 'H']) > 0 else 0

                if promedio_m > 0 or promedio_h > 0:
                    datos_grafico_extra_efectivos.append({
                        'categoria': col_comp,
                        'M': promedio_m,
                        'H': promedio_h,
                        'n_M': len(df_con_comp[df_con_comp[COLS['sexo']] == 'M']),
                        'n_H': len(df_con_comp[df_con_comp[COLS['sexo']] == 'H'])
                    })

        # Dividir en bloques de 20 complementos
        if datos_grafico_extra_efectivos:
            tamanio_bloque = 20
            for idx in range(0, len(datos_grafico_extra_efectivos), tamanio_bloque):
                bloque = datos_grafico_extra_efectivos[idx:idx + tamanio_bloque]
                es_primer_bloque = (idx == 0)

                archivo = f'temp_barras_comp_extrasalariales_efectivos_{idx}.png'
                resultado = crear_grafico_barras(
                    bloque,
                    'Complementos Promedio Efectivos por Complemento Extrasalarial * Sexo',
                    'Promedio (€)',
                    archivo,
                    'horizontal',
                    mostrar_titulo=es_primer_bloque,
                    mostrar_leyenda=es_primer_bloque
                )
                if resultado is not None and os.path.exists(archivo):
                    self.archivos_temp.append(archivo)
                    doc.add_picture(archivo, width=Inches(GRAFICO_CFG['ancho_doc']))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()

        # Equiparados
        datos_grafico_extra_equiparados = []
        for col_comp in cols_pe_base:
            # Buscar la columna equiparada correspondiente
            col_comp_equi = f"{col_comp}_equiparado"
            if col_comp_equi in self.df.columns:
                df_con_comp = df_actual[df_actual[col_comp_equi] > 0]
                if len(df_con_comp) > 0:
                    promedio_m = df_con_comp[df_con_comp[COLS['sexo']] == 'M'][col_comp_equi].mean() if len(df_con_comp[df_con_comp[COLS['sexo']] == 'M']) > 0 else 0
                    promedio_h = df_con_comp[df_con_comp[COLS['sexo']] == 'H'][col_comp_equi].mean() if len(df_con_comp[df_con_comp[COLS['sexo']] == 'H']) > 0 else 0

                    if promedio_m > 0 or promedio_h > 0:
                        datos_grafico_extra_equiparados.append({
                            'categoria': col_comp,
                            'M': promedio_m,
                            'H': promedio_h,
                            'n_M': len(df_con_comp[df_con_comp[COLS['sexo']] == 'M']),
                            'n_H': len(df_con_comp[df_con_comp[COLS['sexo']] == 'H'])
                        })

        # Dividir en bloques de 20 complementos
        if datos_grafico_extra_equiparados:
            tamanio_bloque = 20
            for idx in range(0, len(datos_grafico_extra_equiparados), tamanio_bloque):
                bloque = datos_grafico_extra_equiparados[idx:idx + tamanio_bloque]
                es_primer_bloque = (idx == 0)

                archivo = f'temp_barras_comp_extrasalariales_equiparados_{idx}.png'
                resultado = crear_grafico_barras(
                    bloque,
                    'Complementos Promedio Equiparados por Complemento Extrasalarial * Sexo',
                    'Promedio (€)',
                    archivo,
                    'horizontal',
                    mostrar_titulo=es_primer_bloque,
                    mostrar_leyenda=es_primer_bloque
                )
                if resultado is not None and os.path.exists(archivo):
                    self.archivos_temp.append(archivo)
                    doc.add_picture(archivo, width=Inches(GRAFICO_CFG['ancho_doc']))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()

        # Tablas de complementos extrasalariales por grupo profesional
        # Efectivos
        datos_grupos_extra_efectivos = []
        for grupo in grupos_prof:
            df_grupo = df_actual[df_actual[COLS['grupo_prof']].astype(str) == grupo]

            df_grupo_m = df_grupo[df_grupo[COLS['sexo']] == 'M']
            df_grupo_h = df_grupo[df_grupo[COLS['sexo']] == 'H']

            col_extra_efe = COLS['extra_efectivo']

            # Promedio
            promedio_m = df_grupo_m[col_extra_efe].mean() if len(df_grupo_m) > 0 else 0
            promedio_h = df_grupo_h[col_extra_efe].mean() if len(df_grupo_h) > 0 else 0
            promedio_brecha = calcular_brecha(promedio_h, promedio_m) or 0

            # Mediana
            mediana_m = df_grupo_m[col_extra_efe].median() if len(df_grupo_m) > 0 else 0
            mediana_h = df_grupo_h[col_extra_efe].median() if len(df_grupo_h) > 0 else 0
            mediana_brecha = calcular_brecha(mediana_h, mediana_m) or 0

            datos_grupos_extra_efectivos.append({
                'grupo': grupo,
                'n_M': len(df_grupo_m),
                'n_H': len(df_grupo_h),
                'promedio_M': promedio_m,
                'promedio_H': promedio_h,
                'promedio_brecha': promedio_brecha,
                'mediana_M': mediana_m,
                'mediana_H': mediana_h,
                'mediana_brecha': mediana_brecha
            })

        if datos_grupos_extra_efectivos:
            crear_tabla_complementos_por_grupo(doc, 'Complementos Extrasalariales por Grupos Profesionales - EFECTIVOS', datos_grupos_extra_efectivos)

        # Equiparados
        datos_grupos_extra_equiparados = []
        for grupo in grupos_prof:
            df_grupo = df_actual[df_actual[COLS['grupo_prof']].astype(str) == grupo]

            df_grupo_m = df_grupo[df_grupo[COLS['sexo']] == 'M']
            df_grupo_h = df_grupo[df_grupo[COLS['sexo']] == 'H']

            col_extra_equi = COLS['extra_equiparado']

            # Promedio
            promedio_m = df_grupo_m[col_extra_equi].mean() if len(df_grupo_m) > 0 else 0
            promedio_h = df_grupo_h[col_extra_equi].mean() if len(df_grupo_h) > 0 else 0
            promedio_brecha = calcular_brecha(promedio_h, promedio_m) or 0

            # Mediana
            mediana_m = df_grupo_m[col_extra_equi].median() if len(df_grupo_m) > 0 else 0
            mediana_h = df_grupo_h[col_extra_equi].median() if len(df_grupo_h) > 0 else 0
            mediana_brecha = calcular_brecha(mediana_h, mediana_m) or 0

            datos_grupos_extra_equiparados.append({
                'grupo': grupo,
                'n_M': len(df_grupo_m),
                'n_H': len(df_grupo_h),
                'promedio_M': promedio_m,
                'promedio_H': promedio_h,
                'promedio_brecha': promedio_brecha,
                'mediana_M': mediana_m,
                'mediana_H': mediana_h,
                'mediana_brecha': mediana_brecha
            })

        if datos_grupos_extra_equiparados:
            crear_tabla_complementos_por_grupo(doc, 'Complementos Extrasalariales por Grupos Profesionales - EQUIPARADOS', datos_grupos_extra_equiparados)

        # Gráficos de barras horizontales por grupo profesional - Efectivos
        if datos_grupos_extra_efectivos:
            datos_grafico_grupos_extra_efe = []
            for d in datos_grupos_extra_efectivos:
                if d['promedio_M'] > 0 or d['promedio_H'] > 0:
                    datos_grafico_grupos_extra_efe.append({
                        'categoria': d['grupo'],
                        'M': d['promedio_M'],
                        'H': d['promedio_H'],
                        'n_M': d['n_M'],
                        'n_H': d['n_H']
                    })

            if datos_grafico_grupos_extra_efe:
                archivo = 'temp_barras_comp_extra_grupos_efectivos.png'
                resultado = crear_grafico_barras(
                    datos_grafico_grupos_extra_efe,
                    'Promedio de Complementos Extrasalariales por Grupo Profesional - EFECTIVO',
                    'Promedio (€)',
                    archivo,
                    'horizontal'
                )
                if resultado is not None and os.path.exists(archivo):
                    self.archivos_temp.append(archivo)
                    doc.add_picture(archivo, width=Inches(GRAFICO_CFG['ancho_doc']))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()

        # Gráficos de barras horizontales por grupo profesional - Equiparados
        if datos_grupos_extra_equiparados:
            datos_grafico_grupos_extra_equi = []
            for d in datos_grupos_extra_equiparados:
                if d['promedio_M'] > 0 or d['promedio_H'] > 0:
                    datos_grafico_grupos_extra_equi.append({
                        'categoria': d['grupo'],
                        'M': d['promedio_M'],
                        'H': d['promedio_H'],
                        'n_M': d['n_M'],
                        'n_H': d['n_H']
                    })

            if datos_grafico_grupos_extra_equi:
                archivo = 'temp_barras_comp_extra_grupos_equiparados.png'
                resultado = crear_grafico_barras(
                    datos_grafico_grupos_extra_equi,
                    'Promedio de Complementos Extrasalariales por Grupo Profesional - EQUIPARADO',
                    'Promedio (€)',
                    archivo,
                    'horizontal'
                )
                if resultado is not None and os.path.exists(archivo):
                    self.archivos_temp.append(archivo)
                    doc.add_picture(archivo, width=Inches(GRAFICO_CFG['ancho_doc']))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()

        # Tablas de complementos extrasalariales por agrupaciones (Nivel SVPT) y puestos de trabajo
        # Efectivos
        datos_puestos_extra_efectivos = []
        for comb in combinaciones:
            df_comb = df_actual[df_actual['combinacion'] == comb]

            df_comb_m = df_comb[df_comb[COLS['sexo']] == 'M']
            df_comb_h = df_comb[df_comb[COLS['sexo']] == 'H']

            col_extra_efe = COLS['extra_efectivo']

            # Promedio
            promedio_m = df_comb_m[col_extra_efe].mean() if len(df_comb_m) > 0 else 0
            promedio_h = df_comb_h[col_extra_efe].mean() if len(df_comb_h) > 0 else 0
            promedio_brecha = calcular_brecha(promedio_h, promedio_m) or 0

            # Mediana
            mediana_m = df_comb_m[col_extra_efe].median() if len(df_comb_m) > 0 else 0
            mediana_h = df_comb_h[col_extra_efe].median() if len(df_comb_h) > 0 else 0
            mediana_brecha = calcular_brecha(mediana_h, mediana_m) or 0

            # Puntos promedio
            puntos_promedio = 0
            if COLS['puntos'] in df_comb.columns and len(df_comb) > 0:
                puntos_promedio = df_comb[COLS['puntos']].mean()

            if len(df_comb_m) > 0 or len(df_comb_h) > 0:
                datos_puestos_extra_efectivos.append({
                    'grupo': comb,
                    'n_M': len(df_comb_m),
                    'n_H': len(df_comb_h),
                    'promedio_M': promedio_m,
                    'promedio_H': promedio_h,
                    'promedio_brecha': promedio_brecha,
                    'mediana_M': mediana_m,
                    'mediana_H': mediana_h,
                    'mediana_brecha': mediana_brecha,
                    'puntos': puntos_promedio
                })

        if datos_puestos_extra_efectivos:
            crear_tabla_complementos_por_grupo(doc, 'Complementos Extrasalariales por Agrupaciones (Nivel SVPT) y Puestos de Trabajo - EFECTIVOS', datos_puestos_extra_efectivos)

        # Equiparados
        datos_puestos_extra_equiparados = []
        for comb in combinaciones:
            df_comb = df_actual[df_actual['combinacion'] == comb]

            df_comb_m = df_comb[df_comb[COLS['sexo']] == 'M']
            df_comb_h = df_comb[df_comb[COLS['sexo']] == 'H']

            col_extra_equi = COLS['extra_equiparado']

            # Promedio
            promedio_m = df_comb_m[col_extra_equi].mean() if len(df_comb_m) > 0 else 0
            promedio_h = df_comb_h[col_extra_equi].mean() if len(df_comb_h) > 0 else 0
            promedio_brecha = calcular_brecha(promedio_h, promedio_m) or 0

            # Mediana
            mediana_m = df_comb_m[col_extra_equi].median() if len(df_comb_m) > 0 else 0
            mediana_h = df_comb_h[col_extra_equi].median() if len(df_comb_h) > 0 else 0
            mediana_brecha = calcular_brecha(mediana_h, mediana_m) or 0

            # Puntos promedio
            puntos_promedio = 0
            if COLS['puntos'] in df_comb.columns and len(df_comb) > 0:
                puntos_promedio = df_comb[COLS['puntos']].mean()

            if len(df_comb_m) > 0 or len(df_comb_h) > 0:
                datos_puestos_extra_equiparados.append({
                    'grupo': comb,
                    'n_M': len(df_comb_m),
                    'n_H': len(df_comb_h),
                    'promedio_M': promedio_m,
                    'promedio_H': promedio_h,
                    'promedio_brecha': promedio_brecha,
                    'mediana_M': mediana_m,
                    'mediana_H': mediana_h,
                    'mediana_brecha': mediana_brecha,
                    'puntos': puntos_promedio
                })

        if datos_puestos_extra_equiparados:
            crear_tabla_complementos_por_grupo(doc, 'Complementos Extrasalariales por Agrupaciones (Nivel SVPT) y Puestos de Trabajo - EQUIPARADOS', datos_puestos_extra_equiparados)

        # Gráficos de barras horizontales por puesto - Efectivos
        if datos_puestos_extra_efectivos:
            datos_grafico_puestos_extra_efe = []
            for d in datos_puestos_extra_efectivos:
                if d['promedio_M'] > 0 or d['promedio_H'] > 0:
                    datos_grafico_puestos_extra_efe.append({
                        'categoria': d['grupo'],
                        'M': d['promedio_M'],
                        'H': d['promedio_H'],
                        'n_M': d['n_M'],
                        'n_H': d['n_H']
                    })

            if datos_grafico_puestos_extra_efe:
                archivo = 'temp_barras_comp_extra_puestos_efectivos.png'
                resultado = crear_grafico_barras(
                    datos_grafico_puestos_extra_efe,
                    'Promedio de Complementos Extrasalariales por Puesto de Trabajo - EFECTIVO',
                    'Promedio (€)',
                    archivo,
                    'horizontal'
                )
                if resultado is not None and os.path.exists(archivo):
                    self.archivos_temp.append(archivo)
                    doc.add_picture(archivo, width=Inches(GRAFICO_CFG['ancho_doc']))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()

        # Gráfico dual-eje por puesto - Efectivos (UN SOLO gráfico con todos los puestos)
        if datos_puestos_extra_efectivos:
            datos_grafico_dual = []
            for d in datos_puestos_extra_efectivos:
                if d['promedio_M'] > 0 or d['promedio_H'] > 0:
                    datos_grafico_dual.append({
                        'categoria': d['grupo'],
                        'M': d['promedio_M'],
                        'H': d['promedio_H'],
                        'puntos': d.get('puntos', 0),
                        'n_M': d.get('n_M', 0),
                        'n_H': d.get('n_H', 0)
                    })

            if datos_grafico_dual:
                archivo = 'temp_barras_dual_comp_extra_puestos_efectivos.png'
                resultado = crear_grafico_salarios_con_puntos_condicional(
                    datos_grafico_dual,
                    'Complementos Extrasalariales Medios por Puesto de Trabajo - EFECTIVO',
                    archivo,
                    tipo_grafico='denso'  # Muchas categorías: usar fuentes más pequeñas
                )
                if resultado is not None and os.path.exists(archivo):
                    self.archivos_temp.append(archivo)
                    doc.add_picture(archivo, width=Inches(GRAFICO_CFG['ancho_doc']))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()

        # Gráficos de barras horizontales por puesto - Equiparados
        if datos_puestos_extra_equiparados:
            datos_grafico_puestos_extra_equi = []
            for d in datos_puestos_extra_equiparados:
                if d['promedio_M'] > 0 or d['promedio_H'] > 0:
                    datos_grafico_puestos_extra_equi.append({
                        'categoria': d['grupo'],
                        'M': d['promedio_M'],
                        'H': d['promedio_H'],
                        'n_M': d['n_M'],
                        'n_H': d['n_H']
                    })

            if datos_grafico_puestos_extra_equi:
                archivo = 'temp_barras_comp_extra_puestos_equiparados.png'
                resultado = crear_grafico_barras(
                    datos_grafico_puestos_extra_equi,
                    'Promedio de Complementos Extrasalariales por Puesto de Trabajo - EQUIPARADO',
                    'Promedio (€)',
                    archivo,
                    'horizontal'
                )
                if resultado is not None and os.path.exists(archivo):
                    self.archivos_temp.append(archivo)
                    doc.add_picture(archivo, width=Inches(GRAFICO_CFG['ancho_doc']))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()

        # Gráfico dual-eje por puesto - Equiparados (UN SOLO gráfico con todos los puestos)
        if datos_puestos_extra_equiparados:
            datos_grafico_dual = []
            for d in datos_puestos_extra_equiparados:
                if d['promedio_M'] > 0 or d['promedio_H'] > 0:
                    datos_grafico_dual.append({
                        'categoria': d['grupo'],
                        'M': d['promedio_M'],
                        'H': d['promedio_H'],
                        'puntos': d.get('puntos', 0),
                        'n_M': d.get('n_M', 0),
                        'n_H': d.get('n_H', 0)
                    })

            if datos_grafico_dual:
                archivo = 'temp_barras_dual_comp_extra_puestos_equiparados.png'
                resultado = crear_grafico_salarios_con_puntos_condicional(
                    datos_grafico_dual,
                    'Complementos Extrasalariales Medios por Puesto de Trabajo - EQUIPARADO',
                    archivo,
                    tipo_grafico='denso'  # Muchas categorías: usar fuentes más pequeñas
                )
                if resultado is not None and os.path.exists(archivo):
                    self.archivos_temp.append(archivo)
                    doc.add_picture(archivo, width=Inches(GRAFICO_CFG['ancho_doc']))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()

        # Tablas de complementos extrasalariales por nivel convenio colectivo
        # Efectivos
        datos_niveles_extra_efectivos = []
        for nivel in niveles:
            df_nivel = df_actual[df_actual[COLS['nivel_convenio']].astype(str) == nivel]

            df_nivel_m = df_nivel[df_nivel[COLS['sexo']] == 'M']
            df_nivel_h = df_nivel[df_nivel[COLS['sexo']] == 'H']

            col_extra_efe = COLS['extra_efectivo']

            # Promedio
            promedio_m = df_nivel_m[col_extra_efe].mean() if len(df_nivel_m) > 0 else 0
            promedio_h = df_nivel_h[col_extra_efe].mean() if len(df_nivel_h) > 0 else 0
            promedio_brecha = calcular_brecha(promedio_h, promedio_m) or 0

            # Mediana
            mediana_m = df_nivel_m[col_extra_efe].median() if len(df_nivel_m) > 0 else 0
            mediana_h = df_nivel_h[col_extra_efe].median() if len(df_nivel_h) > 0 else 0
            mediana_brecha = calcular_brecha(mediana_h, mediana_m) or 0

            datos_niveles_extra_efectivos.append({
                'grupo': nivel,
                'n_M': len(df_nivel_m),
                'n_H': len(df_nivel_h),
                'promedio_M': promedio_m,
                'promedio_H': promedio_h,
                'promedio_brecha': promedio_brecha,
                'mediana_M': mediana_m,
                'mediana_H': mediana_h,
                'mediana_brecha': mediana_brecha
            })

        if datos_niveles_extra_efectivos:
            crear_tabla_complementos_por_grupo(doc, 'Complementos Extrasalariales por Nivel Convenio Colectivo - EFECTIVOS', datos_niveles_extra_efectivos)

        # Equiparados
        datos_niveles_extra_equiparados = []
        for nivel in niveles:
            df_nivel = df_actual[df_actual[COLS['nivel_convenio']].astype(str) == nivel]

            df_nivel_m = df_nivel[df_nivel[COLS['sexo']] == 'M']
            df_nivel_h = df_nivel[df_nivel[COLS['sexo']] == 'H']

            col_extra_equi = COLS['extra_equiparado']

            # Promedio
            promedio_m = df_nivel_m[col_extra_equi].mean() if len(df_nivel_m) > 0 else 0
            promedio_h = df_nivel_h[col_extra_equi].mean() if len(df_nivel_h) > 0 else 0
            promedio_brecha = calcular_brecha(promedio_h, promedio_m) or 0

            # Mediana
            mediana_m = df_nivel_m[col_extra_equi].median() if len(df_nivel_m) > 0 else 0
            mediana_h = df_nivel_h[col_extra_equi].median() if len(df_nivel_h) > 0 else 0
            mediana_brecha = calcular_brecha(mediana_h, mediana_m) or 0

            datos_niveles_extra_equiparados.append({
                'grupo': nivel,
                'n_M': len(df_nivel_m),
                'n_H': len(df_nivel_h),
                'promedio_M': promedio_m,
                'promedio_H': promedio_h,
                'promedio_brecha': promedio_brecha,
                'mediana_M': mediana_m,
                'mediana_H': mediana_h,
                'mediana_brecha': mediana_brecha
            })

        if datos_niveles_extra_equiparados:
            crear_tabla_complementos_por_grupo(doc, 'Complementos Extrasalariales por Nivel Convenio Colectivo - EQUIPARADOS', datos_niveles_extra_equiparados)

        # Gráficos de barras horizontales por nivel - Efectivos
        if datos_niveles_extra_efectivos:
            datos_grafico_niveles_extra_efe = []
            for d in datos_niveles_extra_efectivos:
                if d['promedio_M'] > 0 or d['promedio_H'] > 0:
                    datos_grafico_niveles_extra_efe.append({
                        'categoria': d['grupo'],
                        'M': d['promedio_M'],
                        'H': d['promedio_H'],
                        'n_M': d['n_M'],
                        'n_H': d['n_H']
                    })

            if datos_grafico_niveles_extra_efe:
                archivo = 'temp_barras_comp_extra_niveles_efectivos.png'
                resultado = crear_grafico_barras(
                    datos_grafico_niveles_extra_efe,
                    'Promedio de Complementos Extrasalariales por Nivel Convenio Colectivo - EFECTIVO',
                    'Promedio (€)',
                    archivo,
                    'horizontal'
                )
                if resultado is not None and os.path.exists(archivo):
                    self.archivos_temp.append(archivo)
                    doc.add_picture(archivo, width=Inches(GRAFICO_CFG['ancho_doc']))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()

        # Gráficos de barras horizontales por nivel - Equiparados
        if datos_niveles_extra_equiparados:
            datos_grafico_niveles_extra_equi = []
            for d in datos_niveles_extra_equiparados:
                if d['promedio_M'] > 0 or d['promedio_H'] > 0:
                    datos_grafico_niveles_extra_equi.append({
                        'categoria': d['grupo'],
                        'M': d['promedio_M'],
                        'H': d['promedio_H'],
                        'n_M': d['n_M'],
                        'n_H': d['n_H']
                    })

            if datos_grafico_niveles_extra_equi:
                archivo = 'temp_barras_comp_extra_niveles_equiparados.png'
                resultado = crear_grafico_barras(
                    datos_grafico_niveles_extra_equi,
                    'Promedio de Complementos Extrasalariales por Nivel Convenio Colectivo - EQUIPARADO',
                    'Promedio (€)',
                    archivo,
                    'horizontal'
                )
                if resultado is not None and os.path.exists(archivo):
                    self.archivos_temp.append(archivo)
                    doc.add_picture(archivo, width=Inches(GRAFICO_CFG['ancho_doc']))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()

    def _guardar_documento(self, doc, tipo_informe):
        """Guarda el documento Word con timestamp"""
        ruta_base = Path(__file__).parent.parent
        carpeta_salida = ruta_base / '05_INFORMES'
        carpeta_salida.mkdir(exist_ok=True)

        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        nombre_archivo = f'registro_retributivo_{timestamp}_{tipo_informe}.docx'
        self.ruta_salida = carpeta_salida / nombre_archivo

        doc.save(str(self.ruta_salida))

    def _limpiar_temporales(self):
        """Elimina archivos temporales de gráficos"""
        for archivo in self.archivos_temp:
            try:
                if os.path.exists(archivo):
                    os.remove(archivo)
            except:
                pass
        self.archivos_temp = []


# ==================== FUNCIÓN PRINCIPAL ====================

def main():
    """Función principal"""
    print("=" * 60, flush=True)
    print("GENERADOR DE INFORMES DE REGISTRO RETRIBUTIVO", flush=True)
    print("=" * 60, flush=True)
    print(flush=True)

    # Crear generador
    generador = GeneradorInformeOptimizado()

    # Cargar datos
    if not generador.cargar_datos():
        return

    # Si hay argumentos de línea de comandos, usarlos
    import sys
    if len(sys.argv) > 2:
        # Segundo argumento es el tipo de informe
        tipo_arg = sys.argv[2].upper()
        if tipo_arg in ['CONSOLIDADO', 'PROMEDIO', 'MEDIA', 'MEDIANA', 'COMPLEMENTOS', '1', '2', '3', '4']:
            tipos = {
                '1': 'PROMEDIO',
                '2': 'MEDIANA',
                '3': 'COMPLEMENTOS',
                '4': 'CONSOLIDADO',
                'MEDIA': 'PROMEDIO'  # Alias para compatibilidad
            }
            tipo_informe = tipos.get(tipo_arg, tipo_arg)
        else:
            tipo_informe = 'CONSOLIDADO'
    else:
        # Modo interactivo
        print(flush=True)
        print("Seleccione el tipo de informe a generar:", flush=True)
        print("1. PROMEDIO (solo análisis con promedios)", flush=True)
        print("2. MEDIANA (solo análisis con medianas)", flush=True)
        print("3. COMPLEMENTOS (solo análisis de complementos)", flush=True)
        print("4. CONSOLIDADO (incluye PROMEDIO + MEDIANA + COMPLEMENTOS)", flush=True)
        print(flush=True)

        try:
            opcion = input("Opción (1/2/3/4) [Enter para CONSOLIDADO]: ").strip()
            if not opcion:
                opcion = '4'
        except EOFError:
            opcion = '4'

        tipos = {
            '1': 'PROMEDIO',
            '2': 'MEDIANA',
            '3': 'COMPLEMENTOS',
            '4': 'CONSOLIDADO'
        }

        tipo_informe = tipos.get(opcion, 'CONSOLIDADO')

    print(flush=True)
    log(f"Generando informe tipo: {tipo_informe}")
    print(flush=True)

    # Generar informe
    if generador.generar_informe(tipo_informe):
        print(flush=True)
        print("=" * 60, flush=True)
        log(f"INFORME GENERADO EXITOSAMENTE", 'OK')
        print(f"Ubicación: {generador.ruta_salida}", flush=True)
        print("=" * 60, flush=True)
    else:
        print(flush=True)
        log("Error al generar el informe", 'ERROR')


if __name__ == "__main__":
    main()
